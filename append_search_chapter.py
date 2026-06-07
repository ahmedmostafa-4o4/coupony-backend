import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_search_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== SEARCH SYSTEM FEATURE ==================
    doc.add_heading('1.1.14 Global Search System (Meilisearch)', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Global Search System provides a unified, high-performance search experience powered by Laravel Scout and the Meilisearch engine. "
        "It aggregates results across multiple core entities (Products, Stores, and Categories) through a single unified endpoint. "
        "The system supports fast autocomplete suggestions, personal user search history, and global trending search terms. "
        "To ensure lightning-fast response times, search queries are executed purely against Meilisearch indexes in memory, "
        "while search history logging (for analytics and trending algorithms) is offloaded to asynchronous background Queued Jobs."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-108: Unified Global Search - User submits a query. The system queries Meilisearch for Products, Stores, and Categories simultaneously, returning grouped results. The query is asynchronously logged for analytics.",
        "UC-109: Autocomplete Suggestions - As the user types, the system fetches lightweight, limit-constrained suggestions from Meilisearch for immediate UI rendering.",
        "UC-110: View Trending Searches - User views the top 10 globally trending search terms, calculated from the aggregated search_history table.",
        "UC-111: View Personal Search History - Authenticated user retrieves their chronological history of past search terms.",
        "UC-112: Clear Search History - Authenticated user deletes their personal search history records from the database."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-108: Unified Global Search ---
    add_diagram(doc, "Unified Global Search (UC-108)", """sequenceDiagram
    participant User
    participant GSC as GlobalSearchController
    participant GSS as GlobalSearchService
    participant MS as Meilisearch (Scout)
    participant Q as Laravel Queue
    participant DB as Database

    User->>GSC: GET /search?q=iphone&types[]=products,stores
    GSC->>GSS: search("iphone", types, User)
    par Meilisearch Queries
        GSS->>MS: Query Product Index
        MS-->>GSS: Product Hits
        GSS->>MS: Query Store Index
        MS-->>GSS: Store Hits
    end
    GSS->>Q: Dispatch RecordSearchTerm Job
    GSS-->>GSC: Grouped Results Array
    GSC-->>User: 200 OK (Unified JSON)
    
    Note over Q,DB: Asynchronous Background Process
    Q->>DB: Insert into search_history (user_id, term, ip)
""", "sd_unified_search", diagrams, 6.0)

    # --- UC-109: Autocomplete Suggestions ---
    add_diagram(doc, "Autocomplete Suggestions (UC-109)", """sequenceDiagram
    participant User
    participant GSC as GlobalSearchController
    participant GSS as GlobalSearchService
    participant MS as Meilisearch (Scout)

    User->>GSC: GET /search/suggestions?q=iph
    GSC->>GSS: suggestions("iph")
    GSS->>MS: Query Multiple Indexes with take(5)
    MS-->>GSS: Lightweight Hits (Titles/Names only)
    GSS-->>GSC: Flattened Suggestions Array
    GSC-->>User: 200 OK (Suggestions JSON)
""", "sd_autocomplete", diagrams, 5.0)

    # --- UC-110: View Trending Searches ---
    add_diagram(doc, "View Trending Searches (UC-110)", """sequenceDiagram
    participant User
    participant GSC as GlobalSearchController
    participant DB as Database

    User->>GSC: GET /search/trending
    GSC->>DB: Query search_history table
    GSC->>DB: Group by `term`, Order by COUNT(*) DESC
    GSC->>DB: Limit 10
    DB-->>GSC: Top 10 Terms
    GSC-->>User: 200 OK (Trending Terms Array)
""", "sd_trending_searches", diagrams)

    # --- UC-111 & UC-112: Personal Search History ---
    add_diagram(doc, "Personal Search History Management (UC-111, UC-112)", """sequenceDiagram
    participant User
    participant GSC as GlobalSearchController
    participant DB as Database

    alt UC-111: View History
        User->>GSC: GET /search/history
        GSC->>DB: Query search_history where user_id = User.id
        GSC->>DB: Order by created_at DESC, limit 20
        DB-->>GSC: History Records
        GSC-->>User: 200 OK (History Array)
    else UC-112: Clear History
        User->>GSC: DELETE /search/history
        GSC->>DB: DELETE FROM search_history WHERE user_id = User.id
        DB-->>GSC: Rows deleted
        GSC-->>User: 200 OK
    end
""", "sd_search_history", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "Search indexing is managed externally by Meilisearch. The relational database is only used to track analytics and history for trending algorithms."
    )
    schemas = [
        "search_history: Analytics table tracking queries. Fields: id (PK), user_id (UUID, FK to users, nullable for guests), term (string), ip_address (string), results_count (integer, default 0), created_at."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ SEARCH_HISTORY : "generates"
    
    USERS {
        char id PK "UUID"
        string email
    }
    SEARCH_HISTORY {
        bigint id PK
        char user_id FK "Nullable"
        string term
        string ip_address
        int results_count
        timestamp created_at
    }
    
    %% Note: Products, Stores, and Categories are synced to Meilisearch outside standard relational links.
"""
    er_img_path = "search_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Search System chapter appended to {filename}")

if __name__ == "__main__":
    append_search_chapter()
