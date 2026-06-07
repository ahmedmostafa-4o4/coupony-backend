import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_address_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== ADDRESS SYSTEM FEATURE ==================
    doc.add_heading('1.1.12 Address Management System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Address Management System utilizes a highly reusable polymorphic architecture. "
        "Instead of tying addresses directly to a specific user or store table, the core 'addresses' table holds pure geographical "
        "and textual data (linked to strict 'cities' and 'regions' taxonomies). "
        "A polymorphic pivot table named 'addressables' maps these addresses to ANY entity (User, Store, etc.) while "
        "storing contextual flags such as 'is_default', 'is_billing', or 'is_shipping'. "
        "This allows seamless location sharing and uniform geographical querying across the entire platform. "
        "Dedicated domain action classes ensure transactional safety when creating or toggling default statuses."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-92: List User Addresses - Authenticated user retrieves a list of all their saved personal addresses, eager-loaded with City and Region data.",
        "UC-93: Create User Address - User creates a new personal address (home, work, other). The system creates the address and safely links it via the polymorphic pivot.",
        "UC-94: Set Default User Address - User marks a specific address as their primary. The system atomically removes the default flag from all other addresses and sets it on the target.",
        "UC-95: Delete User Address - User removes an address. The system soft-deletes the core record and removes the polymorphic link.",
        "UC-96: List Store Addresses - Store Owner retrieves all operational or billing addresses associated with their store.",
        "UC-97: Create Store Address - Store Owner adds a physical storefront or headquarters address to their store.",
        "UC-98: Admin Manage Store Addresses - Administrators have global override capabilities to list, create, update, or delete physical addresses for any store on the platform."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-93: Create User Address ---
    add_diagram(doc, "Create User Address (UC-93)", """sequenceDiagram
    participant User
    participant UAC as UserAddressController
    participant CUA as CreateUserAddress
    participant DB as Database

    User->>UAC: POST /me/addresses (city_id, address_line1, is_default)
    UAC->>CUA: execute(User, AddressData)
    CUA->>DB: Begin Transaction
    CUA->>DB: Insert into addresses table
    alt is_default == true
        CUA->>DB: UPDATE addressables SET is_default = false WHERE addressable_id = User.id
    end
    CUA->>DB: Insert into addressables (address_id, addressable_id, addressable_type='User', is_default)
    DB-->>CUA: Commit Transaction
    CUA-->>UAC: Address Model
    UAC-->>User: 201 Created
""", "sd_create_user_address", diagrams, 5.5)

    # --- UC-94: Set Default User Address ---
    add_diagram(doc, "Set Default User Address (UC-94)", """sequenceDiagram
    participant User
    participant UAC as UserAddressController
    participant SDUA as SetDefaultUserAddress
    participant DB as Database

    User->>UAC: POST /me/addresses/{address}/default
    UAC->>UAC: Authorize Policy (User owns Address)
    UAC->>SDUA: execute(User, Address)
    SDUA->>DB: Begin Transaction
    SDUA->>DB: UPDATE addressables SET is_default = false WHERE addressable_id = User.id
    SDUA->>DB: UPDATE addressables SET is_default = true WHERE address_id = Target.id
    DB-->>SDUA: Commit Transaction
    SDUA-->>UAC: Success
    UAC-->>User: 200 OK
""", "sd_set_default_address", diagrams, 5.0)

    # --- UC-97: Create Store Address ---
    add_diagram(doc, "Create Store Address (UC-97)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as StoreAddressController
    participant CSA as CreateStoreAddress
    participant DB as Database

    Owner->>SAC: POST /stores/{store}/addresses (type='storefront')
    SAC->>SAC: Gate Check (Can Manage Store)
    SAC->>CSA: execute(Store, AddressData)
    CSA->>DB: Begin Transaction
    CSA->>DB: Insert into addresses table
    alt is_default == true
        CSA->>DB: UPDATE addressables SET is_default = false WHERE addressable_id = Store.id
    end
    CSA->>DB: Insert into addressables (address_id, addressable_id, addressable_type='Store')
    DB-->>CSA: Commit Transaction
    CSA-->>SAC: Address Model
    SAC-->>Owner: 201 Created
""", "sd_create_store_address", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The schema relies on a polymorphic pivot to prevent duplicating address structures for different entities, "
        "while enforcing a strict geographic hierarchy via cities and regions."
    )
    schemas = [
        "cities: Root geographic nodes. Fields: id, name_en, name_ar, is_active.",
        "regions: Sub-areas within a city. Fields: id, city_id (FK), name_en, name_ar, postal_code, is_active.",
        "addresses: Core location data. Fields: id (UUID), title, type (home/work/other/billing/shipping/storefront), city_id (FK), region_id (FK), formatted_address, address_line1, address_line2, postal_code, latitude, longitude.",
        "addressables: Polymorphic pivot mapping addresses to owners. Fields: id, address_id (FK), addressable_id (Morph ID), addressable_type (Morph Type), is_default, is_billing, is_shipping."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    CITIES ||--o{ REGIONS : "contains"
    CITIES ||--o{ ADDRESSES : "located_in"
    REGIONS ||--o{ ADDRESSES : "located_in"
    ADDRESSES ||--o{ ADDRESSABLES : "mapped_via"
    USERS ||--o{ ADDRESSABLES : "morphs_to (addressable)"
    STORES ||--o{ ADDRESSABLES : "morphs_to (addressable)"

    CITIES {
        bigint id PK
        string name_en
        string name_ar
    }
    REGIONS {
        bigint id PK
        bigint city_id FK
        string name_en
        string name_ar
    }
    ADDRESSES {
        uuid id PK
        bigint city_id FK
        bigint region_id FK
        string type
        string address_line1
        decimal latitude
        decimal longitude
    }
    ADDRESSABLES {
        bigint id PK
        uuid address_id FK
        char addressable_id "UUID"
        string addressable_type "Class Name"
        bool is_default
    }
"""
    er_img_path = "address_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Address System chapter appended to {filename}")

if __name__ == "__main__":
    append_address_chapter()
