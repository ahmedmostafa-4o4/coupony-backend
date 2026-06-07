import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_following_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== FOLLOWING SYSTEM FEATURE ==================
    doc.add_heading('1.1.4 Following System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Following System allows users to curate a personalized experience by following their favorite stores. "
        "When a user follows a store, a record is created in the pivot table 'store_followers' along with a timestamp and a notification preference. "
        "Simultaneously, the 'followers_count' on the store record is updated to provide quick analytics. "
        "Users can toggle notifications for specific stores and view a paginated list of all stores they follow, which can be filtered by category. "
        "Store owners have access to a dedicated dashboard that aggregates recent followers across all stores they manage. "
        "A critical feature of this system is the 'Following Feed', which aggregates recent activities (such as new products or offers) from all stores the authenticated user follows into a single personalized timeline."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-25: Follow Store - User follows an active store, enabling notifications by default and incrementing the store's follower count.",
        "UC-26: Unfollow Store - User unfollows a store, safely decreasing the cached follower count.",
        "UC-27: Toggle Follow Notifications - User toggles push/email notifications for updates from a specific followed store.",
        "UC-28: View Followed Stores - User views a list of all stores they currently follow, optionally filtered by store category.",
        "UC-29: View Store Followers - Users or admins retrieve a paginated list of all users following a specific store.",
        "UC-30: Store Owner Followers Dashboard - A store owner views an aggregated list of new followers across all stores they own.",
        "UC-31: Following Feed - User requests a personalized feed timeline containing activities from only the stores they follow."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Follow Store Flow
    doc.add_heading('Sequence Diagram: Follow Store Flow', level=5)
    follow_mermaid = """sequenceDiagram
    participant User
    participant SFC as StoreFollowController
    participant FSA as FollowStore Action
    participant DB as Database
    
    User->>SFC: POST /stores/{store}/follow
    SFC->>FSA: execute(Store, User)
    FSA->>DB: Insert into store_followers (user_id, store_id)
    FSA->>DB: Increment Store followers_count
    FSA-->>SFC: Follow Record
    SFC-->>User: 201 Created (Store followers_count, is_following=true)
"""
    follow_img_path = "follow_store_diagram.png"
    download_image(generate_kroki_url(follow_mermaid), follow_img_path)
    doc.add_picture(follow_img_path, width=Inches(5.0))
    diagrams.append(follow_img_path)

    # 2. Store Owner Followers Dashboard Flow
    doc.add_heading('Sequence Diagram: Store Owner Followers Dashboard Flow', level=5)
    owner_followers_mermaid = """sequenceDiagram
    participant Owner as Store Owner
    participant MFC as MyFollowersController
    participant DB as Database
    
    Owner->>MFC: GET /me/followers/new
    MFC->>DB: Get all store_ids owned by Owner
    MFC->>DB: Query Users INNER JOIN store_followers 
    MFC->>DB: Filter where store_id IN (owned stores)
    MFC->>DB: Order by store_followers.followed_at DESC
    DB-->>MFC: Paginated Users Collection
    MFC-->>Owner: 200 OK (List of new followers with profiles)
"""
    owner_followers_img_path = "owner_followers_diagram.png"
    download_image(generate_kroki_url(owner_followers_mermaid), owner_followers_img_path)
    doc.add_picture(owner_followers_img_path, width=Inches(5.5))
    diagrams.append(owner_followers_img_path)

    # 3. Following Feed Flow
    doc.add_heading('Sequence Diagram: Following Feed Flow', level=5)
    feed_mermaid = """sequenceDiagram
    participant User
    participant FFC as FollowingFeedController
    participant GFA as GetFollowingFeedAction
    participant DB as Database
    
    User->>FFC: GET /feed/following
    FFC->>GFA: execute(User)
    GFA->>DB: Fetch IDs of stores followed by User
    GFA->>DB: Query Activities (Products, Offers) 
    GFA->>DB: Filter where store_id IN (followed IDs)
    GFA->>DB: Order by created_at DESC
    DB-->>GFA: Aggregated Feed Data
    GFA-->>FFC: Formatted Feed
    FFC-->>User: 200 OK (Paginated timeline)
"""
    feed_img_path = "following_feed_diagram.png"
    download_image(generate_kroki_url(feed_mermaid), feed_img_path)
    doc.add_picture(feed_img_path, width=Inches(5.5))
    diagrams.append(feed_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The Following system primarily uses a pivot table to establish a many-to-many relationship between Users and Stores, while storing metadata about the follow action."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "store_followers: Pivot table linking users and stores. Fields: id (BigInt), user_id (UUID, FK to users), store_id (UUID, FK to stores), notification_enabled (Boolean, default true), followed_at (Timestamp). Unique constraint on [user_id, store_id].",
        "stores: The target entity. Contains a cached integer column 'followers_count' which is incremented/decremented to avoid heavy COUNT() queries."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ STORE_FOLLOWERS : "follows"
    STORES ||--o{ STORE_FOLLOWERS : "is_followed_by"
    
    USERS {
        char id PK "UUID"
        string email
    }
    STORES {
        uuid id PK
        string name
        int followers_count "Cached Total"
    }
    STORE_FOLLOWERS {
        bigint id PK
        uuid user_id FK
        uuid store_id FK
        boolean notification_enabled
        timestamp followed_at
    }
"""
    er_img_path = "following_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Following System chapter appended to {filename}")

if __name__ == "__main__":
    append_following_chapter()
