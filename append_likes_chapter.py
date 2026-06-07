import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_likes_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== LIKES SYSTEM FEATURE ==================
    doc.add_heading('1.1.6 Likes System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Likes System is a community engagement feature distinct from Favorites. While Favorites act as a personal save list, Likes serve as a public endorsement metric for Products, Comments, and Banners. "
        "When a user likes a product or a comment, the system verifies the target is active, approved, and visible. It then records the interaction in a specific pivot table (e.g., 'product_likes' or 'product_comment_likes'). "
        "To ensure high performance during read operations (such as listing products in a feed), target entities cache the total 'likes_count'. "
        "Users can also retrieve a paginated list of all products they have liked. The system automatically handles duplicate like attempts by using 'firstOrCreate' logic or throwing specific validation exceptions."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-36: Like Product - User likes an active and approved product, inserting a pivot record and incrementing the product's 'likes_count'.",
        "UC-37: Unlike Product - User removes their like from a product, deleting the pivot record and decrementing the cached count.",
        "UC-38: View Liked Products - User retrieves a paginated list of all products they have previously liked.",
        "UC-39: Like/Unlike Comment - User likes or unlikes a specific product or store comment, tracking engagement at the community discussion level.",
        "UC-40: Like Banner - User likes a promotional banner, helping administrators gauge the popularity of specific marketing campaigns."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Like Product Flow
    doc.add_heading('Sequence Diagram: Like Product Flow', level=5)
    like_product_mermaid = """sequenceDiagram
    participant User
    participant PLC as ProductLikeController
    participant LPA as LikeProduct Action
    participant DB as Database
    
    User->>PLC: POST /products/{product}/like
    PLC->>PLC: Check if Product is ACTIVE & APPROVED
    alt Valid Product
        PLC->>LPA: execute(Product, User)
        LPA->>DB: Insert into product_likes (user_id, product_id)
        LPA->>DB: Increment Product likes_count
        PLC-->>User: 200 OK (is_liked = true, likes_count)
    else Invalid Product
        PLC-->>User: 404 Not Found
    end
"""
    like_product_img_path = "like_product_diagram.png"
    download_image(generate_kroki_url(like_product_mermaid), like_product_img_path)
    doc.add_picture(like_product_img_path, width=Inches(5.5))
    diagrams.append(like_product_img_path)

    # 2. Like Comment Flow
    doc.add_heading('Sequence Diagram: Like Comment Flow', level=5)
    like_comment_mermaid = """sequenceDiagram
    participant User
    participant PCLC as ProductCommentLikeController
    participant DB as Database
    
    User->>PCLC: POST /comments/{comment}/like
    PCLC->>PCLC: Check if Comment is visible & Product is ACTIVE
    alt Valid Comment
        PCLC->>DB: firstOrCreate in product_comment_likes
        PCLC->>DB: Count current likes for comment
        PCLC-->>User: 200 OK (is_liked = true, likes_count)
    else Invalid Comment
        PCLC-->>User: 404 Not Found
    end
"""
    like_comment_img_path = "like_comment_diagram.png"
    download_image(generate_kroki_url(like_comment_mermaid), like_comment_img_path)
    doc.add_picture(like_comment_img_path, width=Inches(5.5))
    diagrams.append(like_comment_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The Likes system mirrors the polymorphic-like pivot table structure of the Favorites system but applies it to community engagement elements, including comments."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "product_likes: Pivot table for product likes. Fields: id, user_id (FK to users), product_id (FK to products), timestamps. Unique constraint on [user_id, product_id].",
        "product_comment_likes: Pivot table for comment likes. Fields: id, user_id (FK to users), comment_id (FK to product_comments), timestamps.",
        "products: Contains a 'likes_count' (unsigned integer, default 0) to quickly sort or display popularity.",
        "product_comments: The target entity for community discussions, allowing dynamic calculation or caching of comment likes."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ PRODUCT_LIKES : "likes"
    PRODUCTS ||--o{ PRODUCT_LIKES : "receives"
    USERS ||--o{ PRODUCT_COMMENT_LIKES : "likes"
    PRODUCT_COMMENTS ||--o{ PRODUCT_COMMENT_LIKES : "receives"
    
    USERS {
        char id PK "UUID"
        string email
    }
    PRODUCTS {
        bigint id PK
        string name
        int likes_count "Cached Total"
    }
    PRODUCT_LIKES {
        bigint id PK
        uuid user_id FK
        bigint product_id FK
        timestamp created_at
    }
    PRODUCT_COMMENTS {
        bigint id PK
        bigint product_id FK
        char user_id FK
        text body
    }
    PRODUCT_COMMENT_LIKES {
        bigint id PK
        uuid user_id FK
        bigint comment_id FK
    }
"""
    er_img_path = "likes_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Likes System chapter appended to {filename}")

if __name__ == "__main__":
    append_likes_chapter()
