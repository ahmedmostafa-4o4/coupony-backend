import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_following_feed_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== FOLLOWING FEED SYSTEM FEATURE ==================
    doc.add_heading('1.1.13 Following Feed System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Following Feed System governs how users interact with stores and consume personalized content. "
        "At its core, it allows users to follow and unfollow stores, and explicitly toggle push notification preferences per store. "
        "For content consumption, the system orchestrates a 'FollowingFeedService' which uses a sophisticated fallback algorithm: "
        "it primarily fetches content from stores the user explicitly follows; if insufficient, it falls back to 'recommended' content based on location (lat/lng), "
        "and finally to 'trending' content. "
        "For store owners (sellers), the system provides dedicated dashboard endpoints to view followers for a specific store, "
        "or an aggregated list of 'new followers' spanning across ALL stores owned by the seller."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-101: Follow a Store - Authenticated user follows an ACTIVE store. System creates the relation, increments followers_count, and enables notifications by default.",
        "UC-102: Unfollow a Store - Authenticated user unfollows a store. System deletes the relation and decrements followers_count.",
        "UC-103: Toggle Follow Notifications - User toggles push notification preferences for a specific store they follow.",
        "UC-104: List Followed Stores - User retrieves a paginated list of all stores they are currently following, with optional filtering by category.",
        "UC-105: List Store Followers - Store Owner/Admin views a paginated list of users following a specific store, ordered by follow date.",
        "UC-106: Get Following Feed - User requests their personalized content feed. The backend utilizes location data and a fallback algorithm (Followed -> Recommended -> Trending) to populate the feed.",
        "UC-107: Get Seller's New Followers - A multi-store Seller views an aggregated, paginated timeline of new followers across ALL stores they currently own."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-101: Follow a Store ---
    add_diagram(doc, "Follow a Store (UC-101)", """sequenceDiagram
    participant User
    participant SFC as StoreFollowController
    participant FS as FollowStore Action
    participant DB as Database

    User->>SFC: POST /stores/{store}/follow
    SFC->>SFC: Validate Store is ACTIVE
    SFC->>FS: execute(Store, User)
    FS->>DB: INSERT INTO store_followers (user_id, store_id)
    Note over DB: Sets notification_enabled = true
    FS->>DB: Increment Store.followers_count
    FS-->>SFC: Follow Result (wasRecentlyCreated)
    SFC-->>User: 201 Created (Follow Status)
""", "sd_follow_store", diagrams)

    # --- UC-102: Unfollow a Store ---
    add_diagram(doc, "Unfollow a Store (UC-102)", """sequenceDiagram
    participant User
    participant SFC as StoreFollowController
    participant US as UnfollowStore Action
    participant DB as Database

    User->>SFC: DELETE /stores/{store}/follow
    SFC->>US: execute(Store, User)
    US->>DB: DELETE FROM store_followers
    alt Record existed
        US->>DB: Decrement Store.followers_count
    end
    US-->>SFC: Success boolean
    SFC-->>User: 200 OK (Unfollowed Status)
""", "sd_unfollow_store", diagrams)

    # --- UC-103: Toggle Follow Notifications ---
    add_diagram(doc, "Toggle Follow Notifications (UC-103)", """sequenceDiagram
    participant User
    participant SFC as StoreFollowController
    participant TN as ToggleStoreFollowNotification
    participant DB as Database

    User->>SFC: PATCH /stores/{store}/follow/notifications
    SFC->>TN: execute(Store, User)
    TN->>DB: Toggle notification_enabled in store_followers
    TN-->>SFC: Updated pivot record
    SFC-->>User: 200 OK (notification_enabled state)
""", "sd_toggle_notifications", diagrams)

    # --- UC-104: List Followed Stores ---
    add_diagram(doc, "List Followed Stores (UC-104)", """sequenceDiagram
    participant User
    participant SFC as StoreFollowController
    participant DB as Database

    User->>SFC: GET /me/followed-stores?category_id=5
    SFC->>DB: Query User->followedStores()
    SFC->>DB: Filter status = ACTIVE
    alt Category provided
        SFC->>DB: Filter by Category ID via WhereHas
    end
    SFC->>DB: Eager load relations (categories, addresses, hours)
    SFC->>DB: Order by Pivot followed_at DESC
    DB-->>SFC: Paginated Stores Collection
    SFC-->>User: 200 OK (FollowedStoreResource)
""", "sd_list_followed_stores", diagrams, 5.5)

    # --- UC-105: List Store Followers ---
    add_diagram(doc, "List Store Followers (UC-105)", """sequenceDiagram
    participant Owner as Store Owner
    participant SFC as StoreFollowController
    participant DB as Database

    Owner->>SFC: GET /stores/{store}/followers
    SFC->>SFC: Validate Store is ACTIVE
    SFC->>DB: Query Store->followerUsers()
    SFC->>DB: Eager load User.profile (for avatar/name)
    SFC->>DB: Order by Pivot followed_at DESC
    DB-->>SFC: Paginated Users Collection
    SFC-->>Owner: 200 OK (StoreFollowerResource)
""", "sd_list_store_followers", diagrams)

    # --- UC-106: Get Following Feed ---
    add_diagram(doc, "Get Following Feed (UC-106)", """sequenceDiagram
    participant User
    participant FFC as FollowingFeedController
    participant GFFA as GetFollowingFeedAction
    participant FFS as FollowingFeedService
    participant DB as Database

    User->>FFC: GET /feed/following?latitude=X&longitude=Y
    FFC->>GFFA: execute(Filters, User)
    GFFA->>FFS: getFeedItems(User, page, perPage, lat, lng)
    FFS->>DB: Query posts/offers from Followed Stores
    alt Results < threshold
        FFS->>DB: Fallback to Recommended Content (using Lat/Lng)
    end
    alt Still < threshold
        FFS->>DB: Fallback to General Trending Content
    end
    DB-->>FFS: Assembled Feed Items
    FFS-->>GFFA: Paginator
    GFFA-->>FFC: Formatted Array
    FFC-->>User: 200 OK (FollowingFeedItemResource)
""", "sd_get_following_feed", diagrams, 6.0)

    # --- UC-107: Get Seller's New Followers ---
    add_diagram(doc, "Get Seller's New Followers (UC-107)", """sequenceDiagram
    participant Seller
    participant MFC as MyFollowersController
    participant DB as Database

    Seller->>MFC: GET /me/followers/new
    MFC->>DB: Fetch all Store IDs owned by Seller
    MFC->>DB: Join `users` table with `store_followers`
    MFC->>DB: Where `store_id` IN (Seller's Store IDs)
    MFC->>DB: Select user data + Pivot followed_at
    MFC->>DB: Order by Pivot followed_at DESC
    DB-->>MFC: Paginated Cross-Store Follower Collection
    MFC->>MFC: Map Pivot data dynamically
    MFC-->>Seller: 200 OK (NewFollowerResource)
""", "sd_get_sellers_followers", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The following logic is driven by a simple but highly indexed pivot table tracking user-to-store relationships."
    )
    schemas = [
        "store_followers: Pivot table. Fields: id, user_id (UUID, FK to users), store_id (UUID, FK to stores), notification_enabled (boolean, default true), followed_at (timestamp)."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ STORE_FOLLOWERS : "follows"
    STORES ||--o{ STORE_FOLLOWERS : "is_followed_by"

    USERS {
        char id PK "UUID"
        string email
    }
    STORES {
        uuid id PK
        string name
        string status
        int followers_count
    }
    STORE_FOLLOWERS {
        bigint id PK
        char user_id FK
        uuid store_id FK
        bool notification_enabled
        timestamp followed_at
    }
"""
    er_img_path = "following_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Following Feed System chapter appended to {filename}")

if __name__ == "__main__":
    append_following_feed_chapter()
