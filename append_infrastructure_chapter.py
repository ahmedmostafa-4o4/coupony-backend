import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Architecture Diagram: {title}', level=3)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_infrastructure_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== CHAPTER 2: INFRASTRUCTURE ==================
    doc.add_page_break()
    doc.add_heading('Chapter 2: Infrastructure & Background Processing', level=1)
    
    intro_text = (
        "While Chapter 1 covers the user-facing features and API logic, Chapter 2 outlines the backend infrastructure that ensures the platform remains highly performant, scalable, and automated. "
        "The architecture relies heavily on Redis for both caching and asynchronous queue processing, managed by Laravel Horizon. "
        "Global database standards ensure data integrity via UUID routing and Soft Deletes, while AWS S3 provides secure cloud storage for media assets."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 2.1 Scheduled Tasks ------------------
    doc.add_heading('2.1 Scheduled Tasks (Cron Jobs)', level=2)
    doc.add_paragraph("The Laravel Scheduler (defined in routes/console.php) acts as the heartbeat of the platform, automating critical state changes without human intervention:")
    tasks = [
        "subscriptions:check-expiring (Runs Daily): Checks for store subscriptions expiring in exactly 3 days and dispatches renewal notifications.",
        "offers:expire-flash (Runs Every Minute): Scans the `product_offers` table for items where `ends_at` has passed and aggressively sets their status to EXPIRED.",
        "banners:archive-expired (Runs Hourly): Automatically downgrades active promotional banners to ARCHIVED once their marketing window closes.",
        "analytics:calculate-trending (Runs Daily at 01:00 AM): A heavy computational job that recalculates the global `trending_score` for all active offers based on the proprietary formula (favorites + views + discounts).",
        "notifications:prune (Runs Weekly): Acts as a database janitor, permanently deleting read notifications older than 30 days to optimize storage."
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    # ------------------ 2.2 Queue Workers ------------------
    doc.add_heading('2.2 Queue Workers & Asynchronous Processing', level=2)
    doc.add_paragraph("To guarantee API response times under 200ms, the platform offloads heavy computational and network I/O tasks to Redis-backed background workers monitored by Laravel Horizon:")
    jobs = [
        "ProcessImageUploadJob: Handles background resizing, WebP compression, and watermark injection for store logos and product images immediately after S3 upload.",
        "SendBulkNotificationJob: When admins send mass broadcasts, this job chunks user collections into batches of 500 to prevent memory exhaustion, executing FCM HTTP requests asynchronously.",
        "GenerateMonthlyAnalyticsReportJob: Runs automatically on the 1st of every month to compile complex store analytics into PDF format and emails them to store owners."
    ]
    for j in jobs:
        doc.add_paragraph(j, style='List Bullet')

    # ------------------ 2.3 Caching Strategy ------------------
    doc.add_heading('2.3 Global Caching Strategy', level=2)
    doc.add_paragraph("The platform aggressively caches high-read, low-write data using Redis to eliminate database bottlenecking on high-traffic endpoints:")
    caches = [
        "Active Banners (TTL: 15 minutes): Caches `customer:banners:active-ids:v1` to ensure the app's home screen loads instantly.",
        "Explore Categories (TTL: 24 hours): Caches the static list of featured categories.",
        "Store Profiles (TTL: Forever): Basic store configuration (contact info, policies) is cached indefinitely and only explicitly invalidated when a store owner updates their profile.",
        "Search Suggestions (TTL: 1 hour): Caches popular and trending search autocomplete keywords."
    ]
    for c in caches:
        doc.add_paragraph(c, style='List Bullet')

    # ------------------ 2.4 Database & Storage ------------------
    doc.add_heading('2.4 Database & Storage Standards', level=2)
    db_standards = [
        "Primary Keys (UUID v4): Almost all core models (User, Store, Product, Banner, Notification) utilize the `HasUuids` trait. This obscures database size from competitors and ensures safe distributed scaling.",
        "Data Retention (Soft Deletes): The `SoftDeletes` trait is standard across the platform. Products, Stores, Offers, and Conversations are never hard-deleted immediately, preserving historical analytics and allowing data recovery.",
        "Cloud Storage (AWS S3): Media assets are securely stored in S3 buckets for production, seamlessly falling back to local storage during development."
    ]
    for d in db_standards:
        doc.add_paragraph(d, style='List Bullet')

    # ======================== ARCHITECTURE DIAGRAMS ========================
    doc.add_heading('System Architecture Diagrams', level=2)

    # --- Async Architecture Diagram ---
    add_diagram(doc, "Asynchronous Queue & Scheduler Architecture", """sequenceDiagram
    participant OS as Server OS (Cron)
    participant Kernel as Laravel Console Kernel
    participant API as API Controllers
    participant Redis as Redis Cache & Queues
    participant Workers as Horizon Queue Workers
    participant S3 as AWS S3 Storage
    participant FCM as Firebase FCM

    Note over OS,Kernel: 1. Scheduled Tasks
    OS->>Kernel: * * * * * (Every Minute)
    Kernel->>Kernel: Check Schedule List
    Kernel->>Database: offers:expire-flash (UPDATE status)
    
    Note over API,Workers: 2. Background Queue Processing
    API->>S3: Upload raw image
    API->>Redis: Dispatch ProcessImageUploadJob
    API-->>Client: 200 OK (Instant Response)
    
    Redis->>Workers: Pop Job
    Workers->>S3: Download, Resize (WebP), Upload
    
    API->>Redis: Dispatch SendBulkNotificationJob (10,000 users)
    Redis->>Workers: Pop Job (Batch 1 of 500)
    Workers->>FCM: Async HTTP Push Delivery
""", "sd_infra_async", diagrams, 6.0)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Infrastructure Chapter appended to {filename}")

if __name__ == "__main__":
    append_infrastructure_chapter()
