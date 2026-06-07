import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_banners_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # Remove old inaccurate chapter
    found = False
    for p in list(doc.paragraphs):
        if p.text == '1.1.16 Banners Management System':
            found = True
        if found:
            p._element.getparent().remove(p._element)
            
    diagrams = []

    # ================== BANNERS SYSTEM FEATURE ==================
    doc.add_heading('1.1.16 Banners & Claims Management System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Banners Management System is a complex promotional engine linking graphical assets to specific Product Offers and physical Store Branches. "
        "It features two distinct banner types: Regular Banners (requested by Store Owners and reviewed by Admins) and Travel Banners (managed entirely by Admins linking directly to products). "
        "When an Admin approves a pending Store Banner request, the system seamlessly auto-approves all linked underlying products and revisions to ensure campaign readiness. "
        "For Customers, banners are heavily cached for performance but maintain dynamic interaction states (likes, favorites). "
        "The most critical feature is the 'Claim' system. Customers can claim a banner, which generates a secure QR Code token and a frozen 'claim_snapshot' "
        "(a JSON copy of the banner and offer at the exact time of claim) to protect the consumer against subsequent malicious price/detail changes by the merchant prior to physical redemption."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-115: Manage Store Banners - Store Owners create promotional banner requests linking specific product offers and physical branch addresses. Banners start in a 'pending' state.",
        "UC-116: Admin Review Banner - Administrators approve or reject Store Banner requests. Approval automatically cascades to approve any pending underlying product revisions.",
        "UC-117: Manage Travel Banners - Administrators exclusively manage 'Travel' banners, linking them directly to active travel products with custom 'save_percent' labels.",
        "UC-118: View Active Banners - Customers retrieve a cached, high-performance feed of active banners. The service appends dynamic user-context flags (is_liked, is_favorited).",
        "UC-119: Interact with Banner - Customers can 'like', 'favorite', or 'share' banners, providing social proof and behavioral analytics.",
        "UC-120: Claim Banner Offer - Customer commits to an offer. The system generates secure redemption tokens (QR codes) and saves a JSON snapshot of the offer to protect against future modifications."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-115 & UC-116: Store Banner Lifecycle ---
    add_diagram(doc, "Store Banner Lifecycle (UC-115, UC-116)", """sequenceDiagram
    participant Owner as Store Owner
    participant Admin
    participant BC as Controllers
    participant BS as BannerService
    participant DB as Database

    Note over Owner,DB: Phase 1: Store Requests Banner
    Owner->>BC: POST /stores/{store}/banners (Image, Offers, Branches)
    BC->>BS: createStoreBanner(data)
    BS->>DB: INSERT INTO banners (status = 'pending')
    BS->>DB: Attach to banner_offers & banner_branches
    BC-->>Owner: 201 Created (Pending)

    Note over Admin,DB: Phase 2: Admin Approves
    Admin->>BC: POST /admin/banners/{id}/approve
    BC->>BS: approveBanner(Banner)
    BS->>DB: UPDATE banners SET status='approved', is_active=1
    loop Every Linked Offer
        BS->>DB: Auto-approve pending product revisions
        BS->>DB: Make product and offer ACTIVE
    end
    BC-->>Admin: 200 OK
""", "sd_banner_lifecycle", diagrams, 6.0)

    # --- UC-118: View Cached Banners ---
    add_diagram(doc, "View Active Banners (UC-118)", """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant BS as BannerService
    participant Cache as Redis Cache
    participant DB as Database

    Customer->>CBC: GET /banners
    CBC->>BS: getActiveBanners(User)
    BS->>Cache: Check 'customer:banners:active-ids:v1'
    alt Cache Miss
        BS->>DB: Query Approved, Active, Unexpired Banners
        BS->>DB: Eager load valid offers & branches
        DB-->>BS: Banner IDs
        BS->>Cache: Store IDs (15 min TTL)
    end
    BS->>DB: Fetch models from IDs
    BS->>DB: Map contextual flags (is_liked, is_favorited) for User
    BS-->>CBC: Hydrated Banners
    CBC-->>Customer: 200 OK
""", "sd_view_cached_banners", diagrams, 5.5)

    # --- UC-120: Claim Banner Offer ---
    add_diagram(doc, "Claim Banner Offer (UC-120)", """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant BS as BannerService
    participant DB as Database

    Customer->>CBC: POST /banners/{id}/claim
    CBC->>BS: createClaim(Banner, User)
    BS->>DB: Verify Banner is Active & Unexpired
    BS->>BS: Generate claim_token (Random String)
    BS->>BS: Generate qr_code_token (Random String)
    BS->>BS: Build claim_snapshot (JSON serialize Banner + Offer)
    BS->>BS: Calculate expires_at (min of Banner or Offer expiry)
    
    BS->>DB: INSERT INTO banner_claims
    DB-->>BS: Claim Record
    BS-->>CBC: Claim details & QR code
    CBC-->>Customer: 201 Created (Claim Token & QR)
""", "sd_claim_banner", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The schema isolates regular store workflows from direct admin promotional structures (travel), and features robust snapshot tracking for claims."
    )
    schemas = [
        "banners: Core store request. Fields: id, store_id, image_url, discount_label, priority, status (pending/approved/rejected), approved_at, end_time.",
        "travel_banners: Admin-direct banners. Links directly to a `product_id` rather than a generic offer. Includes `save_percent`.",
        "banner_offers & banner_branches: Pivot tables linking banners to their target promotional items and physical redemption locations.",
        "banner_claims: Tracks user claims. Fields: status (active/redeemed/expired), claim_token, qr_code_token, claim_snapshot (JSON text), expires_at, redeemed_at.",
        "Interaction Tables: `banner_likes`, `banner_favorites`, `banner_shares`."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    BANNERS ||--o{ BANNER_OFFERS : "promotes"
    BANNERS ||--o{ BANNER_BRANCHES : "redeemable_at"
    PRODUCT_OFFERS ||--o{ BANNER_OFFERS : "linked_via"
    ADDRESSES ||--o{ BANNER_BRANCHES : "linked_via"
    
    BANNERS ||--o{ BANNER_CLAIMS : "generates"
    USERS ||--o{ BANNER_CLAIMS : "owns"
    
    TRAVEL_BANNERS ||--|| PRODUCTS : "promotes_directly"

    BANNERS {
        uuid id PK
        uuid store_id FK
        string status
        timestamp end_time
    }
    BANNER_CLAIMS {
        uuid id PK
        uuid banner_id FK
        char user_id FK
        string status
        json claim_snapshot
        string qr_code_token
    }
    TRAVEL_BANNERS {
        uuid id PK
        uuid product_id FK
        decimal save_percent
    }
"""
    er_img_path = "banners_er_diagram_fixed.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Banners System chapter fixed and appended to {filename}")

if __name__ == "__main__":
    fix_banners_chapter()
