import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_favorites_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== FAVORITES SYSTEM FEATURE ==================
    doc.add_heading('1.1.5 Favorites System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Favorites System allows users to save specific items, primarily Products and promotional Banners, for quick access later. "
        "When a user favorites a product, the system ensures the product is currently active and approved. It then creates a record in the 'product_favorites' pivot table. "
        "To optimize database performance for high-traffic read operations, the system leverages a cached 'favorites_count' column on the target entity (e.g., products, banners) which increments or decrements in sync with user actions. "
        "Users can easily retrieve a paginated list of their favorite products, ordered by the most recently added. "
        "This system drives engagement metrics across the platform, allowing store owners to see how many saves their products or banners have received."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-32: Favorite Product - User marks an active and approved product as a favorite, generating a pivot record and incrementing the product's cached 'favorites_count'.",
        "UC-33: Unfavorite Product - User removes a product from their favorites, securely deleting the pivot record and decrementing the cached count.",
        "UC-34: View Favorite Products - User retrieves a paginated list of all products they have favorited, ordered by the timestamp they were added.",
        "UC-35: Favorite Banner/Offer - User favorites a promotional banner or specific store offer, tracking cross-entity engagement."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Favorite Product Flow
    doc.add_heading('Sequence Diagram: Favorite Product Flow', level=5)
    favorite_mermaid = """sequenceDiagram
    participant User
    participant PFC as ProductFavoriteController
    participant FPA as FavoriteProduct Action
    participant DB as Database
    
    User->>PFC: POST /products/{product}/favorite
    PFC->>PFC: Check if Product is ACTIVE & APPROVED
    alt Valid Product
        PFC->>FPA: execute(Product, User)
        FPA->>DB: Insert into product_favorites (user_id, product_id)
        FPA->>DB: Increment Product favorites_count
        PFC-->>User: 200 OK (is_favorited = true)
    else Invalid Product
        PFC-->>User: 404 Not Found
    end
"""
    favorite_img_path = "favorite_product_diagram.png"
    download_image(generate_kroki_url(favorite_mermaid), favorite_img_path)
    doc.add_picture(favorite_img_path, width=Inches(5.0))
    diagrams.append(favorite_img_path)

    # 2. View Favorite Products Flow
    doc.add_heading('Sequence Diagram: View Favorite Products Flow', level=5)
    view_favorites_mermaid = """sequenceDiagram
    participant User
    participant PFC as ProductFavoriteController
    participant PR as ProductRepository
    participant DB as Database
    
    User->>PFC: GET /me/favorites/products
    PFC->>PR: favoriteProductsPaginate(User)
    PR->>DB: Query Products INNER JOIN product_favorites
    PR->>DB: Filter where user_id = User.id
    PR->>DB: Order by product_favorites.created_at DESC
    DB-->>PR: Paginated Products Collection
    PR-->>PFC: Formatted Collection
    PFC-->>User: 200 OK (Paginated Favorites)
"""
    view_favorites_img_path = "view_favorites_diagram.png"
    download_image(generate_kroki_url(view_favorites_mermaid), view_favorites_img_path)
    doc.add_picture(view_favorites_img_path, width=Inches(5.5))
    diagrams.append(view_favorites_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The Favorites system operates via polymorphic-like relationships using dedicated pivot tables for each entity type to ensure strict referential integrity."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "product_favorites: Pivot table tracking user favorites for products. Fields: id (BigInt), user_id (UUID, FK to users), product_id (BigInt, FK to products), timestamps. Contains a unique compound index on [user_id, product_id].",
        "banner_favorites: Pivot table tracking user favorites for banners. Fields: id (BigInt), user_id (UUID, FK to users), banner_id (BigInt, FK to banners), timestamps.",
        "products: The main product table containing a 'favorites_count' (unsigned integer, default 0) to optimize analytic queries.",
        "banners: The promotional banner table, also tracking 'favorites_count'."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ PRODUCT_FAVORITES : "likes"
    PRODUCTS ||--o{ PRODUCT_FAVORITES : "has_saves"
    USERS ||--o{ BANNER_FAVORITES : "likes"
    BANNERS ||--o{ BANNER_FAVORITES : "has_saves"
    
    USERS {
        char id PK "UUID"
        string email
    }
    PRODUCTS {
        bigint id PK
        string name
        int favorites_count "Cached Total"
    }
    PRODUCT_FAVORITES {
        bigint id PK
        uuid user_id FK
        bigint product_id FK
        timestamp created_at
    }
    BANNERS {
        bigint id PK
        string title
        int favorites_count "Cached Total"
    }
    BANNER_FAVORITES {
        bigint id PK
        uuid user_id FK
        bigint banner_id FK
    }
"""
    er_img_path = "favorites_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Favorites System chapter appended to {filename}")

if __name__ == "__main__":
    append_favorites_chapter()
