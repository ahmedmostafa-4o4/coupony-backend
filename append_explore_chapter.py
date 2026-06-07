import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_explore_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== EXPLORE SYSTEM FEATURE ==================
    doc.add_heading('1.1.15 Explore System (Home Dashboard)', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Explore System serves as the highly customized, data-intensive heart of the application's home screen. "
        "It is engineered to provide a personalized, rich content discovery experience by aggregating complex data through "
        "a 'Bootstrap' endpoint and providing infinite-scroll capabilities through a 'Picks' endpoint. "
        "The Bootstrap action orchestrates parallel queries across multiple domains to fetch personalized 'Interests' (based on user onboarding profiles), "
        "active 'Activities' (store categories), 'Trending' offers (calculated via a dynamic multi-factor formula including views, discounts, and favorites), "
        "time-sensitive 'Flash' offers (expiring within 24 hours), 'Top Stores' (highest rated with best coupons), and 'Nearby' offers using spatial Haversine algorithms. "
        "The Picks action delegates to a dedicated ProductRecommendationService to handle heavy pagination, filtering, and sorting of 'Picked for You' items."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-113: Get Explore Bootstrap - Client requests the main dashboard data. The system executes GetExploreBootstrapAction which aggregates interests, activities, trending, flash, top_stores, and nearby data into a single unified JSON response to hydrate the entire home screen instantly.",
        "UC-114: Get Picked Offers (Picks) - User scrolls down the explore page triggering infinite scroll. The system executes GetExplorePicksAction to return a paginated, filterable (by category, min discount), and sortable list of personalized recommended products."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-113: Get Explore Bootstrap ---
    add_diagram(doc, "Get Explore Bootstrap (UC-113)", """sequenceDiagram
    participant Client
    participant EC as ExploreController
    participant GEBA as GetExploreBootstrapAction
    participant ES as ExploreService
    participant DB as Database

    Client->>EC: GET /explore/bootstrap?lat=X&lng=Y
    EC->>GEBA: execute(Request, User)
    GEBA->>ES: Fetch multiple sections
    
    par Complex Parallel Data Aggregation
        ES->>DB: Fetch Personalized Interests (joins products/offers using User onboarding data)
        DB-->>ES: Interests Array
        
        ES->>DB: Fetch Trending (raw SQL score: favorites*1 + views*0.5 + discount*0.2 + recency)
        DB-->>ES: Trending Array
        
        ES->>DB: Fetch Flash Offers (WHERE ends_at BETWEEN NOW() and NOW()+24h)
        DB-->>ES: Flash Array
        
        ES->>DB: Fetch Top Stores (rating_avg DESC + joins best coupon)
        DB-->>ES: Top Stores Array
        
        alt If Lat/Lng Provided
            ES->>DB: Fetch Nearby (Haversine formula against addressables pivot)
            DB-->>ES: Nearby Array
        end
    end
    
    ES-->>GEBA: Unified Payload
    GEBA-->>EC: Formatted Result
    EC-->>Client: 200 OK (ExploreBootstrapResource)
""", "sd_explore_bootstrap", diagrams, 6.0)

    # --- UC-114: Get Picked Offers ---
    add_diagram(doc, "Get Picked Offers (UC-114)", """sequenceDiagram
    participant Client
    participant EC as ExploreController
    participant GEPA as GetExplorePicksAction
    participant PRS as ProductRecommendationService
    participant DB as Database

    Client->>EC: GET /explore/picks?page=2&sort=trending
    EC->>GEPA: execute(Filters, Pagination, User)
    
    alt User is Authenticated
        GEPA->>PRS: getRecommendedProductIds(User)
        PRS-->>GEPA: Array of Recommended IDs
    end
    
    GEPA->>DB: Query Products + Offers + Stores
    GEPA->>DB: Apply Filters (min_discount raw SQL, category EXISTS)
    alt Recommended IDs exist
        GEPA->>DB: WHERE IN (Recommended IDs)
    else
        GEPA->>DB: Apply fallback sorting (Trending / Newest)
    end
    GEPA->>DB: Execute Paginated Query
    DB-->>GEPA: Paginated Results
    
    GEPA->>DB: Query product_favorites to resolve 'is_favorite' flags
    GEPA-->>EC: Paginator Object
    EC-->>Client: 200 OK (ExplorePicksResource)
""", "sd_explore_picks", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The Explore System relies on complex raw SQL queries and multi-table joins to calculate its sections dynamically, rather than using static 'explore' tables."
    )
    schemas = [
        "interests: Onboarding table containing user preferences (budget ranges, shopping styles, interesting categories) mapped dynamically to products.",
        "trending calculation: A dynamic raw SQL view combining `products.favorites_count`, `product_views.count` (last 7 days), and `product_offers.discount_percent`.",
        "flash offers: Derived dynamically by querying `product_offers.ends_at` relative to `server_time`.",
        "nearby calculation: Derives geospatial data by joining `stores` -> `addressables` -> `addresses.latitude/longitude`.",
        "banners: Independent tables (`banners`, `travel_banners`) injected into the UI via separate related endpoints."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o| INTERESTS : "defines preferences"
    INTERESTS ||--o{ PRODUCTS : "filters recommendations"
    PRODUCTS ||--|| PRODUCT_OFFERS : "contains discount/flash data"
    PRODUCTS ||--o{ PRODUCT_VIEWS : "generates trending score"
    STORES ||--o{ PRODUCTS : "sells"
    STORES ||--o{ ADDRESSABLES : "morphs_to"
    ADDRESSABLES ||--|| ADDRESSES : "provides nearby data"

    INTERESTS {
        uuid user_id PK
        json shopping_style
        json budget_range
    }
    PRODUCT_OFFERS {
        uuid id PK
        uuid product_id FK
        decimal discount_percent
        timestamp ends_at
    }
    PRODUCT_VIEWS {
        bigint id
        uuid product_id FK
        timestamp viewed_at
    }
    ADDRESSES {
        uuid id PK
        decimal latitude
        decimal longitude
    }
"""
    er_img_path = "explore_er_diagram_accurate.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Explore System chapter accurately appended to {filename}")

if __name__ == "__main__":
    append_explore_chapter()
