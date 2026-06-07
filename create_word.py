import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def create_word_file():
    doc = Document()
    
    # Title
    title = doc.add_heading('Application Features Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Chapter 1
    doc.add_heading('1. System Architecture', level=1)
    doc.add_paragraph('This document outlines the core features of the system.')
    
    # Subchapter: Features
    doc.add_heading('1.1 Features', level=2)
    doc.add_paragraph('This subchapter describes the main features implemented in the repository, starting with the Authentication feature.')
    
    # ================== AUTH FEATURE ==================
    doc.add_heading('1.1.1 Authentication Feature', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The authentication system is built using Laravel Sanctum, providing token-based authentication. "
        "The logic supports role-based access control (Admin, Seller, Customer) and verifies user roles during login. "
        "Upon providing valid email and password credentials, the system checks if the user's account is active and verified. "
        "If unverified, the system triggers an OTP (One-Time Password) generation and sends it via email. "
        "For verified users, an access token (short-lived) and a refresh token (30 days TTL) are generated. "
        "The authentication service also records session metadata such as IP address, user agent, and detected device type (mobile, tablet, desktop). "
        "Role assignment can be automatic; for example, if a user requests a 'customer' role during login and doesn't have it, it is assigned dynamically. "
        "The refresh token mechanism allows issuing new access tokens by validating the hashed refresh token stored in the database, invalidating the old session to maintain security. "
        "Additionally, users can register for an account (triggering an OTP), verify their OTP (which automatically logs them in), and authenticate via Google SSO."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-1: User Login - Users authenticate using email and password. Receives Access and Refresh tokens upon success.",
        "UC-2: Role-based Authorization - System validates if the user has the required role (Admin/Seller/Customer) to access specific endpoints.",
        "UC-3: Token Refresh - Client application exchanges a valid refresh token for a new access/refresh token pair without requiring user credentials.",
        "UC-4: Session Management & Logout - Users can view active sessions. They can log out of the current device (revoking the current token) or log out of all devices (revoking all tokens).",
        "UC-5: User Registration - Users can register an account. This automatically triggers an OTP code sent to their email for verification.",
        "UC-6: OTP Verification - Users submit their OTP code to verify their email. Upon successful verification, access tokens are issued and the user is logged in automatically.",
        "UC-7: Google Authentication - Users can register or log in using their Google account. The system verifies the ID token, resolves the user identity, checks onboarding status, and returns access tokens."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    doc.add_heading('Sequence Diagram: User Login Flow', level=5)
    login_mermaid = """sequenceDiagram
    participant C as Client
    participant AC as AuthController
    participant AS as AuthService
    participant DB as Database
    
    C->>AC: POST /auth/login (email, password, role)
    AC->>AS: login(email, password, context)
    AS->>DB: Find User by email
    DB-->>AS: User data
    AS->>AS: Verify Password
    AS->>AS: Check status & verification
    
    alt If unverified
        AS->>C: Return OTP Required error
    else If verified
        AS->>AS: Assert/Assign role
        AS->>AS: Generate AccessToken & RefreshToken
        AS->>DB: Store new session with device metadata
        AC-->>C: Return 200 OK with User data & Tokens
    end
"""
    login_img_path = "login_diagram.png"
    download_image(generate_kroki_url(login_mermaid), login_img_path)
    doc.add_picture(login_img_path, width=Inches(6.0))
    diagrams.append(login_img_path)
    
    doc.add_heading('Sequence Diagram: Token Refresh Flow', level=5)
    refresh_mermaid = """sequenceDiagram
    participant C as Client
    participant AC as AuthController
    participant AS as AuthService
    participant DB as Database
    
    C->>AC: POST /auth/refresh (refresh_token)
    AC->>AS: refreshToken(refresh_token)
    AS->>DB: Find Session by hashed token
    DB-->>AS: Session data
    AS->>AS: Check if session expired
    AS->>DB: Delete old tokens & session
    AS->>AS: Generate new AccessToken & RefreshToken
    AS->>DB: Store new session
    AC-->>C: Return new tokens
"""
    refresh_img_path = "refresh_diagram.png"
    download_image(generate_kroki_url(refresh_mermaid), refresh_img_path)
    doc.add_picture(refresh_img_path, width=Inches(5.0))
    diagrams.append(refresh_img_path)

    doc.add_heading('Sequence Diagram: User Registration Flow', level=5)
    register_mermaid = """sequenceDiagram
    participant C as Client
    participant RC as RegisterController
    participant RU as RegisterUser Action
    participant OS as OtpService
    participant DB as Database
    
    C->>RC: POST /auth/register
    RC->>RU: execute(UserData)
    RU->>DB: Create User & Profile
    RU-->>RC: User Instance
    RC->>OS: generateAndSend(VERIFY_EMAIL)
    OS->>DB: Store OTP Code
    OS->>C: Email sent with OTP
    RC-->>C: 201 Created (OTP Metadata returned)
"""
    register_img_path = "register_diagram.png"
    download_image(generate_kroki_url(register_mermaid), register_img_path)
    doc.add_picture(register_img_path, width=Inches(5.0))
    diagrams.append(register_img_path)

    doc.add_heading('Sequence Diagram: OTP Verification Flow', level=5)
    otp_mermaid = """sequenceDiagram
    participant C as Client
    participant OC as OtpController
    participant OS as OtpService
    participant AS as AuthService
    participant DB as Database
    
    C->>OC: POST /auth/otp/verify (code, email)
    OC->>OS: verify(user, code, VERIFY_EMAIL)
    OS->>DB: Validate OTP code against DB
    
    alt Invalid OTP
        OS-->>OC: Verification Failed
        OC-->>C: 4xx Error Response
    else Valid OTP & Verify Email
        OS->>DB: Mark email as verified
        OS-->>OC: Verification Success
        OC->>AS: issueTokensForUser(User)
        AS->>DB: Generate & Store Session
        AS-->>OC: Access & Refresh Tokens
        OC-->>C: 200 OK (Tokens & User Data)
    end
"""
    otp_img_path = "otp_diagram.png"
    download_image(generate_kroki_url(otp_mermaid), otp_img_path)
    doc.add_picture(otp_img_path, width=Inches(5.0))
    diagrams.append(otp_img_path)

    doc.add_heading('Sequence Diagram: Google Authentication Flow', level=5)
    google_mermaid = """sequenceDiagram
    participant C as Client
    participant GC as GoogleLoginController
    participant GV as GoogleTokenVerifier
    participant AS as AuthService
    participant DB as Database
    
    C->>GC: POST /auth/google (id_token)
    GC->>GV: verifyIdToken(id_token)
    GV->>Google: Validate Token Signature
    GV-->>GC: Google User Payload
    GC->>DB: Find User by provider_id or Email
    alt User not found
        GC->>DB: Create new User & Profile
    else User found
        GC->>DB: Link Google provider info
    end
    GC->>DB: Mark email as verified
    GC->>AS: issueTokensForUser(User)
    AS->>DB: Store session metadata
    GC->>DB: Check onboarding status
    GC-->>C: 200 OK (Tokens, User Data, Onboarding Status)
"""
    google_img_path = "google_diagram.png"
    download_image(generate_kroki_url(google_mermaid), google_img_path)
    doc.add_picture(google_img_path, width=Inches(6.0))
    diagrams.append(google_img_path)


    doc.add_heading('Database Schema', level=4)
    auth_db_text = (
        "The authentication logic relies on several database tables to manage user credentials, metadata, sessions, and verification processes."
    )
    doc.add_paragraph(auth_db_text)
    
    auth_schemas = [
        "users: Core user table storing credentials and status. Fields include id (UUID), email, password_hash, phone_number, verification timestamps, status (active/suspended/deleted), and provider details for social login.",
        "profiles: Stores extended user details. Fields include id, user_id (FK to users), first_name, last_name, date_of_birth, gender, and avatar_url.",
        "sessions: Custom session table for tracking active logins. Fields include id (UUID), user_id (FK to users), token, user_agent, ip_address, device_type, last_activity, and expires_at.",
        "otps: Manages One-Time Passwords for verification. Fields include id, user_id (nullable FK), phone_or_email, otp_hash, purpose, channel, status, attempts, and expires_at."
    ]
    for schema in auth_schemas:
        doc.add_paragraph(schema, style='List Bullet')

    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    auth_er_mermaid = """erDiagram
    USERS ||--o| PROFILES : "has_one"
    USERS ||--o{ SESSIONS : "has_many"
    USERS ||--o{ OTPS : "generates"
    
    USERS {
        char id PK "UUID"
        string email
        string password_hash
        string status
    }
    PROFILES {
        bigint id PK
        uuid user_id FK
        string first_name
        string last_name
    }
    SESSIONS {
        uuid id PK
        uuid user_id FK
        string token
        string ip_address
        string device_type
    }
    OTPS {
        bigint id PK
        uuid user_id FK
        string phone_or_email
        string otp_hash
        string purpose
        string status
    }
"""
    auth_er_img_path = "auth_er_diagram.png"
    download_image(generate_kroki_url(auth_er_mermaid), auth_er_img_path)
    doc.add_picture(auth_er_img_path, width=Inches(6.0))
    diagrams.append(auth_er_img_path)

    # ================== ROLES, PERMISSIONS & STAFF MANAGEMENT FEATURE ==================
    doc.add_heading('1.1.2 Roles, Permissions & Staff Management', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    staff_logic = (
        "The system provides a comprehensive Staff and Role management system divided into two layers: Admin-level Role Management and Store-level Employee Management. "
        "At the Admin level, roles and their associated permissions can be dynamically created, updated, and retrieved using the Spatie Permission package. "
        "At the Store level, store owners can invite other users to join their store as employees using the StoreInvitationController. "
        "Invitations can be sent, accepted, declined, or cancelled. Once a user accepts an invitation and becomes a store employee, the store owner can assign them specific store-related roles (e.g., store_manager, cashier, support_agent). "
        "The system ensures that users are granted the appropriate global Spatie roles based on their store assignments, and automatically cleans up roles when an employee is removed from a store. "
        "Critical actions, such as an owner attempting to remove themselves, are blocked by validation rules."
    )
    doc.add_paragraph(staff_logic)

    doc.add_heading('Use Cases', level=4)
    staff_ucs = [
        "UC-8: Admin Role Management - System administrators can dynamically create, update, and delete roles and assign specific permissions.",
        "UC-9: Invite Store Employee - A store owner sends an invitation to a user to join their store staff.",
        "UC-10: Manage Invitations - Users can view, accept, or decline pending store invitations.",
        "UC-11: Manage Store Employees - Store managers can view, update roles, or remove employees from their store.",
        "UC-12: Update Employee Role - A store owner updates an existing employee's role, automatically revoking old global permissions and assigning new ones.",
        "UC-13: Remove Store Employee - A store owner removes an employee, automatically cleaning up their store-specific access rights.",
        "UC-14: Resend or Cancel Invitations - A store owner resends a pending invitation or cancels it.",
        "UC-15: Decline Store Invitation - A user explicitly declines a pending store invitation.",
        "UC-16: Admin Custom Role Creation - System Administrators create entirely new roles and attach granular Spatie permissions to them."
    ]
    for uc in staff_ucs:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    doc.add_heading('Sequence Diagram: Store Invitation Flow', level=5)
    invite_mermaid = """sequenceDiagram
    participant Owner as Store Owner
    participant IC as StoreInvitationController
    participant IS as StoreInvitationService
    participant DB as Database
    
    Owner->>IC: POST /stores/{store}/invitations (email, role)
    IC->>IS: sendInvitation(store, user, data)
    IS->>DB: Check if user already exists
    IS->>DB: Create StoreInvitation record
    IS->>Owner: Send Notification/Email to invitee
    IC-->>Owner: 201 Created (Invitation details)
"""
    invite_img_path = "invite_diagram.png"
    download_image(generate_kroki_url(invite_mermaid), invite_img_path)
    doc.add_picture(invite_img_path, width=Inches(5.0))
    diagrams.append(invite_img_path)

    doc.add_heading('Sequence Diagram: Accept Invitation Flow', level=5)
    accept_mermaid = """sequenceDiagram
    participant User
    participant IC as StoreInvitationController
    participant IS as StoreInvitationService
    participant DB as Database
    
    User->>IC: POST /invitations/{invitation}/accept
    IC->>IS: acceptInvitation(user, invitation)
    IS->>DB: Mark invitation as accepted
    IS->>DB: Create StoreEmployee record
    IS->>DB: Assign store-specific Spatie role to User
    IC-->>User: 200 OK (Invitation accepted)
"""
    accept_img_path = "accept_diagram.png"
    download_image(generate_kroki_url(accept_mermaid), accept_img_path)
    doc.add_picture(accept_img_path, width=Inches(5.0))
    diagrams.append(accept_img_path)

    doc.add_heading('Sequence Diagram: Update Store Employee Role Flow', level=5)
    update_role_mermaid = """sequenceDiagram
    participant Owner as Store Owner
    participant SEC as StoreEmployeeController
    participant DB as Database
    
    Owner->>SEC: PUT /stores/{store}/employees/{user} (new_role)
    SEC->>DB: Fetch StoreEmployee record
    SEC->>SEC: Check Owner validation rules
    SEC->>DB: Update StoreEmployee role in DB
    
    alt If role has changed
        SEC->>DB: Check if old role still used in other branches
        alt Role not used elsewhere
            SEC->>DB: Revoke old Spatie Role from User
        end
        SEC->>DB: Assign new Spatie Role to User
    end
    SEC-->>Owner: 200 OK (Updated Employee details)
"""
    update_role_img_path = "update_role_diagram.png"
    download_image(generate_kroki_url(update_role_mermaid), update_role_img_path)
    doc.add_picture(update_role_img_path, width=Inches(6.0))
    diagrams.append(update_role_img_path)

    doc.add_heading('Sequence Diagram: Remove Store Employee Flow', level=5)
    remove_emp_mermaid = """sequenceDiagram
    participant Owner as Store Owner
    participant SEC as StoreEmployeeController
    participant DB as Database
    
    Owner->>SEC: DELETE /stores/{store}/employees/{user}
    SEC->>DB: Fetch StoreEmployee record
    SEC->>SEC: Prevent self-deletion if Owner
    SEC->>DB: Delete StoreEmployee record
    
    SEC->>DB: Check if user works at ANY other store
    alt User has no other store roles
        SEC->>DB: Revoke all store-specific Spatie Roles
    end
    SEC-->>Owner: 200 OK (Employee removed)
"""
    remove_emp_img_path = "remove_emp_diagram.png"
    download_image(generate_kroki_url(remove_emp_mermaid), remove_emp_img_path)
    doc.add_picture(remove_emp_img_path, width=Inches(5.0))
    diagrams.append(remove_emp_img_path)

    doc.add_heading('Sequence Diagram: Admin Role Creation Flow', level=5)
    admin_role_mermaid = """sequenceDiagram
    participant Admin
    participant RPC as RolePermissionController
    participant SRA as StoreRoleAction
    participant DB as Database
    
    Admin->>RPC: POST /admin/roles (name, permissions[])
    RPC->>SRA: execute(RoleDTO)
    SRA->>DB: Create new Spatie Role
    SRA->>DB: Fetch specified Permissions
    SRA->>DB: Attach Permissions to new Role
    SRA-->>RPC: Role created
    RPC-->>Admin: 201 Created (Role & Permissions)
"""
    admin_role_img_path = "admin_role_diagram.png"
    download_image(generate_kroki_url(admin_role_mermaid), admin_role_img_path)
    doc.add_picture(admin_role_img_path, width=Inches(5.0))
    diagrams.append(admin_role_img_path)

    doc.add_heading('Database Schema', level=4)
    db_schema_text = (
        "The feature leverages the Spatie Laravel Permission package's standard schema along with custom tables for store context."
    )
    doc.add_paragraph(db_schema_text)
    
    schemas = [
        "roles: Standard Spatie table containing id, name, guard_name, timestamps.",
        "permissions: Standard Spatie table containing id, name, guard_name, timestamps.",
        "model_has_roles: Pivot table linking users (model_id, model_type) to roles (role_id).",
        "model_has_permissions: Pivot table linking users directly to permissions.",
        "role_has_permissions: Pivot table linking roles to permissions.",
        "store_employees: Stores the relationship between a user and a store. Fields: id, store_id (uuid, FK to stores), user_id (FK to users), timestamps. Unique constraint on [store_id, user_id].",
        "store_invitations: Manages pending invitations. Fields: id, store_id (uuid), invited_by_user_id, invitee_user_id, role (string), permissions (json), status (string, default 'pending'), message (text), expires_at, accepted_at, declined_at, timestamps.",
        "user_roles: Custom table for explicit tracking of user roles per store. Fields: id, user_id (uuid), role_id (unsigned bigint), store_id (uuid, nullable), granted_at, granted_by_user_id, expires_at."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ STORE_EMPLOYEES : "has"
    STORES ||--o{ STORE_EMPLOYEES : "employs"
    USERS ||--o{ STORE_INVITATIONS : "invited_by / invitee"
    STORES ||--o{ STORE_INVITATIONS : "has"
    USERS ||--o{ MODEL_HAS_ROLES : "has"
    ROLES ||--o{ MODEL_HAS_ROLES : "assigned_to"
    ROLES ||--o{ ROLE_HAS_PERMISSIONS : "contains"
    PERMISSIONS ||--o{ ROLE_HAS_PERMISSIONS : "assigned_to"
    USERS ||--o{ USER_ROLES : "assigned"
    STORES ||--o{ USER_ROLES : "scopes"
    ROLES ||--o{ USER_ROLES : "defines"
    
    STORE_EMPLOYEES {
        bigint id PK
        uuid store_id FK
        char user_id FK
    }
    STORE_INVITATIONS {
        bigint id PK
        uuid store_id FK
        char invited_by_user_id FK
        char invitee_user_id FK
        string role
        string status
    }
    USER_ROLES {
        bigint id PK
        uuid user_id FK
        bigint role_id FK
        uuid store_id FK
    }
    ROLES {
        bigint id PK
        string name
    }
    PERMISSIONS {
        bigint id PK
        string name
    }
"""
    er_img_path = "er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    filename = "Features_Documentation_Diagrams.docx"
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Word document saved as {filename}")

if __name__ == "__main__":
    create_word_file()
