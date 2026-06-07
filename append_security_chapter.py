import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_security_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 4: SECURITY & ROLES ==================
    doc.add_page_break()
    doc.add_heading('Chapter 4: Security, Roles & Permissions Matrix', level=1)
    
    intro_text = (
        "The backend implements a highly granular, enterprise-grade Role-Based Access Control (RBAC) system powered by Spatie's Laravel Permission package. "
        "The architecture completely abandons legacy hardcoded role checks in favor of a robust Matrix consisting of 13 distinct Roles and over 30 explicit Permissions."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 4.1 The Role Hierarchy ------------------
    doc.add_heading('4.1 The Role Hierarchy', level=2)
    doc.add_paragraph("Users in the system are assigned one or more of the following 13 exact roles:")
    roles = [
        "super_admin & admin: Global platform administrators.",
        "seller & seller_pending: High-level store ownership statuses.",
        "store_owner: The master role for a specific store.",
        "store_manager: Operational manager with broad access.",
        "branch_manager: Restricted manager specifically for physical branch operations.",
        "inventory_manager: Focused on products and offers.",
        "content_manager: Focused on marketing, reviews, and offer creation.",
        "cashier & store_employee: Point-of-sale staff restricted to order viewing and claim redemption.",
        "support_agent: Customer service staff.",
        "customer: The base end-user role."
    ]
    for r in roles:
        doc.add_paragraph(r, style='List Bullet')

    # ------------------ 4.2 Permissions Matrix ------------------
    doc.add_heading('4.2 Store Permissions Matrix', level=2)
    doc.add_paragraph("Store-level access is governed by the following strict permission bindings:")
    
    matrix = [
        "store_owner: Inherits ALL 30+ store.* permissions automatically.",
        "store_manager: dashboard.view, products.manage, offers.manage, claims.manage, orders.manage, employees.view, branches.view, analytics.view, reviews.view.",
        "inventory_manager: products.view, products.create, products.update, offers.view, offers.update.",
        "content_manager: products.manage, offers.manage, reviews.view.",
        "branch_manager: dashboard.view, claims.manage, orders.view, employees.view, branches.view.",
        "cashier & store_employee: Strictly limited to claims.view, claims.redeem, and orders.view.",
        "support_agent: orders.view, claims.view, reviews.view.",
        "seller: dashboard.view, products.manage, orders.view, analytics.view.",
        "seller_pending: Strictly limited to dashboard.view (read-only state until platform approval)."
    ]
    for m in matrix:
        doc.add_paragraph(m, style='List Bullet')

    # ------------------ 4.3 Security Standards ------------------
    doc.add_heading('4.3 Global Security Standards', level=2)
    sec = [
        "Guard: All API authentication and permission resolution is routed strictly through the 'sanctum' guard.",
        "Cache Invalidation: The system aggressively utilizes `forgetCachedPermissions()` during role syncing to prevent privilege escalation vulnerabilities.",
        "Legacy Migration: Hardcoded legacy permissions (e.g., 'edit users', 'delete categories') are actively purged during database seeding to enforce the new structured `domain.entity.action` syntax."
    ]
    for s in sec:
        doc.add_paragraph(s, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"Security Matrix Chapter appended to {filename}")

if __name__ == "__main__":
    append_security_chapter()
