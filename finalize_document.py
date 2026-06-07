import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def finalize_document():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # Extract all headings for the Table of Contents before modifying
    toc_lines = []
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading 1'):
            toc_lines.append((1, p.text))
        elif p.style.name.startswith('Heading 2'):
            toc_lines.append((2, p.text))
        elif p.style.name.startswith('Heading 3'):
            toc_lines.append((3, p.text))

    # We need to insert elements at the top of the document.
    # To maintain order when inserting before the very first element repeatedly,
    # we can gather what we want to insert as a list of actions and apply them.
    # But an easier way is to create a NEW document, add the Title and TOC, 
    # and then append everything from the old document. But python-docx loses images if we just copy paragraphs.
    
    # Safer approach: insert before the first paragraph in reverse order!
    first_p = doc.paragraphs[0]

    # Elements to insert (in normal order). We will apply them by inserting before first_p
    # Actually, we can just use python-docx's low level XML to insert a new paragraph at the beginning.
    # Let's collect all new paragraphs in a dummy document, then move their XML elements.
    
    dummy = Document()
    
    # --- Title Page ---
    title = dummy.add_paragraph("Coupony Platform")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(36)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle = dummy.add_paragraph("Backend Technical Specification & Architecture Master Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    dummy.add_paragraph("")
    dummy.add_paragraph("")
    
    date_p = dummy.add_paragraph(f"Generated Date: {datetime.now().strftime('%Y-%m-%d')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p = dummy.add_paragraph("Version: 1.0 (Comprehensive Build)")
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dummy.add_page_break()

    # --- Table of Contents ---
    toc_title = dummy.add_heading("Table of Contents", level=1)
    
    for level, text in toc_lines:
        if level == 1:
            p = dummy.add_paragraph(text)
            p.runs[0].font.bold = True
        elif level == 2:
            p = dummy.add_paragraph(f"    {text}")
        elif level == 3:
            p = dummy.add_paragraph(f"        {text}")

    dummy.add_page_break()

    # Now, prepend all elements from dummy to the original doc.
    # We take the body elements of the dummy and insert them before the first element of doc.
    doc_body = doc._body._body
    dummy_body = dummy._body._body
    
    first_element = doc_body[0]
    
    # Insert elements from dummy in their original order
    for element in dummy_body:
        # We need to deepcopy the element because it belongs to dummy
        import copy
        new_element = copy.deepcopy(element)
        first_element.addprevious(new_element)

    doc.save(filename)
    print("Title Page and Table of Contents successfully prepended!")

if __name__ == "__main__":
    finalize_document()
