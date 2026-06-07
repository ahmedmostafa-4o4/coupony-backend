import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_localization_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 7: LOCALIZATION ==================
    doc.add_page_break()
    doc.add_heading('Chapter 7: Localization & Internationalization (i18n)', level=1)
    
    intro_text = (
        "The platform is designed from the ground up to support multiple languages seamlessly, "
        "currently fully supporting English ('en') and Arabic ('ar'). The localization strategy ensures that "
        "every API response, validation error, and push notification is delivered in the user's preferred language natively."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 7.1 Dynamic Locale Resolution ------------------
    doc.add_heading('7.1 Dynamic Locale Resolution', level=2)
    doc.add_paragraph("The backend never hardcodes the application locale. Instead, it relies on a sophisticated middleware resolution chain (`SetLocale` and `UseAuthenticatedUserLocale`) that executes on every request:")
    res = [
        "1. Accept-Language Header: The API first checks the `Accept-Language` HTTP header sent by the mobile app or frontend.",
        "2. Database Fallback: If the header is missing, the backend extracts the Bearer token, finds the authenticated User, and falls back to the `language` column stored in the database.",
        "3. System Default: If the user is a guest and provides no header, the system safely falls back to English (`en`) as defined in `config/localization.php`."
    ]
    for r in res:
        doc.add_paragraph(r, style='List Bullet')

    # ------------------ 7.2 Translation Architecture ------------------
    doc.add_heading('7.2 API Messaging & Translations', level=2)
    doc.add_paragraph("No hardcoded strings exist within the Controllers or Actions. All messaging is routed through Laravel's native translation engine:")
    trans = [
        "Language Files: All messages are stored securely in `lang/en/api.php` and `lang/ar/api.php`.",
        "Dynamic Interpolation: Controllers use the `__('api.notifications.deleted_read', ['count' => $deleted])` syntax to inject dynamic variables (like counts or names) into translated strings.",
        "Response Headers: Every API response automatically attaches a `Content-Language` header (e.g., `Content-Language: ar`), ensuring the client application is aware of the dialect the backend chose to respond with."
    ]
    for t in trans:
        doc.add_paragraph(t, style='List Bullet')

    # ------------------ 7.3 Data Localization Strategy ------------------
    doc.add_heading('7.3 Data Localization & Search', level=2)
    data = [
        "User Generated Content: Core user inputs (like a Store's Title or Product Description) are stored as raw strings in the database, allowing Store Owners to input their catalog in their regional dialect.",
        "Arabic Text Normalization: The Native Search system features a dedicated `ArabicTextNormalizer`. This strips out complex Arabic diacritics (Tashkeel) and normalizes character variations (e.g., converting 'أ', 'إ', 'آ' into 'ا') to ensure robust full-text search accuracy regardless of how the user types the query."
    ]
    for d in data:
        doc.add_paragraph(d, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"Localization Chapter appended to {filename}")

if __name__ == "__main__":
    append_localization_chapter()
