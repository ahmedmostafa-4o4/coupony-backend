import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_external_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 3: EXTERNAL INTEGRATIONS ==================
    doc.add_page_break()
    doc.add_heading('Chapter 3: External Integrations & 3rd Party Services', level=1)
    
    intro_text = (
        "This chapter maps out all the external dependencies and 3rd-party services the backend relies on to function. "
        "It provides a clear view of where data leaves our infrastructure and interacts with proprietary external APIs."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 3.1 Google Gemini AI ------------------
    doc.add_heading('3.1 Google Gemini AI (PonyAI Assistant)', level=2)
    doc.add_paragraph("The built-in PonyAI chatbot heavily relies on Google's generative AI models to provide natural language interactions and autonomous tool calling.")
    gemini = [
        "Base URL: `https://generativelanguage.googleapis.com/v1beta`",
        "Text/Vision Model: Utilizes `gemini-2.5-flash` for high-speed conversational responses and visual parsing.",
        "Embedding Model: Utilizes `text-embedding-004` for generating vector embeddings of products to allow semantic search.",
        "Configuration: The timeout is globally capped at 20 seconds, with an explicit mock/fake mode configurable via `.env` for local testing without incurring API costs."
    ]
    for g in gemini:
        doc.add_paragraph(g, style='List Bullet')

    # ------------------ 3.2 HuggingFace ML Recommendation Service ------------------
    doc.add_heading('3.2 Custom ML Recommendation Engine', level=2)
    doc.add_paragraph("Instead of processing heavy machine learning models natively on the PHP backend, the platform offloads this to a dedicated external Python-based Microservice.")
    ml = [
        "Host Provider: HuggingFace Spaces (`ahmedmostafa56-ml-recommendation-service.hf.space`).",
        "Purpose: Calculates complex collaborative-filtering recommendations for the 'Picked for You' user feeds.",
        "Configuration: Strict 10-second timeout, with a seed limit of 20 to prevent runaway recursive calculation."
    ]
    for m in ml:
        doc.add_paragraph(m, style='List Bullet')

    # ------------------ 3.3 Firebase Cloud Messaging (FCM) ------------------
    doc.add_heading('3.3 Firebase Cloud Messaging (FCM)', level=2)
    doc.add_paragraph("The backend connects directly to Google's Firebase network to handle background push delivery.")
    fcm = [
        "Push Mechanism: Converts internal Laravel Notifications into HTTP payloads pushed to FCM.",
        "Targeting: Resolves `device_tokens` internally (mapping user IDs to iOS/Android registration tokens) and pushes the payload."
    ]
    for f in fcm:
        doc.add_paragraph(f, style='List Bullet')

    # ------------------ 3.4 Email & Communication Providers ------------------
    doc.add_heading('3.4 Email & Communication APIs', level=2)
    doc.add_paragraph("The backend dynamically routes automated communication depending on environment configurations in `config/services.php`:")
    email = [
        "AWS SES (Simple Email Service): The primary provider for high-volume transactional emails (like Monthly Analytics PDF Reports or Welcome Emails).",
        "Resend & Postmark: Configured as robust failover options for email delivery.",
        "Slack Integration: Connected via OAuth to push critical administrative alerts (e.g., server issues, massive bulk operations) directly to the company's internal Slack channel."
    ]
    for e in email:
        doc.add_paragraph(e, style='List Bullet')

    # ------------------ 3.5 Cloud Storage (AWS S3) ------------------
    doc.add_heading('3.5 AWS S3 Cloud Storage', level=2)
    doc.add_paragraph("While the database stores pointers and metadata, physical binary files are handled by Amazon Web Services.")
    s3 = [
        "Storage Strategy: Stores all product images, store logos, and generated analytics PDFs.",
        "Optimization: Files are processed asynchronously by Horizon workers (converted to WebP, compressed, watermarked) before being permanently stored on S3."
    ]
    for s in s3:
        doc.add_paragraph(s, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"External Integrations Chapter appended to {filename}")

if __name__ == "__main__":
    append_external_chapter()
