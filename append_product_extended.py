import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_extended_product_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    doc.add_heading('Additional Use Cases & Sequence Diagrams for Product System', level=4)

    # ======================== ADDITIONAL USE CASES ========================
    doc.add_heading('Extended Use Cases', level=4)
    use_cases = [
        "UC-78: Reorder Product Images - Store Owner rearranges the sort_order of their product image gallery.",
        "UC-79: Set Primary Image - Store Owner designates a specific image as the product's primary thumbnail.",
        "UC-80: Create Product Variant - Store Owner adds a new variant (e.g., size/color) with its own SKU, price override, and stock.",
        "UC-81: Update Product Variant - Store Owner modifies a variant's pricing, stock, or active status.",
        "UC-82: Delete Product Variant - Store Owner removes a variant from the product.",
        "UC-83: Replace Variant Attributes - Store Owner bulk-replaces all key-value attribute pairs on a variant.",
        "UC-84: Record Product Share - The system records a share event when a user shares a product to a social platform (WhatsApp, Facebook, Twitter, etc.).",
        "UC-85: Get Product Recommendations - The system returns personalized product recommendations for the authenticated user.",
        "UC-86: View Seller Revisions - Store Owner views their product's revision history to track pending/approved/rejected changes.",
        "UC-87: Create Offer Claim - A customer claims an active product offer, generating a unique QR code token for in-store redemption.",
        "UC-88: Redeem Offer Claim - A store staff member scans the QR code to redeem the claim, decrementing stock and awarding points to both user and store.",
        "UC-89: Record Product View - The system records an analytics event each time a user views a product detail page.",
        "UC-90: Admin Create Product - Admin creates a product that is auto-approved (bypasses the revision workflow entirely).",
        "UC-91: Admin Reject Revision with Structured Feedback - Admin rejects a revision with per-field requested_changes targeting specific variants, images, or attributes.",
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Extended Sequence Diagrams', level=4)

    # --- UC-78: Reorder Product Images ---
    add_diagram(doc, "Reorder Product Images (UC-78)", """sequenceDiagram
    participant Owner as Store Owner
    participant PIC as ProductImageController
    participant Action as ReorderProductImages
    participant DB as Database

    Owner->>PIC: PUT /stores/{store}/products/{product}/images/reorder
    PIC->>PIC: Authorize (Store Ownership)
    PIC->>Action: execute(Product, ordered IDs)
    loop Each Image ID in Order
        Action->>DB: UPDATE product_images SET sort_order = index
    end
    Action-->>PIC: Success
    PIC-->>Owner: 200 OK (Reordered Images)
""", "sd_reorder_images", diagrams)

    # --- UC-79: Set Primary Image ---
    add_diagram(doc, "Set Primary Image (UC-79)", """sequenceDiagram
    participant Owner as Store Owner
    participant PIC as ProductImageController
    participant Action as SetPrimaryProductImage
    participant DB as Database

    Owner->>PIC: POST /images/{image}/set-primary
    PIC->>PIC: Authorize (Store Ownership)
    PIC->>Action: execute(Product, Image)
    Action->>DB: UPDATE product_images SET is_primary = false WHERE product_id
    Action->>DB: UPDATE product_images SET is_primary = true WHERE id = Image.id
    Action-->>PIC: Success
    PIC-->>Owner: 200 OK
""", "sd_set_primary_image", diagrams)

    # --- UC-80: Create Product Variant ---
    add_diagram(doc, "Create Product Variant (UC-80)", """sequenceDiagram
    participant Owner as Store Owner
    participant PVC as ProductVariantController
    participant CVA as CreateProductVariant
    participant RVOP as ResolveVariantOfferPricing
    participant DB as Database

    Owner->>PVC: POST /products/{product}/variants
    PVC->>PVC: Authorize (Store Ownership)
    PVC->>CVA: execute(Product, VariantData)
    CVA->>DB: Insert into product_variants
    CVA->>DB: Insert variant_attributes (key-value pairs)
    CVA->>RVOP: Resolve pricing against active offer
    RVOP->>DB: Calculate final price (fixed/percentage/buy_x_get_y)
    CVA->>DB: Sync product-level base_price from default variant
    CVA-->>PVC: New Variant
    PVC-->>Owner: 201 Created
""", "sd_create_variant", diagrams, 5.5)

    # --- UC-83: Replace Variant Attributes ---
    add_diagram(doc, "Replace Variant Attributes (UC-83)", """sequenceDiagram
    participant Owner as Store Owner
    participant PVC as ProductVariantController
    participant RVA as ReplaceVariantAttributes
    participant DB as Database

    Owner->>PVC: PUT /variants/{variant}/attributes
    PVC->>PVC: Authorize (Store Ownership)
    PVC->>RVA: execute(Variant, NewAttributes[])
    RVA->>DB: DELETE all existing variant_attributes for Variant
    loop Each New Attribute
        RVA->>DB: INSERT variant_attribute (name, value, sort_order)
    end
    RVA-->>PVC: Updated Variant with Attributes
    PVC-->>Owner: 200 OK
""", "sd_replace_attributes", diagrams)

    # --- UC-84: Record Product Share ---
    add_diagram(doc, "Record Product Share (UC-84)", """sequenceDiagram
    participant User
    participant PSC as ProductShareController
    participant DB as Database

    User->>PSC: POST /products/{product}/share (platform=whatsapp)
    PSC->>PSC: Validate platform (whatsapp/facebook/twitter/instagram/copy_link)
    PSC->>DB: Insert into product_shares (product_id, user_id, platform)
    PSC-->>User: 200 OK (Share Recorded)
""", "sd_record_share", diagrams, 4.5)

    # --- UC-85: Get Product Recommendations ---
    add_diagram(doc, "Get Product Recommendations (UC-85)", """sequenceDiagram
    participant User
    participant PRC as ProductRecommendationController
    participant PRS as ProductRecommendationService
    participant DB as Database

    User->>PRC: GET /products/recommendations?limit=10
    PRC->>PRS: recommendFor(User, limit=10)
    PRS->>DB: Analyze user's likes, favorites, views, purchase history
    PRS->>DB: Query similar products (category, tags, engagement scores)
    DB-->>PRS: Ranked Product Collection
    PRS-->>PRC: Recommended Products
    PRC-->>User: 200 OK (Personalized Recommendations)
""", "sd_recommendations", diagrams)

    # --- UC-87: Create Offer Claim ---
    add_diagram(doc, "Create Offer Claim (UC-87)", """sequenceDiagram
    participant Customer
    participant COC as CreateOfferClaim Action
    participant DB as Database

    Customer->>COC: Claim Offer on Product
    COC->>DB: Validate Product is ACTIVE + APPROVED
    COC->>DB: Validate Offer is active and within time window
    alt Offer Type = buy_x_get_y
        COC->>COC: Resolve buy/reward variant selections
        COC->>COC: Apply mix_buy_variants / mix_reward_variants rules
    end
    COC->>COC: Generate unique claim_token
    COC->>COC: Generate unique qr_code_token
    COC->>DB: Insert into offer_claims (status=active, offer_snapshot JSON)
    COC->>COC: Fire OfferClaimCreated Event
    COC-->>Customer: 201 Created (Claim + QR Code)
""", "sd_create_offer_claim", diagrams, 5.5)

    # --- UC-88: Redeem Offer Claim ---
    add_diagram(doc, "Redeem Offer Claim via QR Code (UC-88)", """sequenceDiagram
    participant Staff as Store Staff
    participant ROC as RedeemOfferClaim Action
    participant PS as PointsService
    participant DB as Database

    Staff->>ROC: Scan QR Code Token
    ROC->>DB: Find offer_claim by qr_code_token
    ROC->>ROC: Validate claim belongs to this Store
    ROC->>ROC: Validate status = active and not expired
    ROC->>DB: Begin Transaction
    alt Tracked Inventory Variants
        ROC->>DB: Decrement stock_qty on product_variants
    end
    ROC->>DB: Increment variant.redemption_count
    ROC->>DB: Increment product.redemption_count
    ROC->>DB: Update claim status = redeemed, redeemed_at, redeemed_by
    ROC->>PS: Award Points to Customer
    ROC->>PS: Award Points to Store
    ROC->>ROC: Fire OfferClaimRedeemed Event
    DB-->>ROC: Commit Transaction
    ROC-->>Staff: 200 OK (Redemption Confirmed)
""", "sd_redeem_offer_claim", diagrams, 6.0)

    # --- UC-90: Admin Create Product (Auto-Approved) ---
    add_diagram(doc, "Admin Create Product - Auto-Approved (UC-90)", """sequenceDiagram
    participant Admin
    participant PMC as ProductManagementController
    participant CAP as CreateAdminProduct
    participant DB as Database

    Admin->>PMC: POST /admin/products (store_id, title, pricing)
    PMC->>CAP: execute(Store, ProductData)
    CAP->>DB: Insert product (status=ACTIVE, approval=APPROVED)
    CAP->>DB: Set published_revision_no = 1
    CAP->>DB: Set approved_at = now(), approved_by = Admin
    CAP->>DB: Sync images, variants, offer (same as seller flow)
    Note over CAP: No revision workflow - changes are LIVE immediately
    CAP-->>PMC: Approved Product
    PMC-->>Admin: 201 Created
""", "sd_admin_create_product", diagrams)

    # --- UC-91: Admin Reject Revision with Structured Feedback ---
    add_diagram(doc, "Admin Reject Revision with Structured Feedback (UC-91)", """sequenceDiagram
    participant Admin
    participant PRMC as ProductRevisionManagementController
    participant RR as RejectProductRevision
    participant DB as Database

    Admin->>PRMC: POST /revisions/{rev}/reject (reason, requested_changes)
    Note over Admin: requested_changes contains per-field feedback
    PRMC->>RR: execute(Revision, reason, notes, requested_changes)
    RR->>DB: Update revision status = REJECTED
    RR->>DB: Store rejection_reason, admin_notes
    RR->>DB: Store requested_changes JSON (field selectors)
    alt Product never approved (published_revision_no = 0)
        RR->>DB: Mark product approval_status = REJECTED
    end
    RR->>RR: Fire ProductRevisionRejected Event
    RR-->>PRMC: Rejected Revision
    PRMC-->>Admin: 200 OK
""", "sd_admin_reject_structured", diagrams)

    # ======================== UPDATED ER DIAGRAM ========================
    doc.add_heading('Complete Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    STORES ||--o{ PRODUCTS : "owns"
    PRODUCTS ||--o{ PRODUCT_IMAGES : "has_gallery"
    PRODUCTS ||--o{ PRODUCT_VARIANTS : "has_variants"
    PRODUCT_VARIANTS ||--o{ PRODUCT_VARIANT_ATTRIBUTES : "has_attributes"
    PRODUCTS ||--o{ PRODUCT_REVISIONS : "has_revisions"
    PRODUCTS ||--|| PRODUCT_OFFERS : "has_offer"
    PRODUCT_OFFERS ||--o{ OFFER_VARIANT_TARGETS : "targets"
    PRODUCT_OFFERS ||--o{ OFFER_CLAIMS : "generates"
    PRODUCTS ||--o{ PRODUCT_VIEWS : "tracks_views"
    PRODUCTS ||--o{ PRODUCT_SHARES : "tracks_shares"
    PRODUCTS ||--o{ PRODUCT_CATEGORIES : "categorized_by"

    PRODUCTS {
        uuid id PK
        uuid store_id FK
        string title
        decimal base_price
        string status
        string approval_status
        int published_revision_no
    }
    PRODUCT_VARIANTS {
        uuid id PK
        uuid product_id FK
        decimal original_price
        decimal price
        string inventory_mode
        int stock_qty
    }
    PRODUCT_OFFERS {
        uuid id PK
        uuid product_id FK
        string type "fixed/percentage/buy_x_get_y"
        string status
        timestamp starts_at
        timestamp ends_at
    }
    OFFER_CLAIMS {
        uuid id PK
        uuid offer_id FK
        char user_id FK
        string status "active/redeemed/expired"
        string qr_code_token
    }
    PRODUCT_REVISIONS {
        bigint id PK
        uuid product_id FK
        int revision_no
        string action "create/update/resubmit"
        string status
        json payload
    }
"""
    er_img_path = "product_extended_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Extended Product System diagrams appended to {filename}")

if __name__ == "__main__":
    append_extended_product_diagrams()
