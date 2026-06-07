import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_more_likes_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    doc.add_heading('Additional Sequence Diagrams for Likes System', level=4)
    
    diagrams = []
    
    # 3. Unlike Product Flow
    doc.add_heading('Sequence Diagram: Unlike Product Flow', level=5)
    unlike_product_mermaid = """sequenceDiagram
    participant User
    participant PLC as ProductLikeController
    participant ULA as UnlikeProduct Action
    participant DB as Database
    
    User->>PLC: DELETE /products/{product}/like
    PLC->>PLC: Check if Product is ACTIVE & APPROVED
    alt Valid Product
        PLC->>ULA: execute(Product, User)
        ULA->>DB: Delete from product_likes where user_id & product_id match
        ULA->>DB: Decrement Product likes_count
        PLC-->>User: 200 OK (is_liked = false, likes_count)
    else Invalid Product
        PLC-->>User: 404 Not Found
    end
"""
    unlike_product_img_path = "unlike_product_diagram.png"
    download_image(generate_kroki_url(unlike_product_mermaid), unlike_product_img_path)
    doc.add_picture(unlike_product_img_path, width=Inches(5.5))
    diagrams.append(unlike_product_img_path)

    # 4. View Liked Products Flow
    doc.add_heading('Sequence Diagram: View Liked Products Flow', level=5)
    view_liked_mermaid = """sequenceDiagram
    participant User
    participant PLC as ProductLikeController
    participant PR as ProductRepository
    participant DB as Database
    
    User->>PLC: GET /me/likes/products
    PLC->>PR: likedProductsPaginate(User)
    PR->>DB: Query Products INNER JOIN product_likes
    PR->>DB: Filter where user_id = User.id
    PR->>DB: Order by product_likes.created_at DESC
    DB-->>PR: Paginated Products Collection
    PR-->>PLC: Formatted Collection
    PLC-->>User: 200 OK (Paginated Liked Products)
"""
    view_liked_img_path = "view_liked_diagram.png"
    download_image(generate_kroki_url(view_liked_mermaid), view_liked_img_path)
    doc.add_picture(view_liked_img_path, width=Inches(5.5))
    diagrams.append(view_liked_img_path)

    # 5. Like Banner Flow
    doc.add_heading('Sequence Diagram: Like Banner Flow', level=5)
    like_banner_mermaid = """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant BR as BannerRepository
    participant DB as Database
    
    Customer->>CBC: POST /banners/{banner}/like
    CBC->>BR: like(Banner, User)
    BR->>DB: Insert into banner_likes
    BR->>DB: Increment Banner likes_count
    BR-->>CBC: Updated Banner Instance
    CBC-->>Customer: 200 OK (is_liked = true, likes_count)
"""
    like_banner_img_path = "like_banner_diagram.png"
    download_image(generate_kroki_url(like_banner_mermaid), like_banner_img_path)
    doc.add_picture(like_banner_img_path, width=Inches(5.5))
    diagrams.append(like_banner_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Additional Likes System diagrams appended to {filename}")

if __name__ == "__main__":
    append_more_likes_diagrams()
