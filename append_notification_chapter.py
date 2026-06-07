import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_notification_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== NOTIFICATION SYSTEM FEATURE ==================
    doc.add_heading('1.1.19 Notification & Real-Time System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Notification system is an internal multi-channel messaging engine designed to alert both Users and Administrators of important events. "
        "Unlike the default Laravel Notification system, this platform implements a custom NotificationService and a custom Notification model. "
        "The service checks the receiving user's preferences dynamically before fanning out the message to registered 'Notifiers' (InAppNotifier, PushNotifier, EmailNotifier, SmsNotifier). "
        "If pushed to 'in_app', it persists the record in the database and broadcasts an event to a private real-time WebSocket channel. "
        "If pushed to FCM, it delegates to the PushNotifier to deliver background payload directly to the user's registered devices."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-129: User View Notification History - User retrieves a paginated list of their persisted in-app notifications via the NotificationController.",
        "UC-130: User Manage Notifications - User marks notifications as read/unread or clears their notification history.",
        "UC-131: Admin View Notification History - Administrator views their dedicated notification inbox via the AdminNotificationController.",
        "UC-132: Dispatch Multi-Channel Notification - A domain event triggers the NotificationService, which checks user preferences and dynamically routes the message to In-App, Push, SMS, or Email notifiers.",
        "UC-133: Send Bulk Notification - An automated or admin-triggered event uses the NotificationService to iterate over a collection of users to deliver mass notifications and aggregates a success/failure report."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-129 & UC-130: User Manage Notifications ---
    add_diagram(doc, "User Manage Notifications (UC-129, UC-130)", """sequenceDiagram
    participant User
    participant NC as NotificationController
    participant DB as Database

    alt UC-129: View History
        User->>NC: GET /notifications?page=1
        NC->>DB: Query User Notifications
        DB-->>NC: Paginated Results + Unread Count
        NC-->>User: 200 OK (JSON)
    else UC-130: Manage
        User->>NC: PATCH /notifications/read-all
        NC->>DB: UPDATE notifications SET read_at = NOW() WHERE user_id = User
        NC-->>User: 200 OK
        
        User->>NC: DELETE /notifications/clear-all
        NC->>DB: DELETE FROM notifications WHERE user_id = User
        NC-->>User: 200 OK
    end
""", "sd_user_notifications", diagrams, 5.5)

    # --- UC-131: Admin View Notification History ---
    add_diagram(doc, "Admin View Notification History (UC-131)", """sequenceDiagram
    participant Admin
    participant ANC as AdminNotificationController
    participant DB as Database

    Admin->>ANC: GET /admin/notifications
    ANC->>DB: Query Admin Notifications
    DB-->>ANC: Paginated Results
    ANC-->>Admin: 200 OK (JSON)
""", "sd_admin_notifications", diagrams, 5.0)

    # --- UC-132: Dispatch Multi-Channel Notification ---
    add_diagram(doc, "Dispatch Multi-Channel Notification (UC-132)", """sequenceDiagram
    participant Trigger as Event/Action
    participant NS as NotificationService
    participant IAN as InAppNotifier
    participant PN as PushNotifier
    participant Pusher as WebSocket / FCM
    participant DB as Database

    Trigger->>NS: send(User, title, msg, 'push')
    
    NS->>NS: Check User Preferences
    opt User allows Push
        NS->>DB: INSERT INTO notifications (status='pending')
        NS->>PN: send(Notification, User)
        PN->>PN: getUserDeviceTokens(User)
        PN->>Pusher: POST fcm.googleapis.com/fcm/send
        Pusher-->>PN: 200 OK
        NS->>DB: UPDATE status='sent'
        NS->>Pusher: Broadcast Event (NotificationSent)
    end
""", "sd_dispatch_notification", diagrams, 6.0)

    # --- UC-133: Send Bulk Notification ---
    add_diagram(doc, "Send Bulk Notification (UC-133)", """sequenceDiagram
    participant Trigger
    participant NS as NotificationService
    participant DB as Database

    Trigger->>NS: sendBulk(UsersCollection, title, msg)
    
    loop Every User in Collection
        NS->>NS: Check Preferences & Route to Channel
        alt Success
            NS->>DB: INSERT INTO notifications (status='sent')
        else Failure
            NS->>DB: INSERT INTO notifications (status='failed', failure_reason)
        end
    end
    
    NS->>NS: Aggregate Sent vs Failed Counts
    NS-->>Trigger: Bulk Delivery Report Array
""", "sd_bulk_notification", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The custom notification implementation stores detailed payloads and states for each message."
    )
    schemas = [
        "notifications: Core table. Fields: id, user_id, type, title, message, data (JSON), channel (enum), status (pending/sent/failed), reference_type, reference_id, sent_at, read_at."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ NOTIFICATIONS : "receives"
    
    NOTIFICATIONS {
        bigint id PK
        bigint user_id FK
        string channel "in_app, push, email, sms"
        string status "pending, sent, failed"
        string title
        string message
        json data
        timestamp read_at "Nullable"
    }
"""
    er_img_path = "notification_er_diagram_final.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(4.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Notification System chapter accurately appended to {filename}")

if __name__ == "__main__":
    append_notification_chapter()
