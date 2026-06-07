import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_product_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== PRODUCT SYSTEM FEATURE ==================
    doc.add_heading('1.1.11 Product, Revisions, Offers, Images, Variants & Attributes System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Product System is the core commerce engine. It manages the full lifecycle of a product from creation "
        "through admin approval, revision management, offer scheduling, image galleries, and configurable variants with attributes. "
        "Store Owners create products that enter a 'PENDING' approval queue. Administrators review and Approve or Reject them. "
        "After approval, any update by the Store Owner generates a 'Product Revision' — a JSON snapshot of the proposed changes — "
        "which must be independently approved before being merged back into the live product. "
        "Products support multiple images (with a designated primary), multiple variants (e.g., sizes, colors), each variant having "
        "key-value attributes (bilingual), and time-bound promotional offers with percentage or fixed discounts."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-60: List Products (Public) - Any user browses paginated products with filters (category, store, price range, keyword, min_review_score, sort).",
        "UC-61: View Product Detail - User views a single product with its images, variants, attributes, store info, and personalized is_favorite/is_liked flags.",
        "UC-62: Create Product - Store Owner submits a new product with bilingual name/description, pricing, SKU, stock, images, and variants. Product enters PENDING approval.",
        "UC-63: Update Product - Store Owner updates an existing product's data, images, and variants.",
        "UC-64: Delete Product - Store Owner soft-deletes a product.",
        "UC-65: View My Products - Store Owner lists their own products with status/approval filters.",
        "UC-66: Admin List Products - Admin views all products with advanced filtering (status, approval_status, store, category).",
        "UC-67: Admin Approve Product - Admin approves a pending product, setting approval_status=APPROVED.",
        "UC-68: Admin Reject Product - Admin rejects a product with a mandatory rejection_reason.",
        "UC-69: Admin Suspend Product - Admin sets a live product's status to INACTIVE.",
        "UC-70: Admin Restore Product - Admin re-activates a suspended product back to ACTIVE.",
        "UC-71: Admin List Revisions - Admin views all pending Product Revisions awaiting approval.",
        "UC-72: Admin Approve Revision - Admin approves a revision; the revised_data JSON is merged back into the live product.",
        "UC-73: Admin Reject Revision - Admin rejects a revision with a reason.",
        "UC-74: List Product Offers - User views all active, date-valid promotional offers on a product.",
        "UC-75: Create Product Offer - Store Owner creates a time-bound discount offer (percentage or fixed) on their product.",
        "UC-76: Update Product Offer - Store Owner modifies an existing offer's details or validity window.",
        "UC-77: Delete Product Offer - Store Owner removes an offer.",
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-60: List Products ---
    add_diagram(doc, "List Products (UC-60)", """sequenceDiagram
    participant User
    participant PC as ProductController
    participant DB as Database

    User->>PC: GET /products?category=5&sort=price_asc
    PC->>PC: Validate Filters (category, store, price, keyword)
    PC->>DB: Query Products (ACTIVE, APPROVED)
    PC->>DB: Apply Filters & Sorting
    PC->>DB: Eager Load -> images, store, category
    DB-->>PC: Paginated Collection
    PC-->>User: 200 OK (Products + Meta)
""", "sd_list_products", diagrams)

    # --- UC-61: View Product Detail ---
    add_diagram(doc, "View Product Detail (UC-61)", """sequenceDiagram
    participant User
    participant PC as ProductController
    participant DB as Database

    User->>PC: GET /products/{product}
    PC->>DB: Find Product (ACTIVE, APPROVED)
    PC->>DB: Eager Load -> images, variants.attributes, store, category
    alt User is Authenticated
        PC->>DB: Check product_favorites pivot (is_favorite)
        PC->>DB: Check product_likes pivot (is_liked)
    end
    PC-->>User: 200 OK (Full Product Detail)
""", "sd_view_product", diagrams)

    # --- UC-62: Create Product ---
    add_diagram(doc, "Create Product (UC-62)", """sequenceDiagram
    participant Owner as Store Owner
    participant PC as ProductController
    participant DB as Database

    Owner->>PC: POST /products (name, price, images[], variants[])
    PC->>PC: Verify Store Ownership
    PC->>PC: Validate Input (bilingual fields, pricing, SKU)
    PC->>DB: Insert into products (status=ACTIVE, approval=PENDING)
    PC->>PC: syncImages(product, images[])
    loop Each Image
        PC->>DB: Insert into product_images
    end
    PC->>PC: syncVariants(product, variants[])
    loop Each Variant
        PC->>DB: Insert into product_variants
        loop Each Attribute
            PC->>DB: Insert into variant_attributes
        end
    end
    PC-->>Owner: 201 Created (Product Data)
""", "sd_create_product", diagrams, 6.0)

    # --- UC-63: Update Product ---
    add_diagram(doc, "Update Product (UC-63)", """sequenceDiagram
    participant Owner as Store Owner
    participant PC as ProductController
    participant DB as Database

    Owner->>PC: PUT /products/{product}
    PC->>PC: Verify Store Ownership
    PC->>PC: Validate Updated Fields
    PC->>DB: Update products record
    alt Images Provided
        PC->>DB: Delete existing product_images
        PC->>DB: Insert new product_images
    end
    alt Variants Provided
        PC->>DB: Delete existing variants & attributes
        PC->>DB: Insert new product_variants & variant_attributes
    end
    PC-->>Owner: 200 OK (Updated Product)
""", "sd_update_product", diagrams)

    # --- UC-64: Delete Product ---
    add_diagram(doc, "Delete Product (UC-64)", """sequenceDiagram
    participant Owner as Store Owner
    participant PC as ProductController
    participant DB as Database

    Owner->>PC: DELETE /products/{product}
    PC->>PC: Verify Store Ownership
    PC->>DB: Soft Delete Product (set deleted_at)
    PC-->>Owner: 200 OK (Product Deleted)
""", "sd_delete_product", diagrams, 4.5)

    # --- UC-65: View My Products ---
    add_diagram(doc, "View My Products (UC-65)", """sequenceDiagram
    participant Owner as Store Owner
    participant PC as ProductController
    participant DB as Database

    Owner->>PC: GET /my-products?status=active
    PC->>PC: Resolve Authenticated User's Store
    PC->>DB: Query Products where store_id = Owner's Store
    PC->>DB: Apply status/approval_status filters
    DB-->>PC: Paginated Collection
    PC-->>Owner: 200 OK (My Products + Meta)
""", "sd_my_products", diagrams)

    # --- UC-67: Admin Approve Product ---
    add_diagram(doc, "Admin Approve Product (UC-67)", """sequenceDiagram
    participant Admin
    participant PMC as ProductManagementController
    participant DB as Database

    Admin->>PMC: POST /admin/products/{product}/approve
    PMC->>DB: Find Product
    PMC->>DB: Update approval_status = APPROVED
    PMC->>DB: Set approved_at = now()
    PMC-->>Admin: 200 OK (Approved Product)
""", "sd_approve_product", diagrams, 4.5)

    # --- UC-68: Admin Reject Product ---
    add_diagram(doc, "Admin Reject Product (UC-68)", """sequenceDiagram
    participant Admin
    participant PMC as ProductManagementController
    participant DB as Database

    Admin->>PMC: POST /admin/products/{product}/reject (reason)
    PMC->>DB: Find Product
    PMC->>DB: Update approval_status = REJECTED
    PMC->>DB: Set rejection_reason, rejected_at = now()
    PMC-->>Admin: 200 OK (Rejected Product)
""", "sd_reject_product", diagrams, 4.5)

    # --- UC-69: Admin Suspend Product ---
    add_diagram(doc, "Admin Suspend Product (UC-69)", """sequenceDiagram
    participant Admin
    participant PMC as ProductManagementController
    participant DB as Database

    Admin->>PMC: POST /admin/products/{product}/suspend
    PMC->>DB: Find Product
    PMC->>DB: Update status = INACTIVE
    PMC-->>Admin: 200 OK (Suspended Product)
""", "sd_suspend_product", diagrams, 4.0)

    # --- UC-72: Admin Approve Revision ---
    add_diagram(doc, "Admin Approve Revision (UC-72)", """sequenceDiagram
    participant Admin
    participant PRC as ProductRevisionController
    participant DB as Database

    Admin->>PRC: POST /admin/revisions/{revision}/approve
    PRC->>DB: Find Revision with Parent Product
    PRC->>DB: Read revised_data JSON snapshot
    PRC->>DB: Merge revised_data into live Product record
    PRC->>DB: Update Revision approval_status = APPROVED
    PRC->>DB: Set reviewed_by = Admin, reviewed_at = now()
    PRC-->>Admin: 200 OK (Updated Product)
""", "sd_approve_revision", diagrams)

    # --- UC-73: Admin Reject Revision ---
    add_diagram(doc, "Admin Reject Revision (UC-73)", """sequenceDiagram
    participant Admin
    participant PRC as ProductRevisionController
    participant DB as Database

    Admin->>PRC: POST /admin/revisions/{revision}/reject (reason)
    PRC->>DB: Find Revision
    PRC->>DB: Update approval_status = REJECTED
    PRC->>DB: Set rejection_reason, reviewed_by, reviewed_at
    PRC-->>Admin: 200 OK (Rejected Revision)
""", "sd_reject_revision", diagrams, 4.5)

    # --- UC-75: Create Product Offer ---
    add_diagram(doc, "Create Product Offer (UC-75)", """sequenceDiagram
    participant Owner as Store Owner
    participant POC as ProductOfferController
    participant DB as Database

    Owner->>POC: POST /products/{product}/offers
    POC->>POC: Verify Store Ownership
    POC->>POC: Validate (title, discount_type, value, dates)
    POC->>DB: Insert into product_offers
    POC-->>Owner: 201 Created (Offer Data)
""", "sd_create_offer", diagrams, 4.5)

    # --- UC-74: List Product Offers ---
    add_diagram(doc, "List Product Offers (UC-74)", """sequenceDiagram
    participant User
    participant POC as ProductOfferController
    participant DB as Database

    User->>POC: GET /products/{product}/offers
    POC->>DB: Query product_offers where product_id
    POC->>DB: Filter -> is_active = true
    POC->>DB: Filter -> valid_from <= now AND valid_until >= now
    DB-->>POC: Active Offers Collection
    POC-->>User: 200 OK (Offers List)
""", "sd_list_offers", diagrams)

    # --- UC-77: Delete Product Offer ---
    add_diagram(doc, "Delete Product Offer (UC-77)", """sequenceDiagram
    participant Owner as Store Owner
    participant POC as ProductOfferController
    participant DB as Database

    Owner->>POC: DELETE /products/{product}/offers/{offer}
    POC->>POC: Verify Store Ownership
    POC->>DB: Delete product_offers record
    POC-->>Owner: 200 OK (Offer Deleted)
""", "sd_delete_offer", diagrams, 4.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The Product system uses a normalized set of tables to separate core product data from its visual assets, "
        "configurable variations, promotional offers, and revision history."
    )
    schemas = [
        "products: Core table. Fields: id, store_id (FK), category_id (FK), name_en, name_ar, slug, description_en, description_ar, brand, price, compare_at_price, discount_percentage, sku, barcode, stock_quantity, status (active/inactive), approval_status (pending/approved/rejected), rejection_reason, favorites_count, likes_count, rating_avg, rating_count.",
        "product_images: Gallery table. Fields: id, product_id (FK), image_url, alt_text, sort_order, is_primary (boolean).",
        "product_variants: Variant configurations. Fields: id, product_id (FK), name_en, name_ar, sku, price (override), stock_quantity, is_active.",
        "variant_attributes: Key-value pairs per variant. Fields: id, variant_id (FK), attribute_name_en, attribute_name_ar, attribute_value_en, attribute_value_ar.",
        "product_revisions: Proposed changes snapshot. Fields: id, product_id (FK), revised_data (JSON), submitted_by (FK), approval_status, rejection_reason, reviewed_by (FK), reviewed_at.",
        "product_offers: Time-bound discounts. Fields: id, product_id (FK), title_en, title_ar, description, discount_type (percentage/fixed), discount_value, valid_from, valid_until, is_active."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    STORES ||--o{ PRODUCTS : "owns"
    CATEGORIES ||--o{ PRODUCTS : "categorizes"
    PRODUCTS ||--o{ PRODUCT_IMAGES : "has_gallery"
    PRODUCTS ||--o{ PRODUCT_VARIANTS : "has_variants"
    PRODUCT_VARIANTS ||--o{ VARIANT_ATTRIBUTES : "has_attributes"
    PRODUCTS ||--o{ PRODUCT_REVISIONS : "has_revisions"
    PRODUCTS ||--o{ PRODUCT_OFFERS : "has_offers"

    PRODUCTS {
        bigint id PK
        uuid store_id FK
        bigint category_id FK
        string name_en
        string name_ar
        decimal price
        string status
        string approval_status
    }
    PRODUCT_IMAGES {
        bigint id PK
        bigint product_id FK
        string image_url
        bool is_primary
    }
    PRODUCT_VARIANTS {
        bigint id PK
        bigint product_id FK
        string name_en
        decimal price
        int stock_quantity
    }
    VARIANT_ATTRIBUTES {
        bigint id PK
        bigint variant_id FK
        string attribute_name_en
        string attribute_value_en
    }
    PRODUCT_REVISIONS {
        bigint id PK
        bigint product_id FK
        json revised_data
        string approval_status
    }
    PRODUCT_OFFERS {
        bigint id PK
        bigint product_id FK
        string discount_type
        decimal discount_value
        timestamp valid_until
    }
"""
    er_img_path = "product_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Product System chapter appended to {filename}")

if __name__ == "__main__":
    append_product_chapter()
