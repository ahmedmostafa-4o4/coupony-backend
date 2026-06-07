import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_nearby_offers_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== NEARBY OFFERS SYSTEM FEATURE ==================
    doc.add_heading('1.1.18 Nearby Offers System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Nearby Offers System provides geospatial intelligence, allowing users to discover active product offers sorted strictly by physical proximity to their current location. "
        "The core functionality is powered by the Haversine formula implemented directly in raw SQL via the NearbyService. "
        "Because stores can have multiple branches (linked via the polymorphic 'addressables' pivot to the 'addresses' table), the system dynamically calculates the distance "
        "from the user's provided latitude and longitude to every single active branch belonging to a store that has active offers. "
        "To optimize for front-end rendering, the system separates the logic into two distinct endpoints: one returning full paginated offer data for list views, "
        "and another returning a heavily stripped-down payload of coordinates specifically designed for rendering Map Markers without latency."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-126: List Nearby Offers - User provides coordinates and an optional search radius (e.g., 50km). The system executes the Haversine formula, filtering out offers beyond the radius, and returns a paginated list of full offer details sorted closest-first.",
        "UC-127: Load Map Markers - User opens a map interface. The system calculates proximity but returns only essential data (Latitude, Longitude, Store Name, and Offer ID) to ensure fast rendering of hundreds of map pins simultaneously."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-126: List Nearby Offers ---
    add_diagram(doc, "List Nearby Offers (UC-126)", """sequenceDiagram
    participant User
    participant NOC as NearbyOffersController
    participant NS as NearbyService
    participant DB as Database

    User->>NOC: GET /nearby/offers?lat=X&lng=Y&radius=50
    NOC->>NS: getNearbyOffers(lat, lng, radius, category_id)
    
    NS->>DB: Build Base Query (products, product_offers, stores)
    NS->>DB: Join `addressables` (owner_type='Store')
    NS->>DB: Join `addresses`
    
    Note over NS,DB: Inject Raw SQL Haversine Formula
    NS->>DB: SELECT (6371 * acos(cos(lat)*cos(lat)...)) AS distance
    NS->>DB: HAVING distance <= 50
    NS->>DB: ORDER BY distance ASC
    
    DB-->>NS: Paginated Collection
    NS-->>NOC: Result Array
    NOC-->>User: 200 OK (NearbyOfferResource)
""", "sd_list_nearby_offers", diagrams, 6.0)

    # --- UC-127: Load Map Markers ---
    add_diagram(doc, "Load Map Markers (UC-127)", """sequenceDiagram
    participant User
    participant NOC as NearbyOffersController
    participant NS as NearbyService
    participant DB as Database

    User->>NOC: GET /nearby/map-markers?lat=X&lng=Y&radius=50
    NOC->>NS: getMapMarkers(lat, lng, radius, category_id)
    
    NS->>DB: Execute Identical Haversine Join Query
    Note over NS: Exclude heavy eager loading (No Images/Descriptions)
    NS->>DB: Select ONLY (distance, lat, lng, store_name, offer_id)
    
    DB-->>NS: Lightweight Array Collection
    NS-->>NOC: Result Array
    NOC-->>User: 200 OK (MapMarkerResource)
""", "sd_load_map_markers", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The Nearby System acts as an aggregation layer spanning 5 core tables to execute its geographical queries."
    )
    schemas = [
        "products & product_offers: Filtered to ensure only active, valid promotions are shown.",
        "stores: The entity that owns the products.",
        "addressables: The polymorphic pivot resolving the many-to-many relationship between stores and their physical branches.",
        "addresses: Contains the crucial 'latitude' and 'longitude' columns required to execute the raw SQL Haversine calculation."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    PRODUCTS ||--|| PRODUCT_OFFERS : "has"
    STORES ||--o{ PRODUCTS : "sells"
    STORES ||--o{ ADDRESSABLES : "morphs_to"
    ADDRESSABLES ||--|| ADDRESSES : "maps_to"

    PRODUCTS {
        uuid id PK
        string status
    }
    PRODUCT_OFFERS {
        uuid id PK
        uuid product_id FK
    }
    STORES {
        uuid id PK
        string name
    }
    ADDRESSABLES {
        bigint id PK
        char owner_id FK "UUID"
        string owner_type
        bigint address_id FK
    }
    ADDRESSES {
        bigint id PK
        decimal latitude
        decimal longitude
    }
"""
    er_img_path = "nearby_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Nearby Offers System chapter appended to {filename}")

if __name__ == "__main__":
    append_nearby_offers_chapter()
