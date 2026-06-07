import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_more_review_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    doc.add_heading('Additional Sequence Diagrams for Review System', level=4)
    
    diagrams = []
    
    # 3. Reply to Review Flow
    doc.add_heading('Sequence Diagram: Reply to Review Flow', level=5)
    reply_mermaid = """sequenceDiagram
    participant User
    participant PCC as ProductCommentController
    participant DB as Database
    
    User->>PCC: POST /products/{product}/comments/{comment}/reply (body)
    PCC->>PCC: Validate Input (Rating PROHIBITED)
    PCC->>PCC: Check Parent Comment is visible
    alt Valid Request
        PCC->>DB: Insert into product_comments (body, parent_id, user_id)
        PCC-->>User: 201 Created (Reply Data)
    else Invalid Request
        PCC-->>User: 422 Unprocessable Entity
    end
"""
    reply_img_path = "reply_review_diagram.png"
    download_image(generate_kroki_url(reply_mermaid), reply_img_path)
    doc.add_picture(reply_img_path, width=Inches(5.0))
    diagrams.append(reply_img_path)

    # 4. Update Review Flow
    doc.add_heading('Sequence Diagram: Update Review Flow', level=5)
    update_mermaid = """sequenceDiagram
    participant User
    participant PCC as ProductCommentController
    participant DB as Database
    
    User->>PCC: PUT /comments/{comment}
    PCC->>PCC: Check if User owns the comment
    alt isTopLevelReview() == true
        PCC->>PCC: Require Rating
    else isReply() == true
        PCC->>PCC: Prohibit Rating
    end
    PCC->>DB: Update Comment Data
    PCC-->>User: 200 OK (Updated Comment)
"""
    update_img_path = "update_review_diagram.png"
    download_image(generate_kroki_url(update_mermaid), update_img_path)
    doc.add_picture(update_img_path, width=Inches(5.0))
    diagrams.append(update_img_path)

    # 5. View Paginated Reviews Flow
    doc.add_heading('Sequence Diagram: View Paginated Reviews Flow', level=5)
    view_reviews_mermaid = """sequenceDiagram
    participant Frontend
    participant PCC as ProductCommentController
    participant DB as Database
    
    Frontend->>PCC: GET /products/{product}/comments?page=1
    PCC->>DB: Query ProductComments
    PCC->>DB: Filter -> topLevel(), visible()
    PCC->>DB: Eager Load -> user.profile, visibleReplies, likesCount
    PCC->>DB: Order -> latest()
    DB-->>PCC: Paginated Collection
    PCC-->>Frontend: 200 OK (Reviews + Nested Replies)
"""
    view_reviews_img_path = "view_reviews_diagram.png"
    download_image(generate_kroki_url(view_reviews_mermaid), view_reviews_img_path)
    doc.add_picture(view_reviews_img_path, width=Inches(5.5))
    diagrams.append(view_reviews_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Additional Review System diagrams appended to {filename}")

if __name__ == "__main__":
    append_more_review_diagrams()
