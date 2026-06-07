import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_review_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== REVIEW SYSTEM FEATURE ==================
    doc.add_heading('1.1.7 Review System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Review System manages user feedback on Products and Stores through a robust hierarchical commenting engine. "
        "It distinguishes between 'Top-Level Reviews' (which require a 1-5 rating) and 'Replies' (which prohibit a rating and act as threaded discussions). "
        "Users can submit, update, or delete their own reviews. To maintain community standards, Administrators and Store Owners have the authority to 'Hide' inappropriate reviews from public visibility without permanently deleting them. "
        "The system relies on 'product_comments' and 'store_comments' tables which support self-referencing relationships (via a parent_id) to create infinite or constrained reply threads."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-41: Submit Review - User submits a top-level review on a Product or Store, providing a mandatory rating (1-5) and a text body.",
        "UC-42: Reply to Review - User replies to an existing review to start a discussion. Ratings are strictly prohibited in replies.",
        "UC-43: Update Review - User modifies their previously submitted review (updating rating/body) or reply.",
        "UC-44: Delete Review - User deletes their own review, or an Admin force-deletes a review.",
        "UC-45: Hide Review - A Store Owner or Administrator marks a review as 'HIDDEN', preserving the data but removing it from public API responses."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Submit Product Review Flow
    doc.add_heading('Sequence Diagram: Submit Review Flow', level=5)
    submit_review_mermaid = """sequenceDiagram
    participant User
    participant PCC as ProductCommentController
    participant DB as Database
    
    User->>PCC: POST /products/{product}/comments (rating=5, body="Great!")
    PCC->>PCC: Validate Input (Rating is required)
    PCC->>PCC: Check Product is ACTIVE & APPROVED
    alt Valid Request
        PCC->>DB: Insert into product_comments (rating, body, user_id)
        PCC-->>User: 201 Created (Review Data)
    else Validation Failed
        PCC-->>User: 422 Unprocessable Entity
    end
"""
    submit_review_img_path = "submit_review_diagram.png"
    download_image(generate_kroki_url(submit_review_mermaid), submit_review_img_path)
    doc.add_picture(submit_review_img_path, width=Inches(5.0))
    diagrams.append(submit_review_img_path)

    # 2. Hide Review Flow
    doc.add_heading('Sequence Diagram: Hide Review Flow', level=5)
    hide_review_mermaid = """sequenceDiagram
    participant Owner as Store Owner
    participant PCC as ProductCommentController
    participant DB as Database
    
    Owner->>PCC: POST /comments/{comment}/hide
    PCC->>DB: Load Comment & Product & Store
    PCC->>PCC: Authorize (Is Admin OR Is Store Owner?)
    alt Authorized
        PCC->>DB: Update status='HIDDEN', hidden_by=Owner.id, hidden_at=now()
        PCC-->>Owner: 200 OK (Hidden Comment Data)
    else Unauthorized
        PCC-->>Owner: 403 Forbidden
    end
"""
    hide_review_img_path = "hide_review_diagram.png"
    download_image(generate_kroki_url(hide_review_mermaid), hide_review_img_path)
    doc.add_picture(hide_review_img_path, width=Inches(5.0))
    diagrams.append(hide_review_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The Review system uses dedicated tables for Product and Store comments, supporting hierarchical data via self-referencing foreign keys."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "product_comments: Fields include id, product_id (FK), user_id (FK), parent_id (FK to self, nullable for replies), rating (tinyint, 1-5, nullable for replies), body (text), status (enum: visible/hidden), hidden_by (FK), hidden_at (timestamp).",
        "store_comments: Mirrors product_comments structure but links to store_id (FK)."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ PRODUCT_COMMENTS : "writes"
    PRODUCTS ||--o{ PRODUCT_COMMENTS : "has_reviews"
    PRODUCT_COMMENTS ||--o{ PRODUCT_COMMENTS : "has_replies (parent_id)"
    USERS ||--o{ STORE_COMMENTS : "writes"
    STORES ||--o{ STORE_COMMENTS : "has_reviews"
    STORE_COMMENTS ||--o{ STORE_COMMENTS : "has_replies (parent_id)"
    
    PRODUCT_COMMENTS {
        bigint id PK
        bigint product_id FK
        char user_id FK
        bigint parent_id FK "Nullable"
        tinyint rating "1-5, Nullable"
        text body
        string status "VISIBLE/HIDDEN"
        char hidden_by FK "Nullable"
    }
"""
    er_img_path = "reviews_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Review System chapter appended to {filename}")

if __name__ == "__main__":
    append_review_chapter()
