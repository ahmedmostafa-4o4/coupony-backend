import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Architecture Diagram: {title}', level=3)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_infrastructure_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # Remove the generic Chapter 2 if it exists
    in_chapter_2 = False
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        if p.text == 'Chapter 2: Infrastructure & Background Processing':
            in_chapter_2 = True
        
        if in_chapter_2:
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    # Re-add Chapter 2 with ACTUAL backend logic
    doc.add_heading('Chapter 2: Infrastructure & Background Processing', level=1)
    
    intro_text = (
        "While Chapter 1 covers user-facing features, Chapter 2 outlines the backend infrastructure that ensures the platform remains highly performant, automated, and scalable. "
        "The architecture leverages heavily customized caching strategies, diverse queue prioritizations, and standardized data storage practices."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 2.1 Scheduled Tasks ------------------
    doc.add_heading('2.1 Scheduled Tasks (Cron Jobs)', level=2)
    doc.add_paragraph("The Laravel Scheduler, orchestrated through `routes/console.php`, manages daily automations. Notably, many Subscription commands reside in the `Jobs` domain directory but technically extend `Illuminate\\Console\\Command`.")
    tasks = [
        "Subscription Lifecycle Management (Runs Daily): Automatically transitions subscriptions through various states via commands like `subscription:transition-to-grace`, `subscription:transition-to-degraded`, and `subscription:transition-to-suspended`.",
        "Subscription Expiration Alerts (Runs Daily): `subscription:send-expiring-notifications` triggers alerts for owners whose plans are near expiration.",
        "Daily Analytics Summaries (Runs Daily at 22:00): Computes and aggregates the day's analytical data into structured formats.",
        "Milestone Checks (Runs Daily at 22:30): `notifications:check-milestones` detects if users or stores have reached targeted milestones.",
        "Housekeeping Tasks: Custom commands like `CleanupExpiredOtps`, `PonyEmbedProducts`, and `PonyPurgeImageQueries` maintain data integrity and AI state."
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    # ------------------ 2.2 Queue Workers ------------------
    doc.add_heading('2.2 Queue Workers & Asynchronous Processing', level=2)
    doc.add_paragraph("The application utilizes Laravel Queues with `database` as the default driver, but scales up using prioritized Redis queues configured in `config/queue.php`:")
    jobs = [
        "redis-high (Retry: 90s): Critical priority lane for Payments and Orders to ensure zero drops.",
        "redis-default (Retry: 300s): Normal priority lane for Emails, Notifications, and standard Event Listeners (e.g., `SendWelcomeEmail`, `SendOfferRedeemedNotifications`).",
        "redis-low (Retry: 600s): Low priority lane for Analytics generation and machine learning tasks (e.g., `RegenerateProductEmbeddingsJob`, `PurgeOldPonyImageUploadsJob`)."
    ]
    for j in jobs:
        doc.add_paragraph(j, style='List Bullet')

    # ------------------ 2.3 Caching Strategy ------------------
    doc.add_heading('2.3 Global Caching Strategy', level=2)
    doc.add_paragraph("The platform aggressively caches high-read data to reduce database querying, utilizing `database` by default and dedicated Redis stores (`redis-sessions`, `redis-recommendations`).")
    caches = [
        "Repositories Caching (TTL: 1 hour): UserRepository and StoreRepository cache single record lookups (e.g., `user.by_id.{id}`).",
        "Heavy Analytics Caching: GetProductAnalyticsAction and GetSellerDashboardAction extensively cache expensive aggregation queries (e.g., `seller_analytics:{store_id}:{period}`).",
        "Banners Caching (TTL: 15 minutes): Active banners stored in `CUSTOMER_BANNERS_CACHE_KEY`.",
        "PonyAI Embedding Caching: Vector embeddings and AI extracted intents are cached to significantly reduce costly API calls to Google Gemini."
    ]
    for c in caches:
        doc.add_paragraph(c, style='List Bullet')

    # ------------------ 2.4 Database & Storage ------------------
    doc.add_heading('2.4 Database & Storage Standards', level=2)
    db_standards = [
        "Primary Keys (UUIDs): The system uniformly uses UUIDs across major tables. Notably, it avoids Laravel's default `HasUuids` trait. Instead, generation happens directly in the database (`DEFAULT (UUID())`) and inside model `static::creating` events via `Str::uuid()`.",
        "Data Retention (Soft Deletes): Applied extensively to `users`, `stores`, `products`, and `comments` at the schema level. However, for `users` and `stores`, the Eloquent `SoftDeletes` trait itself is intentionally omitted from the model, managing visibility strictly through explicit logic.",
        "Cloud Storage Config: Configures `local`, `public`, and AWS `s3` disks, using `local` by default to simplify developer setups."
    ]
    for d in db_standards:
        doc.add_paragraph(d, style='List Bullet')

    # ======================== ARCHITECTURE DIAGRAMS ========================
    doc.add_heading('System Architecture Diagrams', level=2)

    # --- Async Architecture Diagram ---
    add_diagram(doc, "Asynchronous Queue & Scheduler Architecture", """sequenceDiagram
    participant OS as Server OS (Cron)
    participant Kernel as Console Kernel (routes/console.php)
    participant API as API Controllers / Event Listeners
    participant Redis as Redis Queues
    participant Workers as Queue Workers
    participant S3 as AWS S3 / Local Storage
    participant External as 3rd Party APIs (FCM/Gemini)

    Note over OS,Kernel: 1. Scheduled Tasks
    OS->>Kernel: * * * * * (Every Minute)
    Kernel->>Kernel: Check Schedule List
    Kernel->>Database: subscription:transition-to-degraded (Daily)
    
    Note over API,Workers: 2. Background Queue Processing
    API->>API: Fire Domain Event (e.g., NotificationSent)
    API->>Redis: Dispatch Listener (redis-default queue)
    API-->>Client: 200 OK (Instant Response)
    
    Redis->>Workers: Pop Job (Priority execution)
    Workers->>External: Execute API Requests asynchronously
""", "sd_infra_async_fixed", diagrams, 6.0)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Infrastructure Chapter correctly appended to {filename}")

if __name__ == "__main__":
    fix_infrastructure_chapter()
