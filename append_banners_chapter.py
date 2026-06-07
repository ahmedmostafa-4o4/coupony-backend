import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_banners_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== BANNERS SYSTEM FEATURE ==================
    doc.add_heading('1.1.16 Banners Management System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Banners Management System governs the creation, review, display, and analytics tracking of promotional banners. "
        "It operates on a strict tri-actor model: Store Owners create banners linking to their product offers; "
        "Administrators act as gatekeepers, reviewing and setting display priorities; and Customers interact with active banners. "
        "To ensure high performance while tracking rich analytics, the system uses an event-driven interaction architecture. "
        "When a customer views, clicks, likes, or claims a banner, a dedicated InteractWithBannerAction records the raw interaction event in a pivot table "
        "while simultaneously incrementing denormalized counter columns on the main Banner model. "
        "The system also features parallel architecture for standard retail banners and specialized Travel Banners."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-115: View Active Banners - Customers retrieve prioritized, active banners. The system eager-loads associated offers/destinations.",
        "UC-116: Interact with Banner - Customers view, click, like, share, or claim a banner. The system inserts an interaction record and increments the banner's counter.",
        "UC-117: Create Store Banner - Store Owners submit banner images and link product offers. The banner defaults to PENDING_REVIEW.",
        "UC-118: Manage Store Banners - Store Owners edit or delete banners. Editing an ACTIVE banner safely downgrades its status back to PENDING_REVIEW.",
        "UC-119: View Banner Analytics - Store Owners retrieve aggregated interaction counters (views, clicks, claims) to gauge campaign performance.",
        "UC-120: Admin Review Banner - Administrators approve or reject pending banners, providing rejection reasons if necessary.",
        "UC-121: Admin Manage Banner Priority - Administrators manually dictate the display sorting order of approved active banners."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-115: View Active Banners ---
    add_diagram(doc, "View Active Banners (UC-115)", """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant DB as Database

    Customer->>CBC: GET /banners
    CBC->>DB: Query banners WHERE status = 'ACTIVE'
    CBC->>DB: ORDER BY priority ASC
    CBC->>DB: Eager Load banner_offers (Products)
    DB-->>CBC: Paginated Banners
    CBC-->>Customer: 200 OK (BannerResource)
""", "sd_view_banners", diagrams, 5.0)

    # --- UC-116: Interact with Banner ---
    add_diagram(doc, "Interact with Banner (UC-116)", """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant IBA as InteractWithBannerAction
    participant DB as Database

    Customer->>CBC: POST /banners/{id}/interact (action: 'like')
    CBC->>IBA: execute(Banner, action, User)
    
    IBA->>DB: Begin Transaction
    IBA->>DB: Check if User already 'liked' (prevent duplicates)
    alt New Interaction
        IBA->>DB: INSERT INTO banner_likes (banner_id, user_id)
        IBA->>DB: Increment Banner.likes_count
        DB-->>IBA: Commit Transaction
        IBA-->>CBC: Success
        CBC-->>Customer: 200 OK
    else Already Interacted
        IBA-->>CBC: Error Duplicate
        CBC-->>Customer: 422 Unprocessable Entity
    end
""", "sd_interact_banner", diagrams, 6.0)

    # --- UC-117 & UC-118: Create/Manage Store Banners ---
    add_diagram(doc, "Create & Edit Store Banner (UC-117, UC-118)", """sequenceDiagram
    participant Owner as Store Owner
    participant SBC as StoreBannerController
    participant Storage
    participant DB as Database

    alt Create Banner
        Owner->>SBC: POST /stores/{store}/banners (Image, Offer IDs)
        SBC->>Storage: Upload Image URL
        SBC->>DB: INSERT INTO banners (status = 'PENDING_REVIEW')
        SBC->>DB: Attach to product_offers via banner_offers
        SBC-->>Owner: 201 Created
    else Edit Active Banner
        Owner->>SBC: PUT /stores/{store}/banners/{id}
        SBC->>DB: UPDATE banner fields
        Note over DB: State Machine Rule
        SBC->>DB: UPDATE status = 'PENDING_REVIEW'
        SBC-->>Owner: 200 OK
    end
""", "sd_manage_store_banner", diagrams, 6.0)

    # --- UC-119: View Banner Analytics ---
    add_diagram(doc, "View Banner Analytics (UC-119)", """sequenceDiagram
    participant Owner as Store Owner
    participant SBC as StoreBannerController
    participant DB as Database

    Owner->>SBC: GET /stores/{store}/banners/{id}/analytics
    SBC->>DB: Fetch Banner model
    Note over SBC: Extracts denormalized counters
    SBC->>SBC: Map views_count, clicks_count, claims_count
    SBC-->>Owner: 200 OK (Analytics JSON)
""", "sd_view_banner_analytics", diagrams, 5.0)

    # --- UC-120 & UC-121: Admin Review & Priority ---
    add_diagram(doc, "Admin Review & Priority (UC-120, UC-121)", """sequenceDiagram
    participant Admin
    participant ABC as AdminBannerController
    participant RBA as ReviewBannerAction
    participant DB as Database

    alt UC-120: Review Banner
        Admin->>ABC: POST /admin/banners/{id}/review (status: ACTIVE)
        ABC->>RBA: execute(Banner, status, reason)
        RBA->>DB: UPDATE status = 'ACTIVE'
        RBA->>Owner: Dispatch Approval Notification
        ABC-->>Admin: 200 OK
    else UC-121: Set Priority
        Admin->>ABC: PATCH /admin/banners/{id}/priority (priority: 1)
        ABC->>DB: UPDATE banners SET priority = 1
        ABC-->>Admin: 200 OK
    end
""", "sd_admin_review_banner", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The schema isolates the heavy analytical tracking into separate polymorphic-style pivot tables while keeping the main entity lightweight using denormalized counters."
    )
    schemas = [
        "banners: Main entity. Fields: id, store_id, title, image_url, status (pending/active/rejected), priority, views_count, clicks_count, likes_count, claims_count.",
        "banner_offers: Links a banner to specific product offers in the system.",
        "travel_banners & travel_banner_destinations: Mirrored tables catering specifically to Travel/Tourism agencies.",
        "Interaction Tables: Dedicated tables (e.g., banner_views, banner_clicks, banner_likes) containing banner_id, user_id, and timestamps for deep behavioral tracking."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    STORES ||--o{ BANNERS : "creates"
    BANNERS ||--o{ BANNER_OFFERS : "promotes"
    PRODUCT_OFFERS ||--o{ BANNER_OFFERS : "featured_in"
    
    BANNERS ||--o{ BANNER_VIEWS : "receives"
    BANNERS ||--o{ BANNER_LIKES : "receives"
    USERS ||--o{ BANNER_VIEWS : "generates (optional)"
    USERS ||--o{ BANNER_LIKES : "generates"

    BANNERS {
        uuid id PK
        uuid store_id FK
        string status
        int priority
        int views_count
        int likes_count
    }
    BANNER_OFFERS {
        uuid banner_id FK
        uuid product_offer_id FK
    }
    BANNER_LIKES {
        uuid banner_id FK
        char user_id FK
        timestamp created_at
    }
"""
    er_img_path = "banners_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Banners System chapter appended to {filename}")

if __name__ == "__main__":
    append_banners_chapter()
