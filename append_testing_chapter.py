import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_testing_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 9: TESTING & CI/CD ==================
    doc.add_page_break()
    doc.add_heading('Chapter 9: Testing Strategy & CI/CD Pipeline', level=1)
    
    intro_text = (
        "To ensure absolute stability, the backend relies on an extensive automated testing suite and Continuous Integration (CI) gates. "
        "Every code push is automatically verified before it can be merged, ensuring no critical regressions reach production."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 9.1 Testing Framework ------------------
    doc.add_heading('9.1 Testing Framework & Coverage', level=2)
    doc.add_paragraph("The platform uses standard PHPUnit rather than Pest or Dusk, splitting the test suites into two distinct categories:")
    test = [
        "Feature Tests (HTTP Layer): Heavy integration tests that hit the actual API endpoints. The largest suites are `ProductTest.php` (91KB), `ProductRevisionRoutesTest.php`, and `AuthenticationTest.php`.",
        "Unit Tests (Domain Logic): Highly isolated tests targeting core business rules without hitting the HTTP layer. Notable suites include `SubscriptionStateMachineTest`, `ProductRevisionWorkflowTest`, and `PaymobServiceTest`."
    ]
    for t in test:
        doc.add_paragraph(t, style='List Bullet')

    # ------------------ 9.2 Factory State Generation ------------------
    doc.add_heading('9.2 Factory State Generation', level=2)
    doc.add_paragraph("Test data is dynamically generated using over 17 core Model Factories. To mock complex database scenarios rapidly, the factories employ advanced `->state()` modifiers:")
    fact = [
        "User State Simulation: `UserFactory::new()->unverified()->suspended()->create()` instantly generates a highly specific user condition without manual database seeding.",
        "Store & Product States: Factories like `ProductFactory` and `BannerFactory` contain explicit states for `active`, `pending`, and `expired` to rapidly test edge-case business logic."
    ]
    for f in fact:
        doc.add_paragraph(f, style='List Bullet')

    # ------------------ 9.3 GitHub Actions CI/CD ------------------
    doc.add_heading('9.3 GitHub Actions CI/CD Pipeline', level=2)
    doc.add_paragraph("The `.github/workflows/tests.yml` configuration acts as the strict automated gatekeeper for the repository. While deployment is handled manually, the integration pipeline operates flawlessly:")
    ci = [
        "Triggers: The pipeline automatically fires on `push` to master, any `pull_request`, and on a daily schedule (Cron).",
        "Execution Matrix: The suite is executed across an array of PHP versions (8.2, 8.3, and 8.4) to ensure absolute compatibility against future server upgrades.",
        "In-Memory Database: To guarantee the tests run rapidly in the cloud, the GitHub Action automatically sets up an SQLite in-memory database, bypassing the need for a heavy MySQL container."
    ]
    for c in ci:
        doc.add_paragraph(c, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"Testing Chapter appended to {filename}")

if __name__ == "__main__":
    append_testing_chapter()
