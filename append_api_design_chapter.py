import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_api_design_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 6: API DESIGN ==================
    doc.add_page_break()
    doc.add_heading('Chapter 6: API Design & Standardization Guidelines', level=1)
    
    intro_text = (
        "To ensure frontend, mobile, and third-party developers have a seamless integration experience, the platform enforces strict API design standards. "
        "Every endpoint adheres to a unified architecture for authentication, response formatting, localization, and rate limiting."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 6.1 Unified Response Formatting ------------------
    doc.add_heading('6.1 Unified Response Wrapper', level=2)
    doc.add_paragraph("Directly returning raw Eloquent models is prohibited. All endpoints extend a base Controller that forces responses through the `localizedJson()` wrapper. This guarantees a predictable JSON structure globally:")
    res = [
        "Success Payloads: Always wrapped in a `{ \"data\": [...] }` object.",
        "Pagination Metadata: If a response is paginated, the data wrapper is accompanied by a `{ \"meta\": { \"current_page\", \"last_page\", \"total\" } }` object to ensure the frontend can build universal pagination interceptors.",
        "Status Messaging: Action endpoints (like marking a notification as read or confirming a payment) return a `{ \"message\": \"...\" }` string translated into the user's active locale."
    ]
    for r in res:
        doc.add_paragraph(r, style='List Bullet')

    # ------------------ 6.2 Global Middleware & Localization ------------------
    doc.add_heading('6.2 Global Middleware Pipeline', level=2)
    doc.add_paragraph("The Laravel 11 `bootstrap/app.php` configuration appends a mandatory middleware pipeline to every single API request before it ever reaches a controller:")
    mid = [
        "PerformanceMonitoring: Automatically logs the execution time and memory usage of heavy endpoints.",
        "SetLocale & UseAuthenticatedUserLocale: Intercepts the `Accept-Language` header and the user's saved DB preferences, dynamically shifting the application's locale on the fly. The response automatically includes a matching `Content-Language` header.",
        "UpdateUserSession: Ensures the user's `last_active_at` timestamp is updated efficiently without bottlenecking the database."
    ]
    for m in mid:
        doc.add_paragraph(m, style='List Bullet')

    # ------------------ 6.3 Authentication & Rate Limiting ------------------
    doc.add_heading('6.3 Authentication & Rate Limiting', level=2)
    auth = [
        "Sanctum Tokens: Authentication is handled entirely statelessly via Laravel Sanctum. The `auth:sanctum` middleware is required for all secured routes, expecting a `Bearer {token}` in the Authorization header.",
        "Specialized Rate Limiting: While standard endpoints have general throttle rules, heavy computational endpoints (like PonyAI) are protected by a dedicated `pony.throttle` rate limiter to prevent abuse and API cost overrun.",
        "Subscription Gatekeeping: Specialized endpoints (like Store Analytics or Bulk Notifications) sit behind a custom `subscription` middleware that instantly rejects requests with a 403 Forbidden if the store owner's payment plan has expired."
    ]
    for a in auth:
        doc.add_paragraph(a, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"API Design Chapter appended to {filename}")

if __name__ == "__main__":
    append_api_design_chapter()
