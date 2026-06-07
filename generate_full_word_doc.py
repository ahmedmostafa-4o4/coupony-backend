import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def create_full_document():
    doc = Document()
    diagrams = []
    
    # Title
    title = doc.add_heading('Application Features Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Chapter 1
    doc.add_heading('1. System Architecture', level=1)
    doc.add_paragraph('This document outlines the core features of the system.')
    
    # Subchapter: Features
    doc.add_heading('1.1 Features', level=2)
    doc.add_paragraph('This subchapter describes the main features implemented in the repository.')
    
    # ================== AUTH FEATURE ==================
    doc.add_heading('1.1.1 Authentication Feature', level=3)
    doc.add_heading('Description of Logic', level=4)
    doc.add_paragraph("The authentication system is built using Laravel Sanctum, providing token-based authentication. "
        "The logic supports role-based access control (Admin, Seller, Customer) and verifies user roles during login. "
        "Upon providing valid email and password credentials, the system checks if the user's account is active and verified. "
        "If unverified, the system triggers an OTP (One-Time Password) generation and sends it via email. "
        "For verified users, an access token (short-lived) and a refresh token (30 days TTL) are generated. "
        "The authentication service also records session metadata such as IP address, user agent, and detected device type (mobile, tablet, desktop). "
        "Role assignment can be automatic; for example, if a user requests a 'customer' role during login and doesn't have it, it is assigned dynamically. "
        "The refresh token mechanism allows issuing new access tokens by validating the hashed refresh token stored in the database, invalidating the old session to maintain security. "
        "Additionally, users can register for an account (triggering an OTP), verify their OTP (which automatically logs them in), and authenticate via Google SSO.")
    
    doc.add_heading('Use Cases', level=4)
    for uc in [
        "UC-1: User Login - Users authenticate using email and password. Receives Access and Refresh tokens upon success.",
        "UC-2: Role-based Authorization - System validates if the user has the required role (Admin/Seller/Customer) to access specific endpoints.",
        "UC-3: Token Refresh - Client application exchanges a valid refresh token for a new access/refresh token pair without requiring user credentials.",
        "UC-4: Session Management & Logout - Users can view active sessions. They can log out of the current device (revoking the current token) or log out of all devices (revoking all tokens).",
        "UC-5: User Registration - Users can register an account. This automatically triggers an OTP code sent to their email for verification.",
        "UC-6: OTP Verification - Users submit their OTP code to verify their email. Upon successful verification, access tokens are issued and the user is logged in automatically.",
        "UC-7: Google Authentication - Users can register or log in using their Google account. The system verifies the ID token, resolves the user identity, checks onboarding status, and returns access tokens."
    ]: doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    auth_flows = [
        ('Sequence Diagram: User Login Flow', """sequenceDiagram
    participant C as Client
    participant AC as AuthController
    participant AS as AuthService
    participant DB as Database
    C->>AC: POST /auth/login (email, password, role)
    AC->>AS: login(email, password, context)
    AS->>DB: Find User by email
    DB-->>AS: User data
    AS->>AS: Verify Password
    AS->>AS: Check status & verification
    alt If unverified
        AS->>C: Return OTP Required error
    else If verified
        AS->>AS: Assert/Assign role
        AS->>AS: Generate AccessToken & RefreshToken
        AS->>DB: Store new session with device metadata
        AC-->>C: Return 200 OK with User data & Tokens
    end""", "login_diagram.png", 6.0),
        ('Sequence Diagram: Token Refresh Flow', """sequenceDiagram
    participant C as Client
    participant AC as AuthController
    participant AS as AuthService
    participant DB as Database
    C->>AC: POST /auth/refresh (refresh_token)
    AC->>AS: refreshToken(refresh_token)
    AS->>DB: Find Session by hashed token
    DB-->>AS: Session data
    AS->>AS: Check if session expired
    AS->>DB: Delete old tokens & session
    AS->>AS: Generate new AccessToken & RefreshToken
    AS->>DB: Store new session
    AC-->>C: Return new tokens""", "refresh_diagram.png", 5.0),
        ('Sequence Diagram: User Registration Flow', """sequenceDiagram
    participant C as Client
    participant RC as RegisterController
    participant RU as RegisterUser Action
    participant OS as OtpService
    participant DB as Database
    C->>RC: POST /auth/register
    RC->>RU: execute(UserData)
    RU->>DB: Create User & Profile
    RU-->>RC: User Instance
    RC->>OS: generateAndSend(VERIFY_EMAIL)
    OS->>DB: Store OTP Code
    OS->>C: Email sent with OTP
    RC-->>C: 201 Created (OTP Metadata returned)""", "register_diagram.png", 5.0),
        ('Sequence Diagram: OTP Verification Flow', """sequenceDiagram
    participant C as Client
    participant OC as OtpController
    participant OS as OtpService
    participant AS as AuthService
    participant DB as Database
    C->>OC: POST /auth/otp/verify (code, email)
    OC->>OS: verify(user, code, VERIFY_EMAIL)
    OS->>DB: Validate OTP code against DB
    alt Invalid OTP
        OS-->>OC: Verification Failed
        OC-->>C: 4xx Error Response
    else Valid OTP & Verify Email
        OS->>DB: Mark email as verified
        OS-->>OC: Verification Success
        OC->>AS: issueTokensForUser(User)
        AS->>DB: Generate & Store Session
        AS-->>OC: Access & Refresh Tokens
        OC-->>C: 200 OK (Tokens & User Data)
    end""", "otp_diagram.png", 5.0),
        ('Sequence Diagram: Google Authentication Flow', """sequenceDiagram
    participant C as Client
    participant GC as GoogleLoginController
    participant GV as GoogleTokenVerifier
    participant AS as AuthService
    participant DB as Database
    C->>GC: POST /auth/google (id_token)
    GC->>GV: verifyIdToken(id_token)
    GV->>Google: Validate Token Signature
    GV-->>GC: Google User Payload
    GC->>DB: Find User by provider_id or Email
    alt User not found
        GC->>DB: Create new User & Profile
    else User found
        GC->>DB: Link Google provider info
    end
    GC->>DB: Mark email as verified
    GC->>AS: issueTokensForUser(User)
    AS->>DB: Store session metadata
    GC->>DB: Check onboarding status
    GC-->>C: 200 OK (Tokens, User Data, Onboarding Status)""", "google_diagram.png", 6.0)
    ]
    for title, text, path, width in auth_flows:
        doc.add_heading(title, level=5)
        download_image(generate_kroki_url(text), path)
        doc.add_picture(path, width=Inches(width))
        diagrams.append(path)

    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph("The authentication logic relies on several database tables to manage user credentials, metadata, sessions, and verification processes.")
    for schema in [
        "users: Core user table storing credentials and status. Fields include id (UUID), email, password_hash, phone_number, verification timestamps, status (active/suspended/deleted), and provider details for social login.",
        "profiles: Stores extended user details. Fields include id, user_id (FK to users), first_name, last_name, date_of_birth, gender, and avatar_url.",
        "sessions: Custom session table for tracking active logins. Fields include id (UUID), user_id (FK to users), token, user_agent, ip_address, device_type, last_activity, and expires_at.",
        "otps: Manages One-Time Passwords for verification. Fields include id, user_id (nullable FK), phone_or_email, otp_hash, purpose, channel, status, attempts, and expires_at."
    ]: doc.add_paragraph(schema, style='List Bullet')

    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    auth_er_path = "auth_er_diagram.png"
    download_image(generate_kroki_url("""erDiagram
    USERS ||--o| PROFILES : "has_one"
    USERS ||--o{ SESSIONS : "has_many"
    USERS ||--o{ OTPS : "generates"
    USERS { char id PK string email }
    PROFILES { bigint id PK uuid user_id FK }
    SESSIONS { uuid id PK uuid user_id FK }
    OTPS { bigint id PK uuid user_id FK }
"""), auth_er_path)
    doc.add_picture(auth_er_path, width=Inches(5.0))
    diagrams.append(auth_er_path)


    # ================== ROLES FEATURE ==================
    doc.add_heading('1.1.2 Roles, Permissions & Staff Management', level=3)
    doc.add_heading('Description of Logic', level=4)
    doc.add_paragraph("The system provides a comprehensive Staff and Role management system divided into two layers: Admin-level Role Management and Store-level Employee Management. "
        "At the Admin level, roles and their associated permissions can be dynamically created, updated, and retrieved using the Spatie Permission package. "
        "At the Store level, store owners can invite other users to join their store as employees using the StoreInvitationController. "
        "Invitations can be sent, accepted, declined, or cancelled. Once a user accepts an invitation and becomes a store employee, the store owner can assign them specific store-related roles (e.g., store_manager, cashier, support_agent). "
        "The system ensures that users are granted the appropriate global Spatie roles based on their store assignments, and automatically cleans up roles when an employee is removed from a store. "
        "Critical actions, such as an owner attempting to remove themselves, are blocked by validation rules.")
    
    doc.add_heading('Use Cases', level=4)
    for uc in [
        "UC-8: Admin Role Management - System administrators can dynamically create, update, and delete roles and assign specific permissions.",
        "UC-9: Invite Store Employee - A store owner sends an invitation to a user to join their store staff.",
        "UC-10: Manage Invitations - Users can view, accept, or decline pending store invitations.",
        "UC-11: Manage Store Employees - Store managers can view, update roles, or remove employees from their store.",
        "UC-12: Update Employee Role - A store owner updates an existing employee's role, automatically revoking old global permissions and assigning new ones.",
        "UC-13: Remove Store Employee - A store owner removes an employee, automatically cleaning up their store-specific access rights.",
        "UC-14: Resend or Cancel Invitations - A store owner resends a pending invitation or cancels it.",
        "UC-15: Decline Store Invitation - A user explicitly declines a pending store invitation.",
        "UC-16: Admin Custom Role Creation - System Administrators create entirely new roles and attach granular Spatie permissions to them."
    ]: doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    roles_flows = [
        ('Sequence Diagram: Store Invitation Flow', """sequenceDiagram
    participant Owner as Store Owner
    participant IC as StoreInvitationController
    participant IS as StoreInvitationService
    participant DB as Database
    Owner->>IC: POST /stores/{store}/invitations
    IC->>IS: sendInvitation(store, user, data)
    IS->>DB: Check if user exists & Create StoreInvitation
    IS->>Owner: Send Notification/Email
    IC-->>Owner: 201 Created""", "invite_diagram.png", 5.0),
        ('Sequence Diagram: Accept Invitation Flow', """sequenceDiagram
    participant User
    participant IC as StoreInvitationController
    participant IS as StoreInvitationService
    participant DB as Database
    User->>IC: POST /invitations/{inv}/accept
    IC->>IS: acceptInvitation(user, invitation)
    IS->>DB: Mark accepted & Create StoreEmployee
    IS->>DB: Assign store-specific Spatie role
    IC-->>User: 200 OK""", "accept_diagram.png", 5.0),
        ('Sequence Diagram: Update Store Employee Role Flow', """sequenceDiagram
    participant Owner as Store Owner
    participant SEC as StoreEmployeeController
    participant DB as Database
    Owner->>SEC: PUT /stores/{store}/employees/{user}
    SEC->>DB: Fetch & Update StoreEmployee role
    alt If role has changed
        SEC->>DB: Check if old role used elsewhere
        alt Not used
            SEC->>DB: Revoke old Spatie Role
        end
        SEC->>DB: Assign new Spatie Role
    end
    SEC-->>Owner: 200 OK""", "update_role_diagram.png", 6.0),
        ('Sequence Diagram: Remove Store Employee Flow', """sequenceDiagram
    participant Owner as Store Owner
    participant SEC as StoreEmployeeController
    participant DB as Database
    Owner->>SEC: DELETE /stores/{store}/employees/{user}
    SEC->>DB: Prevent self-deletion
    SEC->>DB: Delete StoreEmployee record
    SEC->>DB: Check if user works at ANY other store
    alt No other store roles
        SEC->>DB: Revoke all store-specific Spatie Roles
    end
    SEC-->>Owner: 200 OK""", "remove_emp_diagram.png", 5.0),
        ('Sequence Diagram: Admin Role Creation Flow', """sequenceDiagram
    participant Admin
    participant RPC as RolePermissionController
    participant SRA as StoreRoleAction
    participant DB as Database
    Admin->>RPC: POST /admin/roles
    RPC->>SRA: execute(RoleDTO)
    SRA->>DB: Create Spatie Role & Attach Permissions
    SRA-->>RPC: Role created
    RPC-->>Admin: 201 Created""", "admin_role_diagram.png", 5.0)
    ]
    for title, text, path, width in roles_flows:
        doc.add_heading(title, level=5)
        download_image(generate_kroki_url(text), path)
        doc.add_picture(path, width=Inches(width))
        diagrams.append(path)

    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph("The feature leverages the Spatie Laravel Permission package's standard schema along with custom tables for store context.")
    for schema in [
        "roles: Standard Spatie table containing id, name, guard_name, timestamps.",
        "permissions: Standard Spatie table containing id, name, guard_name, timestamps.",
        "model_has_roles: Pivot table linking users (model_id, model_type) to roles (role_id).",
        "model_has_permissions: Pivot table linking users directly to permissions.",
        "role_has_permissions: Pivot table linking roles to permissions.",
        "store_employees: Stores the relationship between a user and a store. Fields: id, store_id (uuid, FK to stores), user_id (FK to users), timestamps. Unique constraint on [store_id, user_id].",
        "store_invitations: Manages pending invitations. Fields: id, store_id (uuid), invited_by_user_id, invitee_user_id, role (string), permissions (json), status (string, default 'pending'), message (text), expires_at, accepted_at, declined_at, timestamps.",
        "user_roles: Custom table for explicit tracking of user roles per store. Fields: id, user_id (uuid), role_id (unsigned bigint), store_id (uuid, nullable), granted_at, granted_by_user_id, expires_at."
    ]: doc.add_paragraph(schema, style='List Bullet')

    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    roles_er_path = "roles_er_diagram.png"
    download_image(generate_kroki_url("""erDiagram
    USERS ||--o{ STORE_EMPLOYEES : "has"
    STORES ||--o{ STORE_EMPLOYEES : "employs"
    USERS ||--o{ STORE_INVITATIONS : "invited_by / invitee"
    STORES ||--o{ STORE_INVITATIONS : "has"
    USERS ||--o{ MODEL_HAS_ROLES : "has"
    ROLES ||--o{ MODEL_HAS_ROLES : "assigned_to"
    ROLES ||--o{ ROLE_HAS_PERMISSIONS : "contains"
    PERMISSIONS ||--o{ ROLE_HAS_PERMISSIONS : "assigned_to"
    USERS ||--o{ USER_ROLES : "assigned"
    STORES ||--o{ USER_ROLES : "scopes"
    ROLES ||--o{ USER_ROLES : "defines"
    STORE_EMPLOYEES { bigint id PK uuid store_id FK char user_id FK }
    STORE_INVITATIONS { bigint id PK uuid store_id FK string role }
    USER_ROLES { bigint id PK uuid user_id FK }
"""), roles_er_path)
    doc.add_picture(roles_er_path, width=Inches(6.0))
    diagrams.append(roles_er_path)


    # ================== STORE FEATURE ==================
    doc.add_heading('1.1.3 Store & Profile Management', level=3)
    doc.add_heading('Description of Logic', level=4)
    doc.add_paragraph("Store & Profile Management allows users to manage their personal identities and create business entities (Stores). "
        "Users can update their personal profiles (like first name, last name, and avatar) via the '/me' endpoint. "
        "When a user creates a new Store, the system records the core store information and handles the upload of store assets (logo and banner) as well as verification documents. "
        "Store owners and managers can separately update public-facing store profiles and manage business hours, social links, and locations. "
        "To provide analytics to store owners, the system automatically logs a 'Store Profile View' record with the viewer's IP address every time a public store profile is visited. "
        "Additionally, System Administrators review submitted stores and verification documents through the Admin Dashboard, enabling them to approve, reject, or suspend stores based on platform guidelines.")
    
    doc.add_heading('Use Cases', level=4)
    for uc in [
        "UC-17: User Profile Update - A user updates their personal profile information (name, avatar, bio) via the authenticated /me endpoint.",
        "UC-18: Create Store - A user creates a new store, providing store details and securely uploading initial logos and verification documents.",
        "UC-19: Update Store Profile - A store manager explicitly updates the public-facing store profile assets (logo, banner) separate from core store business details.",
        "UC-20: View Public Store Profile - Customers view a public store profile. This action automatically logs a view record for store analytics.",
        "UC-21: Update Verification Documents - Store managers submit or replace business verification documents to achieve a 'Verified' status.",
        "UC-22: Search & Filter Public Stores - Customers browse active stores using advanced filters (category, city, verification status, rating) and sorting logic.",
        "UC-23: Admin Store Approval - System Administrators review pending store registrations and either approve or reject them based on submitted details.",
        "UC-24: Admin Verification Review - System Administrators review uploaded legal documents and mark them as approved or rejected."
    ]: doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    store_flows = [
        ('Sequence Diagram: Update User Profile Flow', """sequenceDiagram
    participant User
    participant LC as LoginController
    participant DB as Database
    participant Storage
    User->>LC: PATCH /me (first_name, avatar)
    LC->>DB: Fetch authenticated User & Profile
    alt Avatar provided
        LC->>Storage: Delete old avatar if exists
        LC->>Storage: Upload new avatar
    end
    LC->>DB: Update user_profiles table
    LC-->>User: 200 OK""", "profile_diagram.png", 5.0),
        ('Sequence Diagram: Create Store Flow', """sequenceDiagram
    participant User
    participant SC as StoreController
    participant CSA as CreateStore Action
    participant Storage
    participant DB as Database
    User->>SC: POST /stores (StoreData, logo, docs)
    SC->>CSA: execute(StoreData)
    CSA->>DB: Insert new Store record
    CSA-->>SC: Store Instance
    SC->>Storage: Upload logo_url & banner_url
    SC->>DB: Update Store with logo paths
    SC->>Storage: Upload Verification Documents
    SC->>DB: Insert Store Verifications
    SC-->>User: 201 Created""", "create_store_diagram.png", 6.0),
        ('Sequence Diagram: View Public Store Profile Flow', """sequenceDiagram
    participant Client
    participant SC as StoreController
    participant DB as Database
    Client->>SC: GET /public/stores/{store}
    SC->>DB: Fetch Store with public relations
    alt Store is Active
        SC->>DB: Insert into store_profile_views (IP, store_id, user_id)
        SC-->>Client: 200 OK (Public Store Data)
    else Store is Inactive/Deleted
        SC-->>Client: 404 Not Found
    end""", "view_store_diagram.png", 5.0),
        ('Sequence Diagram: Search & Filter Public Stores Flow', """sequenceDiagram
    participant Client
    participant SC as StoreController
    participant DB as Database
    Client->>SC: GET /public/stores?category_id=x&city=y&sort_by=popular
    SC->>DB: Base Query: Stores where status=ACTIVE
    SC->>DB: Apply Joins/WhereHas for categories & addresses
    SC->>DB: Apply ordering logic (followers_count, rating_avg)
    DB-->>SC: Paginated Collection
    SC-->>Client: 200 OK (Paginated Public Stores)""", "search_store_diagram.png", 6.0),
        ('Sequence Diagram: Admin Store Approval Flow', """sequenceDiagram
    participant Admin
    participant ASC as AdminStoreController
    participant ASA as ApproveStore Action
    participant DB as Database
    Admin->>ASC: POST /admin/stores/{store}/approve
    ASC->>ASA: execute(Store, notes)
    ASA->>DB: Validate status == PENDING
    ASA->>DB: Update Store status to ACTIVE
    ASA-->>ASC: Approved Store Instance
    ASC-->>Admin: 200 OK""", "admin_store_approve_diagram.png", 5.0),
        ('Sequence Diagram: Admin Verification Review Flow', """sequenceDiagram
    participant Admin
    participant ASC as AdminStoreController
    participant AVD as ApproveVerification Action
    participant DB as Database
    Admin->>ASC: POST /admin/stores/{store}/verifications/{id}/approve
    ASC->>AVD: execute(Verification, notes)
    AVD->>DB: Update Verification status to APPROVED
    AVD-->>ASC: Approved Verification Instance
    ASC-->>Admin: 200 OK""", "admin_verification_diagram.png", 5.0)
    ]
    for title, text, path, width in store_flows:
        doc.add_heading(title, level=5)
        download_image(generate_kroki_url(text), path)
        doc.add_picture(path, width=Inches(width))
        diagrams.append(path)

    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph("Store and Profile Management utilizes specific tables to handle multi-tenant business data alongside user demographics and engagement tracking.")
    for schema in [
        "profiles: Stores extended user details. Fields: id, user_id (FK to users), first_name, last_name, avatar_url, bio.",
        "stores: Main business entity. Fields: id (UUID), owner_user_id (FK to users), name, description, logo_url, banner_url, status, is_verified, rating_avg, followers_count.",
        "store_verifications: Tracks uploaded legal documents. Fields: id, store_id, document_type, document_path, status.",
        "store_profile_views: Analytics table tracking public views. Fields: id, store_id (FK to stores), user_id (nullable FK for logged-in users), ip_address."
    ]: doc.add_paragraph(schema, style='List Bullet')

    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    store_er_path = "store_er_diagram.png"
    download_image(generate_kroki_url("""erDiagram
    USERS ||--o| PROFILES : "has_one"
    USERS ||--o{ STORES : "owns"
    STORES ||--o{ STORE_VERIFICATIONS : "validates_via"
    STORES ||--o{ STORE_PROFILE_VIEWS : "receives"
    USERS ||--o{ STORE_PROFILE_VIEWS : "generates (optional)"
    USERS { char id PK "UUID" string email }
    PROFILES { bigint id PK uuid user_id FK string avatar_url }
    STORES { uuid id PK char owner_user_id FK string name string status }
    STORE_VERIFICATIONS { bigint id PK uuid store_id FK string status }
    STORE_PROFILE_VIEWS { bigint id PK uuid store_id FK }
"""), store_er_path)
    doc.add_picture(store_er_path, width=Inches(6.0))
    diagrams.append(store_er_path)

    filename = "Features_Documentation_Diagrams.docx"
    doc.save(filename)
    
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Full document generated successfully: {filename}")

if __name__ == "__main__":
    create_full_document()
