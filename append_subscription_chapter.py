import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_subscription_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== SUBSCRIPTIONS SYSTEM FEATURE ==================
    doc.add_heading('1.1.9 Subscriptions System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Subscriptions System governs the billing tier for Stores. It links Stores to predefined Subscription Plans using "
        "a robust, event-driven state machine. Subscriptions transition between various states (ACTIVE, SUSPENDED, ARCHIVED) using "
        "dedicated Action classes (like TransitionSubscriptionAction and ConfirmPaymentAction) to guarantee that every status change is "
        "recorded in historical audit logs. Payment Sessions act as a gateway; a subscription is only activated or renewed once a "
        "payment session is marked as PAID. Administrators can also manually bypass the payment gateway by issuing a 'Mock $0 Override Session' to directly assign a plan."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-51: View Subscriptions - Admin views a paginated list of all store subscriptions, optionally filtering by status, store, or plan.",
        "UC-52: View Subscription Details - Admin fetches a specific subscription, eager-loading its historical audit logs and plan details.",
        "UC-53: Manually Assign Subscription - Admin assigns a plan to a Store without requiring actual payment. The system generates an immediate $0 PAID payment session.",
        "UC-54: Cancel Subscription - Admin transitions an active subscription to 'ARCHIVED', requiring a cancellation reason.",
        "UC-55: Suspend Subscription - Admin transitions an active subscription to 'SUSPENDED'.",
        "UC-56: Approve Payment Session - Admin manually approves a pending gateway Payment Session, triggering the 'ConfirmPaymentAction' to safely activate the subscription."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Manually Assign Subscription Flow
    doc.add_heading('Sequence Diagram: Manually Assign Subscription Flow', level=5)
    assign_mermaid = """sequenceDiagram
    participant Admin
    participant SSMC as StoreSubscriptionManagementController
    participant CPA as ConfirmPaymentAction
    participant DB as Database
    
    Admin->>SSMC: POST /stores/{store}/subscriptions/assign (plan_id)
    SSMC->>DB: Begin Transaction
    SSMC->>DB: Create Mock PaymentSession (Amount=0, Status=PAID)
    SSMC->>CPA: execute(Store, Session ID)
    CPA->>DB: Process Subscription Update/Creation
    CPA->>DB: Write to subscription_history & audit_logs
    DB-->>SSMC: Commit Transaction
    SSMC-->>Admin: 200 OK (New Subscription Data)
"""
    assign_img_path = "assign_sub_diagram.png"
    download_image(generate_kroki_url(assign_mermaid), assign_img_path)
    doc.add_picture(assign_img_path, width=Inches(5.5))
    diagrams.append(assign_img_path)

    # 2. Transition Subscription Flow
    doc.add_heading('Sequence Diagram: Transition Subscription Flow (Cancel/Suspend)', level=5)
    transition_mermaid = """sequenceDiagram
    participant Admin
    participant SSMC as StoreSubscriptionManagementController
    participant TSA as TransitionSubscriptionAction
    participant DB as Database
    
    Admin->>SSMC: POST /subscriptions/{id}/cancel (reason)
    SSMC->>DB: FindOrFail(Subscription)
    SSMC->>TSA: execute(Subscription, ARCHIVED, reason)
    TSA->>DB: Update Subscription Status = ARCHIVED
    TSA->>DB: Insert into subscription_audit_logs (Old Status, New Status, Reason)
    TSA->>DB: Insert into subscription_history (Snapshot)
    TSA-->>SSMC: Updated Subscription
    SSMC-->>Admin: 200 OK
"""
    transition_img_path = "transition_sub_diagram.png"
    download_image(generate_kroki_url(transition_mermaid), transition_img_path)
    doc.add_picture(transition_img_path, width=Inches(5.5))
    diagrams.append(transition_img_path)

    # 3. Approve Payment Session Flow
    doc.add_heading('Sequence Diagram: Approve Payment Session Flow', level=5)
    approve_payment_mermaid = """sequenceDiagram
    participant Admin
    participant PSMC as PaymentSessionManagementController
    participant CPA as ConfirmPaymentAction
    participant DB as Database
    
    Admin->>PSMC: POST /payment-sessions/{id}/approve
    PSMC->>DB: Find Pending PaymentSession
    PSMC->>DB: Begin Transaction
    PSMC->>CPA: execute(Store, Session ID)
    CPA->>DB: Update Session Status = PAID
    CPA->>DB: Create/Renew Subscription based on Session data
    DB-->>PSMC: Commit Transaction
    PSMC-->>Admin: 200 OK (Subscription Activated)
"""
    approve_payment_img_path = "approve_payment_diagram.png"
    download_image(generate_kroki_url(approve_payment_mermaid), approve_payment_img_path)
    doc.add_picture(approve_payment_img_path, width=Inches(5.5))
    diagrams.append(approve_payment_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The schema isolates available billing tiers (plans) from active store states (subscriptions), while utilizing independent tables for deep auditing."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "subscription_plans: Defines the available billing tiers.",
        "subscriptions: Links a store_id to a plan_id. Fields include status, started_at, ends_at, auto_renew.",
        "payment_sessions: A temporary staging table mapping Paymob gateways to plans before they convert into Subscriptions.",
        "subscription_history: Periodical snapshots of a store's subscription data for historical reporting.",
        "subscription_audit_logs: An immutable ledger recording state changes (e.g., ACTIVE -> SUSPENDED) along with reasons and timestamps."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    SUBSCRIPTION_PLANS ||--o{ SUBSCRIPTIONS : "has"
    SUBSCRIPTION_PLANS ||--o{ PAYMENT_SESSIONS : "priced_by"
    STORES ||--|| SUBSCRIPTIONS : "owns"
    STORES ||--o{ PAYMENT_SESSIONS : "initiates"
    SUBSCRIPTIONS ||--o{ SUBSCRIPTION_HISTORY : "snapshotted_into"
    SUBSCRIPTIONS ||--o{ SUBSCRIPTION_AUDIT_LOGS : "logs_state_changes"
    
    SUBSCRIPTION_PLANS {
        bigint id PK
        string name
        decimal price
    }
    SUBSCRIPTIONS {
        bigint id PK
        uuid store_id FK
        bigint plan_id FK
        string status
        timestamp ends_at
    }
    PAYMENT_SESSIONS {
        bigint id PK
        uuid store_id FK
        bigint plan_id FK
        string status "PENDING/PAID"
    }
    SUBSCRIPTION_AUDIT_LOGS {
        bigint id PK
        bigint subscription_id FK
        string old_status
        string new_status
        string reason
    }
"""
    er_img_path = "subs_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Subscriptions System chapter appended to {filename}")

if __name__ == "__main__":
    append_subscription_chapter()
