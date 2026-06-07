import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_address_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # Remove old inaccurate chapter
    found = False
    for p in list(doc.paragraphs):
        if p.text == '1.1.12 Address Management System':
            found = True
        if found:
            p._element.getparent().remove(p._element)
            
    diagrams = []

    # ================== ADDRESS SYSTEM FEATURE ==================
    doc.add_heading('1.1.12 Address Management System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Address Management System utilizes a reusable polymorphic architecture. "
        "The core 'addresses' table holds geographical and textual data (such as address lines, city, postal code, and coordinates) and is detached from any specific owner. "
        "A polymorphic pivot table named 'addressables' maps these addresses to ANY entity (User, Store) using an 'owner_type' and 'owner_id'. "
        "This pivot table also stores contextual flags such as 'label' (e.g., home, work), 'is_default_shipping', and 'is_default_billing'. "
        "Dedicated domain action classes and private controller helpers ensure transactional safety when creating or toggling default statuses, preventing multiple defaults for the same owner."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-92: List User Addresses - Authenticated user retrieves a list of all their saved personal addresses, including labels and default flags.",
        "UC-93: Create User Address - User creates a new personal address. The system creates the address record and safely links it via the polymorphic pivot, syncing default flags if requested.",
        "UC-94: Update User Address - User updates an existing personal address. The system updates the base address and pivot fields, syncing default shipping/billing flags.",
        "UC-95: Delete User Address - User removes an address. The system detaches the pivot, and if the address is orphaned (no other owners), it deletes the address record.",
        "UC-96: List Store Addresses - Store Owner retrieves all operational or billing addresses associated with their store.",
        "UC-97: Create Store Address - Store Owner adds a physical storefront or headquarters address to their store via CreateStoreAddress action.",
        "UC-98: Update Store Address - Store Owner modifies a store address.",
        "UC-99: Delete Store Address - Store Owner deletes a store address, detaching the pivot and cleaning up orphans.",
        "UC-100: Admin Manage Store Addresses - Administrators have global override capabilities to list, create, update, or delete physical addresses for any store on the platform bypassing standard gates."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-93: Create User Address ---
    add_diagram(doc, "Create User Address (UC-93)", """sequenceDiagram
    participant User
    participant MAC as MeAddressController
    participant DB as Database

    User->>MAC: POST /me/addresses (address fields, label, is_default_shipping)
    MAC->>MAC: Validate payload
    MAC->>DB: Begin Transaction
    MAC->>DB: Insert into addresses table
    MAC->>DB: Attach to User via addressables (owner_type, owner_id)
    alt is_default_shipping == true
        MAC->>DB: syncDefaultFlags(User) -> Set all other addresses to false
    end
    DB-->>MAC: Commit Transaction
    MAC-->>User: 201 Created (AddressResource)
""", "sd_create_user_address", diagrams, 5.5)

    # --- UC-97: Create Store Address ---
    add_diagram(doc, "Create Store Address (UC-97)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as StoreAddressController
    participant CSA as CreateStoreAddress Action
    participant DB as Database

    Owner->>SAC: POST /stores/{store}/addresses
    SAC->>SAC: Gate Check (Can Manage Store)
    SAC->>CSA: execute(Store, AddressData)
    CSA->>DB: Begin Transaction
    CSA->>DB: Insert into addresses table
    CSA->>DB: Attach to Store via addressables
    alt is_default_billing == true
        CSA->>DB: syncDefaultFlags(Store) -> Set all other addresses to false
    end
    DB-->>CSA: Commit Transaction
    CSA-->>SAC: Address Model
    SAC-->>Owner: 201 Created
""", "sd_create_store_address", diagrams, 5.5)

    # --- UC-99: Delete Store Address ---
    add_diagram(doc, "Delete Store Address & Cleanup Orphans (UC-99)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as StoreAddressController
    participant DSA as DeleteStoreAddress Action
    participant DB as Database

    Owner->>SAC: DELETE /stores/{store}/addresses/{address}
    SAC->>DSA: execute(Store, Address)
    DSA->>DB: Detach from addressables for this Store
    DSA->>DB: Check if Address has other owners in addressables
    alt No other owners (Orphaned)
        DSA->>DB: Delete from addresses table
    end
    DSA-->>SAC: Success
    SAC-->>Owner: 200 OK
""", "sd_delete_store_address", diagrams, 5.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The schema relies on a polymorphic pivot to prevent duplicating address structures for different entities. "
    )
    schemas = [
        "addresses: Core location data. Fields: id (UUID), first_name, last_name, company, address_line1, address_line2, city, state_province, postal_code, country_code, phone_number, latitude, longitude, delivery_instructions.",
        "addressables: Polymorphic pivot mapping addresses to owners. Fields: id, address_id (FK), owner_type (Morph String), owner_id (Morph ID), label, is_default_shipping, is_default_billing."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    ADDRESSES ||--o{ ADDRESSABLES : "mapped_via"
    USERS ||--o{ ADDRESSABLES : "morphs_to (owner)"
    STORES ||--o{ ADDRESSABLES : "morphs_to (owner)"

    ADDRESSES {
        bigint id PK
        string address_line1
        string city
        string state_province
        string postal_code
        string country_code
        decimal latitude
        decimal longitude
    }
    ADDRESSABLES {
        bigint id PK
        bigint address_id FK
        char owner_id "UUID"
        string owner_type "Class Name"
        string label
        bool is_default_shipping
        bool is_default_billing
    }
"""
    er_img_path = "address_er_diagram_fixed.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Address System chapter fixed and appended to {filename}")

if __name__ == "__main__":
    fix_address_chapter()
