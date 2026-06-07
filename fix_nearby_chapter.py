import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_nearby_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # Remove old inaccurate chapter
    found = False
    for p in list(doc.paragraphs):
        if p.text == '1.1.18 Nearby Offers System':
            found = True
        if found:
            p._element.getparent().remove(p._element)
            
    diagrams = []

    # ================== NEARBY OFFERS SYSTEM FEATURE ==================
    doc.add_heading('1.1.18 Geolocation & Nearby Offers Logic', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Geolocation & Nearby Offers capability is not an isolated endpoint; rather, it is a cross-cutting 'Shared Concern' woven deeply into the application's most critical discovery systems. "
        "The core mathematical logic resides in a centralized 'HaversineCalculator' helper class, which generates complex raw SQL. Notably, it incorporates LEAST(1, GREATEST(-1,...)) clamping functions "
        "around the ACOS calculations to prevent floating-point precision errors (which cause NULL outputs in MySQL) when comparing incredibly close distances. "
        "This shared geospatial engine powers the 'Nearby' section of the Explore Dashboard, the proximity-based ranking in the Native Search System, and the local-fallback algorithm in the Following Feed. "
        "It accomplishes this by joining five tables: products, product_offers, stores, addressables (polymorphic), and finally addresses."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-126: Explore Nearby Section - During the Explore Bootstrap load, if user coordinates are provided, the system executes the Haversine formula to return a distinct payload of geographically sorted 'Nearby' offers alongside trending and flash items.",
        "UC-127: Search by Proximity - When executing a Native Search, the user can apply a `quick_filter=nearby`. The SearchOfferService utilizes the shared geolocation logic to rank matching search results strictly by closest physical distance.",
        "UC-128: Feed Popular Nearby Fallback - When a user views their Following Feed, if they run out of content from stores they actually follow, the algorithm automatically switches to a 'Popular Nearby' tier, using geolocation to recommend highly-rated local offers."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-126: Explore Nearby Section ---
    add_diagram(doc, "Explore Nearby Section (UC-126)", """sequenceDiagram
    participant User
    participant EC as ExploreController
    participant ES as ExploreService
    participant HC as HaversineCalculator
    participant DB as Database

    User->>EC: GET /explore/bootstrap?lat=X&lng=Y
    EC->>ES: Gather Bootstrap Data
    ES->>ES: Check if Lat/Lng present
    
    opt If Coordinates Provided
        ES->>HC: generateHaversineSql(lat, lng)
        HC-->>ES: Raw SQL String with LEAST/GREATEST Clamping
        ES->>DB: Build Query (products + offers + stores + addresses)
        ES->>DB: Inject Haversine SQL & ORDER BY distance ASC
        DB-->>ES: Nearby Collection
    end
    
    ES-->>EC: Unified Payload (Including Nearby Data)
    EC-->>User: 200 OK (JSON)
""", "sd_explore_nearby", diagrams, 6.0)

    # --- UC-127: Search by Proximity ---
    add_diagram(doc, "Search by Proximity (UC-127)", """sequenceDiagram
    participant User
    participant SOC as SearchOffersController
    participant SOS as SearchOfferService
    participant DB as Database

    User->>SOC: GET /search/offers?q=shoes&lat=X&lng=Y&quick_filter=nearby
    SOC->>SOS: executeSearch(filters)
    
    SOS->>DB: Match keyword 'shoes' against titles/categories
    SOS->>SOS: Detect quick_filter = 'nearby'
    SOS->>DB: Join `addressables` & `addresses`
    SOS->>DB: Inject Shared Haversine SQL
    SOS->>DB: Override Default Sorting -> ORDER BY distance ASC
    
    DB-->>SOS: Proximity Sorted Results
    SOS-->>SOC: Paginated Search Results
    SOC-->>User: 200 OK
""", "sd_search_nearby", diagrams, 6.0)

    # --- UC-128: Feed Popular Nearby Fallback ---
    add_diagram(doc, "Feed Popular Nearby Fallback (UC-128)", """sequenceDiagram
    participant User
    participant FFC as FollowingFeedController
    participant FFS as FollowingFeedService
    participant DB as Database

    User->>FFC: GET /feed?page=1&lat=X&lng=Y
    FFC->>FFS: getFeedOffers(User, lat, lng)
    
    FFS->>DB: Attempt to fetch 'Followed' Store Offers
    DB-->>FFS: Results (Empty or Insufficient)
    
    FFS->>FFS: Trigger Tier 2 Fallback (Popular Nearby)
    FFS->>DB: Query Active Products
    FFS->>DB: Inject Shared Haversine SQL
    FFS->>DB: Filter where Distance <= Radius
    FFS->>DB: ORDER BY Store Rating DESC, Distance ASC
    DB-->>FFS: Local Recommendations
    
    FFS->>FFS: Assign `recommendation_reason = 'popular_nearby'`
    FFS-->>FFC: Hydrated Feed
    FFC-->>User: 200 OK
""", "sd_feed_nearby_fallback", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "Geospatial calculation requires traversing the polymorphic relationship from Stores down to Addresses."
    )
    schemas = [
        "products & product_offers: Core promotional entity.",
        "stores: Joined via `product.store_id`.",
        "addressables: Polymorphic table joined where `owner_id = stores.id` AND `owner_type = 'App\\\\Domain\\\\Store\\\\Models\\\\Store'`.",
        "addresses: Final table providing the non-nullable `latitude` and `longitude` fields to the raw SQL Haversine query."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    SHARED_GEOLOCATION_ENGINE ||--o{ EXPLORE_SYSTEM : "powers nearby section"
    SHARED_GEOLOCATION_ENGINE ||--o{ SEARCH_SYSTEM : "powers proximity sorting"
    SHARED_GEOLOCATION_ENGINE ||--o{ FEED_SYSTEM : "powers local fallbacks"
    
    SHARED_GEOLOCATION_ENGINE ||--|| HAVERSINE_CALCULATOR : "utilizes"
    
    HAVERSINE_CALCULATOR ||--o{ ADDRESSES : "queries coordinates"
    ADDRESSES ||--|| ADDRESSABLES : "mapped via"
    ADDRESSABLES ||--|| STORES : "belongs to"
    STORES ||--o{ PRODUCTS : "owns"

    HAVERSINE_CALCULATOR {
        string raw_sql "LEAST/GREATEST Clamped ACOS"
    }
    ADDRESSABLES {
        bigint id PK
        char owner_id FK "UUID"
        string owner_type
        bigint address_id FK
    }
    ADDRESSES {
        bigint id PK
        decimal latitude
        decimal longitude
    }
"""
    er_img_path = "nearby_er_diagram_fixed.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Nearby Logic chapter fixed and appended to {filename}")

if __name__ == "__main__":
    fix_nearby_chapter()
