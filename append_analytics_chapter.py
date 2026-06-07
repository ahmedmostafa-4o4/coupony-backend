import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_analytics_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== ANALYTICS SYSTEM FEATURE ==================
    doc.add_heading('1.1.17 Analytics & Reporting System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Analytics & Reporting System is a powerful tool designed exclusively for Store Owners to track campaign performance, gauge user interest, and optimize conversion funnels. "
        "Unlike basic denormalized counters (which only show a total lifetime number), this system relies on dedicated interaction logging tables "
        "(such as product_views and offer_claims) that capture granular timestamps, user IDs, and IP addresses. "
        "This allows the backend AnalyticsService to perform complex GROUP BY date aggregations, enabling dynamic time-series charts on the frontend. "
        "Store owners can view a high-level dashboard encompassing all their active offers, drill down into funnel metrics for a specific offer, "
        "and export their raw data directly to CSV or PDF for offline reporting."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-122: View Store Analytics Dashboard - Store Owner requests aggregated metrics (Overview Cards, Time-Series Chart Data, Top Offers) spanning a specific date range (e.g., 7d, 30d).",
        "UC-123: View Specific Offer Performance - Store Owner requests granular data for a single offer, viewing deep funnel metrics (Unique Viewers -> Claims -> Redemptions) and a geographic breakdown by branch.",
        "UC-124: Export Analytics Report - Store Owner triggers an export action. The backend formats the data and returns a downloadable CSV or PDF file."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-122: View Store Analytics Dashboard ---
    add_diagram(doc, "View Store Analytics Dashboard (UC-122)", """sequenceDiagram
    participant Owner as Store Owner
    participant AC as AnalyticsController
    participant GDA as GetStoreDashboardMetricsAction
    participant DB as Database

    Owner->>AC: GET /analytics/stores/{store}/dashboard?date_range=30d
    AC->>AC: Gate Check (manageAnalytics)
    AC->>GDA: execute(Store, "30d")
    
    par Metric Aggregation
        GDA->>DB: COUNT(product_views), SUM(likes) GROUP BY DATE
        DB-->>GDA: Time-Series Data
        GDA->>DB: Fetch Active Offers Count
        DB-->>GDA: Scalar Count
        GDA->>DB: Query Top 5 Offers by claim conversion rate
        DB-->>GDA: Collection of Offers
    end
    
    GDA-->>AC: Formatted Array [overview, chart_data, top_offers]
    AC-->>Owner: 200 OK (JSON)
""", "sd_analytics_dashboard", diagrams, 6.0)

    # --- UC-123: View Specific Offer Performance ---
    add_diagram(doc, "View Specific Offer Performance (UC-123)", """sequenceDiagram
    participant Owner as Store Owner
    participant AC as AnalyticsController
    participant GOP as GetOfferPerformanceAction
    participant DB as Database

    Owner->>AC: GET /analytics/stores/{store}/offers/{offerId}
    AC->>AC: Gate Check (manageAnalytics)
    AC->>GOP: execute(Store, Offer)
    
    GOP->>DB: Fetch daily product_views for Offer
    GOP->>DB: Calculate Unique Viewers (Distinct IPs/UserIDs)
    GOP->>DB: Fetch offer_claims (Status: redeemed vs expired)
    GOP->>DB: Calculate Claim-to-Redemption Funnel Ratio
    GOP->>DB: Group claims by branch location
    
    DB-->>GOP: Raw Data Sets
    GOP->>GOP: Format Funnel & Geographic Data
    GOP-->>AC: Structured Array
    AC-->>Owner: 200 OK (JSON)
""", "sd_analytics_offer_performance", diagrams, 6.0)

    # --- UC-124: Export Analytics Report ---
    add_diagram(doc, "Export Analytics Report (UC-124)", """sequenceDiagram
    participant Owner as Store Owner
    participant AC as AnalyticsController
    participant EAR as ExportAnalyticsReportAction
    participant FS as File System

    Owner->>AC: GET /analytics/stores/{store}/export?format=csv
    AC->>AC: Gate Check (manageAnalytics)
    AC->>EAR: execute(Store, "csv")
    
    EAR->>Database: Fetch Raw Aggregated Data
    Database-->>EAR: Data Collection
    
    EAR->>EAR: Format Data into Tabular Format
    alt Format == CSV
        EAR->>FS: Generate Temporary .csv File
    else Format == PDF
        EAR->>FS: Generate Temporary .pdf File (via PDF library)
    end
    
    FS-->>EAR: File Stream / Path
    EAR-->>AC: Binary File Response
    AC-->>Owner: 200 OK (File Download)
""", "sd_analytics_export", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The analytics engine relies heavily on dedicated logging tables rather than simple integers on the main models. This ensures chronological tracking."
    )
    schemas = [
        "product_views: Granular tracking of offer views. Fields: id, product_id, user_id (nullable for guests), ip_address, viewed_at.",
        "store_profile_views: Granular tracking of store visits. Fields: id, store_id, user_id, ip_address, viewed_at.",
        "offer_claims: Used natively as a conversion metric. Tracks the status (redeemed vs expired) and the exact timestamps of claiming and redemption."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    PRODUCTS ||--o{ PRODUCT_VIEWS : "logs"
    STORES ||--o{ STORE_PROFILE_VIEWS : "logs"
    PRODUCT_OFFERS ||--o{ OFFER_CLAIMS : "tracks conversion"
    
    USERS ||--o{ PRODUCT_VIEWS : "generates"
    USERS ||--o{ STORE_PROFILE_VIEWS : "generates"
    USERS ||--o{ OFFER_CLAIMS : "claims"

    PRODUCT_VIEWS {
        bigint id PK
        uuid product_id FK
        char user_id FK "Nullable"
        string ip_address
        timestamp viewed_at
    }
    STORE_PROFILE_VIEWS {
        bigint id PK
        uuid store_id FK
        char user_id FK "Nullable"
        string ip_address
        timestamp viewed_at
    }
    OFFER_CLAIMS {
        uuid id PK
        uuid offer_id FK
        char user_id FK
        string status
        timestamp claimed_at
        timestamp redeemed_at
    }
"""
    er_img_path = "analytics_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Analytics System chapter appended to {filename}")

if __name__ == "__main__":
    append_analytics_chapter()
