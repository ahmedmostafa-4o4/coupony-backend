import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_payment_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== PAYMENT SYSTEM FEATURE ==================
    doc.add_heading('1.1.10 Payment System (Paymob Integration)', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Payment System manages the checkout process for Store Subscriptions using the external Paymob Gateway. "
        "It employs a robust dual-verification architecture: Users can manually trigger a confirmation once they complete the frontend checkout, "
        "while a server-to-server Webhook runs in the background to catch automated callbacks. "
        "The core data model is the 'PaymentSession', which tracks the lifecycle of a checkout attempt (PENDING -> PAID/FAILED) and prevents duplicate processing."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-57: Initiate Checkout - A Store Owner selects a subscription plan. The backend generates a Paymob order, acquires a client_secret, creates a PENDING PaymentSession, and returns the data for the frontend iframe.",
        "UC-58: Manual Payment Confirmation - After checking out, the Store Owner clicks confirm. The system verifies the transaction status directly with Paymob, marks the session as PAID, and instantly activates the subscription.",
        "UC-59: Automated Webhook Processing - Paymob fires an asynchronous background callback. The Webhook controller securely verifies the HMAC signature, locates the PaymentSession, and processes the subscription activation as a failsafe."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Initiate Checkout Flow (UC-57)
    doc.add_heading('Sequence Diagram: Initiate Checkout Flow (UC-57)', level=5)
    initiate_mermaid = """sequenceDiagram
    participant StoreOwner
    participant SC as SubscriptionController
    participant IPA as InitiatePaymentAction
    participant PM as Paymob API
    participant DB as Database
    
    StoreOwner->>SC: POST /initiate-payment (plan_id, cycle)
    SC->>IPA: execute(Store, Plan)
    IPA->>PM: Request Auth Token & Order ID
    PM-->>IPA: Order Details
    IPA->>PM: Request Payment Key (client_secret)
    PM-->>IPA: Payment Key Token
    IPA->>DB: Insert PaymentSession (Status=PENDING, paymob_order_id)
    IPA-->>SC: PaymentSession + client_secret
    SC-->>StoreOwner: 200 OK (Iframe URL / client_secret)
"""
    initiate_img_path = "initiate_payment_diagram.png"
    download_image(generate_kroki_url(initiate_mermaid), initiate_img_path)
    doc.add_picture(initiate_img_path, width=Inches(5.5))
    diagrams.append(initiate_img_path)

    # 2. Manual Payment Confirmation Flow (UC-58)
    doc.add_heading('Sequence Diagram: Manual Payment Confirmation Flow (UC-58)', level=5)
    confirm_mermaid = """sequenceDiagram
    participant StoreOwner
    participant SC as SubscriptionController
    participant CPA as ConfirmPaymentAction
    participant DB as Database
    
    StoreOwner->>SC: POST /confirm-payment (session_id)
    SC->>CPA: execute(Store, session_id)
    CPA->>DB: Lock PaymentSession For Update
    alt Session is already PAID
        CPA-->>SC: Throw PaymentSessionAlreadyUsedException
        SC-->>StoreOwner: 409 Conflict
    else Session is PENDING
        CPA->>DB: Update Session Status = PAID
        CPA->>DB: Activate/Renew Subscription
        DB-->>CPA: Commit Transaction
        CPA-->>SC: Updated Subscription
        SC-->>StoreOwner: 200 OK
    end
"""
    confirm_img_path = "confirm_payment_diagram.png"
    download_image(generate_kroki_url(confirm_mermaid), confirm_img_path)
    doc.add_picture(confirm_img_path, width=Inches(5.5))
    diagrams.append(confirm_img_path)

    # 3. Automated Webhook Processing Flow (UC-59)
    doc.add_heading('Sequence Diagram: Automated Webhook Processing Flow (UC-59)', level=5)
    webhook_mermaid = """sequenceDiagram
    participant Paymob
    participant WC as WebhookController
    participant PWA as ProcessWebhookAction
    participant CPA as ConfirmPaymentAction
    participant DB as Database
    
    Paymob->>WC: POST /webhooks/paymob (Payload, HMAC Header)
    WC->>PWA: execute(payload, hmac)
    PWA->>PWA: Verify HMAC Signature securely
    alt Invalid HMAC
        PWA-->>WC: Throw HttpException(401)
        WC-->>Paymob: 401 Unauthorized
    else Valid HMAC
        PWA->>DB: Find PaymentSession by Order ID
        PWA->>CPA: execute(Store, Session ID)
        CPA->>DB: Update Session to PAID
        CPA->>DB: Activate/Renew Subscription
        CPA-->>PWA: Success
        PWA-->>WC: Success
        WC-->>Paymob: 200 OK (Acknowledged)
    end
"""
    webhook_img_path = "webhook_payment_diagram.png"
    download_image(generate_kroki_url(webhook_mermaid), webhook_img_path)
    doc.add_picture(webhook_img_path, width=Inches(5.5))
    diagrams.append(webhook_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The checkout process relies entirely on the temporary 'payment_sessions' table, mapping our internal billing structures to external gateway references."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "payment_sessions: Fields: id (PK), store_id (FK), plan_id (FK), billing_cycle, amount (decimal), currency (string), status (enum: PENDING, PAID, FAILED), paymob_order_id (string), paymob_transaction_id (string), payment_url, expires_at (timestamp), paid_at (timestamp), failure_reason (text)."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    STORES ||--o{ PAYMENT_SESSIONS : "initiates"
    SUBSCRIPTION_PLANS ||--o{ PAYMENT_SESSIONS : "prices"
    PAYMENT_SESSIONS ||--o| SUBSCRIPTIONS : "converts_to"
    
    PAYMENT_SESSIONS {
        bigint id PK
        uuid store_id FK
        bigint plan_id FK
        string billing_cycle
        decimal amount
        string status "PENDING/PAID"
        string paymob_order_id
        string paymob_transaction_id
    }
"""
    er_img_path = "payment_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Payment System chapter appended to {filename}")

if __name__ == "__main__":
    append_payment_chapter()
