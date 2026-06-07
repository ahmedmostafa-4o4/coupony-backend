import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_ponyai_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== PONYAI SYSTEM ==================
    doc.add_heading('1.1.20 PonyAI Intelligent Assistant System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "PonyAI is the platform's built-in AI-powered conversational assistant, backed by Google Gemini. "
        "It provides context-aware chat sessions where users can discover offers, get store help, or ask general questions. "
        "The system's most advanced feature is autonomous Tool Calling: when the conversation context is 'offer_discovery', "
        "the PonyAIService injects a function definition for 'search_offers(query, category, max_price)' into the Gemini API request. "
        "If Gemini decides the user's intent requires a product search, it returns a structured tool_call instead of plain text. "
        "The service then intercepts this, executes the search against the real SearchOfferService, feeds the raw results back to Gemini, "
        "and the AI synthesizes them into a natural-language summary with recommendations. "
        "Conversation history is managed via a sliding window of the last N messages (configurable via pony_ai_settings) to stay within token limits. "
        "A ContextPromptBuilder tailors the system prompt per context type, and Admins can dynamically update prompts, model parameters (temperature, max_tokens), "
        "and review conversation quality through a dedicated management dashboard."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-134: Start AI Conversation - User initiates a new PonyAI chat session, selecting a context (offer_discovery, store_help, general). The system creates a conversation record and initializes it with a tailored system prompt.",
        "UC-135: Send Message & Receive AI Response - User sends a text message. The PonyAIService loads conversation history, builds the prompt array, calls Gemini, and persists both the user message and the AI response.",
        "UC-136: AI Autonomous Tool Call (Offer Discovery) - During an offer_discovery conversation, Gemini autonomously decides to invoke 'search_offers'. The service executes the real search, feeds results back to Gemini, and returns a natural-language product recommendation.",
        "UC-137: View Conversation History - User retrieves the full message thread of a specific past conversation.",
        "UC-138: List Past Conversations - User browses a paginated list of their previous PonyAI conversations with preview snippets.",
        "UC-139: Delete Conversation - User soft-deletes a conversation and all its associated messages.",
        "UC-140: Rate AI Response - User submits a thumbs up/down rating and optional feedback text on a specific AI message for quality tracking.",
        "UC-141: Admin View PonyAI Dashboard - Admin retrieves aggregate usage statistics: total conversations, total messages, average satisfaction rating, top contexts, and daily usage chart data.",
        "UC-142: Admin Review Conversations - Admin browses and inspects any user's conversation thread for quality assurance, with filtering by context, date, and rating.",
        "UC-143: Admin Manage AI Settings - Admin updates the system prompt template, Gemini model name, temperature, max_tokens, and max_history_messages via the pony_ai_settings table."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-134: Start AI Conversation ---
    add_diagram(doc, "Start AI Conversation (UC-134)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant PS as PonyAIService
    participant CPB as ContextPromptBuilder
    participant DB as Database

    User->>PC: POST /pony-ai/conversations (context: 'offer_discovery')
    PC->>PS: createConversation(User, 'offer_discovery')
    PS->>CPB: build('offer_discovery')
    CPB-->>PS: Tailored System Prompt
    PS->>DB: INSERT INTO pony_ai_conversations
    PS->>DB: INSERT system prompt as first pony_ai_messages (role: 'system')
    DB-->>PS: Conversation Record
    PS-->>PC: Conversation ID
    PC-->>User: 201 Created (conversation_id)
""", "sd_ponyai_start", diagrams, 6.0)

    # --- UC-135: Send Message & Receive AI Response ---
    add_diagram(doc, "Send Message and Receive AI Response (UC-135)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant PS as PonyAIService
    participant GC as GeminiClient
    participant DB as Database

    User->>PC: POST /pony-ai/conversations/{id}/messages (message)
    PC->>PS: processMessage(conversationId, message)
    PS->>DB: Load last N messages (sliding history window)
    PS->>PS: Build prompt array [system, ...history, user_msg]
    PS->>GC: chat(messages, tools, temperature, maxTokens)
    GC->>GC: POST generativelanguage.googleapis.com
    GC-->>PS: AI Text Response
    PS->>DB: INSERT user message (role: 'user')
    PS->>DB: INSERT AI response (role: 'assistant', tokens_used)
    PS-->>PC: AI Response
    PC-->>User: 200 OK (assistant message)
""", "sd_ponyai_send_message", diagrams, 6.0)

    # --- UC-136: AI Autonomous Tool Call ---
    add_diagram(doc, "AI Autonomous Tool Call - Offer Discovery (UC-136)", """sequenceDiagram
    participant User
    participant PS as PonyAIService
    participant GC as GeminiClient
    participant SOS as SearchOfferService
    participant DB as Database

    User->>PS: "Find me cheap shoes near downtown"
    PS->>GC: chat(messages, tools=[search_offers])
    GC-->>PS: tool_call: search_offers(query='shoes', max_price=50)
    
    Note over PS: Intercept Tool Call
    PS->>SOS: execute('shoes', max_price=50)
    SOS->>DB: Haversine + keyword search
    DB-->>SOS: Matching Offers
    SOS-->>PS: Search Results Array
    
    PS->>GC: chat([...history, tool_result: results])
    Note over GC: AI synthesizes natural language summary
    GC-->>PS: "I found 3 shoe deals near you..."
    
    PS->>DB: Save messages (tool_calls JSON, tool_results JSON)
    PS-->>User: Natural Language Recommendation
""", "sd_ponyai_tool_call", diagrams, 6.0)

    # --- UC-137: View Conversation History ---
    add_diagram(doc, "View Conversation History (UC-137)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant DB as Database

    User->>PC: GET /pony-ai/conversations/{id}
    PC->>DB: Fetch conversation WHERE user_id = User
    PC->>DB: Fetch all pony_ai_messages ORDER BY created_at
    DB-->>PC: Messages Collection
    PC-->>User: 200 OK (Full Thread)
""", "sd_ponyai_view_history", diagrams, 5.0)

    # --- UC-138: List Past Conversations ---
    add_diagram(doc, "List Past Conversations (UC-138)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant DB as Database

    User->>PC: GET /pony-ai/conversations?page=1
    PC->>DB: Query pony_ai_conversations WHERE user_id = User
    PC->>DB: ORDER BY last_message_at DESC, Paginate
    DB-->>PC: Paginated Conversations with Previews
    PC-->>User: 200 OK (Conversation List)
""", "sd_ponyai_list_conversations", diagrams, 5.0)

    # --- UC-139: Delete Conversation ---
    add_diagram(doc, "Delete Conversation (UC-139)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant DB as Database

    User->>PC: DELETE /pony-ai/conversations/{id}
    PC->>DB: Verify ownership (user_id = User)
    PC->>DB: Soft-delete conversation (SET deleted_at)
    PC->>DB: Soft-delete associated messages
    DB-->>PC: Done
    PC-->>User: 200 OK
""", "sd_ponyai_delete", diagrams, 5.0)

    # --- UC-140: Rate AI Response ---
    add_diagram(doc, "Rate AI Response (UC-140)", """sequenceDiagram
    participant User
    participant PC as PonyAIController
    participant DB as Database

    User->>PC: POST /pony-ai/messages/{id}/rate (rating: 1, feedback: 'Helpful!')
    PC->>DB: Verify message belongs to User's conversation
    PC->>DB: UPDATE pony_ai_messages SET rating=1, rating_feedback='Helpful!'
    DB-->>PC: Updated
    PC-->>User: 200 OK
""", "sd_ponyai_rate", diagrams, 5.0)

    # --- UC-141: Admin View Dashboard ---
    add_diagram(doc, "Admin View PonyAI Dashboard (UC-141)", """sequenceDiagram
    participant Admin
    participant AMC as PonyAIManagementController
    participant DB as Database

    Admin->>AMC: GET /admin/pony-ai/dashboard
    par Aggregate Metrics
        AMC->>DB: COUNT total conversations
        AMC->>DB: COUNT total messages
        AMC->>DB: AVG rating from pony_ai_messages
        AMC->>DB: GROUP BY context (top contexts)
        AMC->>DB: GROUP BY DATE (daily usage chart)
    end
    DB-->>AMC: Aggregated Data
    AMC-->>Admin: 200 OK (Dashboard JSON)
""", "sd_ponyai_admin_dashboard", diagrams, 5.5)

    # --- UC-142: Admin Review Conversations ---
    add_diagram(doc, "Admin Review Conversations (UC-142)", """sequenceDiagram
    participant Admin
    participant AMC as PonyAIManagementController
    participant DB as Database

    Admin->>AMC: GET /admin/pony-ai/conversations?context=offer_discovery
    AMC->>DB: Query ALL conversations with filters
    DB-->>AMC: Paginated List
    AMC-->>Admin: 200 OK

    Admin->>AMC: GET /admin/pony-ai/conversations/{id}
    AMC->>DB: Fetch full message thread
    DB-->>AMC: Messages with tool_calls and ratings
    AMC-->>Admin: 200 OK (Full Thread for QA)
""", "sd_ponyai_admin_review", diagrams, 5.5)

    # --- UC-143: Admin Manage AI Settings ---
    add_diagram(doc, "Admin Manage AI Settings (UC-143)", """sequenceDiagram
    participant Admin
    participant AMC as PonyAIManagementController
    participant DB as Database

    alt Get Current Settings
        Admin->>AMC: GET /admin/pony-ai/settings
        AMC->>DB: SELECT * FROM pony_ai_settings
        DB-->>AMC: Key-Value Pairs
        AMC-->>Admin: 200 OK (Settings JSON)
    else Update Settings
        Admin->>AMC: PUT /admin/pony-ai/settings (temperature: 0.7, model: 'gemini-2.0-flash')
        AMC->>DB: UPSERT pony_ai_settings for each key
        DB-->>AMC: Updated
        AMC-->>Admin: 200 OK
    end
""", "sd_ponyai_admin_settings", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "PonyAI maintains its own isolated schema to track conversations, messages (including tool interactions), and dynamic configuration."
    )
    schemas = [
        "pony_ai_conversations: Tracks chat sessions. Fields: id (uuid), user_id, context (offer_discovery/store_help/general), title (auto-generated), last_message_at, is_active, metadata (JSON), deleted_at.",
        "pony_ai_messages: Stores every message in the thread. Fields: id (uuid), conversation_id, role (user/assistant/system/tool), content (longText), attachments (JSON), tool_calls (JSON), tool_results (JSON), tokens_used, rating (nullable: 1/-1), rating_feedback.",
        "pony_ai_settings: Admin-configurable key-value store. Keys include: system_prompt, model_name, temperature, max_tokens, max_history_messages."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o{ PONY_AI_CONVERSATIONS : "starts"
    PONY_AI_CONVERSATIONS ||--o{ PONY_AI_MESSAGES : "contains"

    PONY_AI_CONVERSATIONS {
        uuid id PK
        char user_id FK
        string context
        string title
        timestamp last_message_at
        json metadata
    }
    PONY_AI_MESSAGES {
        uuid id PK
        uuid conversation_id FK
        enum role "user, assistant, system, tool"
        longtext content
        json tool_calls "Nullable"
        json tool_results "Nullable"
        int tokens_used
        tinyint rating "Nullable"
    }
    PONY_AI_SETTINGS {
        bigint id PK
        string key "Unique"
        text value
    }
"""
    er_img_path = "ponyai_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"PonyAI System chapter appended to {filename}")

if __name__ == "__main__":
    append_ponyai_chapter()
