import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_analytics_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # Remove old inaccurate chapter
    found = False
    for p in list(doc.paragraphs):
        if p.text == '1.1.17 Analytics & Reporting System':
            found = True
        if found:
            p._element.getparent().remove(p._element)
            
    diagrams = []

    # ================== ANALYTICS SYSTEM FEATURE ==================
    doc.add_heading('1.1.17 Analytics & Reporting System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Analytics System is designed to provide actionable intelligence at both the macro platform level and the micro merchant level. "
        "For System Administrators, it tracks platform-wide financial health, calculating MRR (Monthly Recurring Revenue), total active subscriptions, and churn rates. "
        "For Store Owners and Staff (secured by the ANALYTICS_VIEW permission), it provides an advanced dashboard powered by complex time-series aggregations. "
        "Sellers can set and track monthly redemption goals, analyze redemption heatmaps to identify peak traffic times, and drill down into specific products. "
        "The Product Analytics engine tracks granular 'Impressions' using dedicated logging tables, logging the exact 'Traffic Source' (e.g., search, explore, profile) "
        "and utilizing joins against User Profiles to generate detailed Audience Demographics (Age buckets, Gender distribution) to help sellers target their campaigns."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-122: Admin Platform Metrics - System Administrators view macro-level statistics regarding platform revenue (from PAID sessions), MRR, and churn count over the past 30 days.",
        "UC-123: View Store Dashboard - Store Owners monitor store health by evaluating performance against a Monthly Goal, tracking follower growth, viewing redemption heatmaps, and identifying top offers.",
        "UC-124: Update Monthly Goal - Store Owners update their target monthly redemption goal.",
        "UC-125: View Product Analytics - Store Owners drill into a specific product to understand reach (Traffic Sources), engagement trends, and deep Audience Demographics (Age, Gender)."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-122: Admin Platform Metrics ---
    add_diagram(doc, "Admin Platform Metrics (UC-122)", """sequenceDiagram
    participant Admin
    participant SAC as SubscriptionAnalyticsController
    participant DB as Database

    Admin->>SAC: GET /admin/analytics/statistics
    par Calculate Subscription Metrics
        SAC->>DB: COUNT active subscriptions
        SAC->>DB: SUM revenue from PAID sessions
        SAC->>DB: Calculate MRR based on active plans
        SAC->>DB: COUNT churned (ARCHIVED/SUSPENDED) past 30d
    end
    DB-->>SAC: Metrics Array
    SAC-->>Admin: 200 OK (Platform Statistics JSON)
""", "sd_admin_analytics", diagrams, 5.0)

    # --- UC-123 & UC-124: Store Dashboard & Goals ---
    add_diagram(doc, "Store Dashboard & Goals (UC-123, UC-124)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as SellerAnalyticsController
    participant GDA as GetSellerDashboardAction
    participant UMGA as UpdateMonthlyGoalAction
    participant DB as Database

    alt UC-123: View Dashboard
        Owner->>SAC: GET /seller/stores/{id}/dashboard
        SAC->>GDA: execute(Store, Period)
        GDA->>DB: Fetch claims vs Monthly Goal
        GDA->>DB: Calculate Follower/Visit Growth
        GDA->>DB: Aggregate Peak Redemption Heatmap
        GDA->>DB: Query Top 10 Offers
        GDA-->>SAC: Dashboard Array
        SAC-->>Owner: 200 OK
    else UC-124: Update Goal
        Owner->>SAC: PUT /seller/stores/{id}/monthly-goal
        SAC->>UMGA: execute(Store, new_goal)
        UMGA->>DB: UPDATE stores SET monthly_goal = X
        UMGA-->>SAC: Success
        SAC-->>Owner: 200 OK
    end
""", "sd_seller_dashboard", diagrams, 6.0)

    # --- UC-125: View Product Analytics ---
    add_diagram(doc, "View Product Analytics (UC-125)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as SellerAnalyticsController
    participant GPA as GetProductAnalyticsAction
    participant DB as Database

    Owner->>SAC: GET /seller/stores/{store}/products/{id}/analytics
    SAC->>GPA: execute(Product, Period)
    
    par Complex Analytics Aggregation
        GPA->>DB: Fetch Engagement (Views, Likes, Comments)
        GPA->>DB: Group product_views by `source` (Traffic Sources)
        GPA->>DB: Join users & profiles
        Note over GPA,DB: Demographics Calculation
        GPA->>DB: Calculate Age Buckets from profiles.date_of_birth
        GPA->>DB: Group by profiles.gender
    end
    
    DB-->>GPA: Raw Metric Data
    GPA->>GPA: Format Chart Series & Normalizations
    GPA-->>SAC: Formatted Payload
    SAC-->>Owner: 200 OK (Product Analytics JSON)
""", "sd_product_analytics", diagrams, 6.0)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "To support deep demographic and traffic analysis, the system introduces dedicated granular tracking tables enriched with metadata."
    )
    schemas = [
        "product_views: Core impression table. Tracks `product_id`, `user_id` (nullable), `ip_address`, `user_agent`, and importantly, the `source` (e.g., search, explore, profile).",
        "store_profile_views: Tracks overall store traffic. Fields: `store_id`, `user_id`, `ip_address`.",
        "stores.monthly_goal: A simple integer on the main stores table utilized by the dashboard to measure periodic campaign success."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    PRODUCTS ||--o{ PRODUCT_VIEWS : "tracks impressions"
    STORES ||--o{ STORE_PROFILE_VIEWS : "tracks visits"
    
    USERS ||--o{ PRODUCT_VIEWS : "generates"
    USERS ||--|| PROFILES : "provides demographics"
    
    PRODUCT_VIEWS {
        bigint id PK
        uuid product_id FK
        char user_id FK "Nullable"
        string source
        timestamp viewed_at
    }
    PROFILES {
        bigint id PK
        date date_of_birth
        string gender
    }
    STORES {
        uuid id PK
        int monthly_goal
    }
"""
    er_img_path = "analytics_er_diagram_fixed.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Analytics System chapter fixed and appended to {filename}")

if __name__ == "__main__":
    fix_analytics_chapter()
