import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_missing_address_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    doc.add_heading('Additional Address System Sequence Diagrams', level=4)

    # --- UC-92: List User Addresses ---
    add_diagram(doc, "List User Addresses (UC-92)", """sequenceDiagram
    participant User
    participant MAC as MeAddressController
    participant DB as Database

    User->>MAC: GET /me/addresses
    MAC->>DB: Query addressables where owner_id = User.id
    MAC->>DB: Apply search filters (city, label, etc.)
    MAC->>DB: Eager load Address model
    DB-->>MAC: Collection of Addresses with pivot data
    MAC-->>User: 200 OK (AddressResource Collection)
""", "sd_list_user_addresses", diagrams)

    # --- UC-94: Update User Address ---
    add_diagram(doc, "Update User Address (UC-94)", """sequenceDiagram
    participant User
    participant MAC as MeAddressController
    participant DB as Database

    User->>MAC: PUT /me/addresses/{address}
    MAC->>MAC: Validate payload
    MAC->>DB: Begin Transaction
    MAC->>DB: UPDATE addresses SET updated fields
    MAC->>DB: UPDATE addressables SET label, flags WHERE address_id
    alt is_default_shipping == true
        MAC->>DB: syncDefaultFlags(User) -> Update other pivots to false
    end
    DB-->>MAC: Commit Transaction
    MAC-->>User: 200 OK
""", "sd_update_user_address", diagrams)

    # --- UC-95: Delete User Address ---
    add_diagram(doc, "Delete User Address (UC-95)", """sequenceDiagram
    participant User
    participant MAC as MeAddressController
    participant DB as Database

    User->>MAC: DELETE /me/addresses/{address}
    MAC->>DB: Begin Transaction
    MAC->>DB: Detach User from addressables pivot
    MAC->>DB: Check if Address has other owners (addressables count)
    alt No other owners (Orphaned)
        MAC->>DB: Delete from addresses table
    end
    DB-->>MAC: Commit Transaction
    MAC-->>User: 200 OK
""", "sd_delete_user_address", diagrams)

    # --- UC-96: List Store Addresses ---
    add_diagram(doc, "List Store Addresses (UC-96)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as StoreAddressController
    participant LSA as ListStoreAddresses Action
    participant DB as Database

    Owner->>SAC: GET /stores/{store}/addresses
    SAC->>SAC: Gate Check (Can Manage Store)
    SAC->>LSA: execute(Store)
    LSA->>DB: Query Store addresses via addressables relation
    DB-->>LSA: Address Collection
    LSA-->>SAC: Address Collection
    SAC-->>Owner: 200 OK (StoreAddressResource Collection)
""", "sd_list_store_addresses", diagrams)

    # --- UC-98: Update Store Address ---
    add_diagram(doc, "Update Store Address (UC-98)", """sequenceDiagram
    participant Owner as Store Owner
    participant SAC as StoreAddressController
    participant USA as UpdateStoreAddress Action
    participant DB as Database

    Owner->>SAC: PUT /stores/{store}/addresses/{address}
    SAC->>SAC: Gate Check (Can Manage Store)
    SAC->>USA: execute(Store, Address, NewData)
    USA->>DB: Begin Transaction
    USA->>DB: UPDATE addresses fields
    USA->>DB: UPDATE addressables pivot for this Store
    alt is_default_billing == true
        USA->>DB: syncDefaultFlags(Store)
    end
    DB-->>USA: Commit Transaction
    USA-->>SAC: Updated Address
    SAC-->>Owner: 200 OK
""", "sd_update_store_address", diagrams)

    # --- UC-100: Admin Manage Store Addresses ---
    add_diagram(doc, "Admin Manage Store Addresses (UC-100)", """sequenceDiagram
    participant Admin
    participant SAMC as Admin\StoreAddressManagementController
    participant Actions as StoreAddress Actions
    participant DB as Database

    Admin->>SAMC: POST /admin/stores/{store}/addresses
    Note over SAMC: Bypasses Store Owner Gates (uses admin middleware)
    SAMC->>Actions: CreateStoreAddress::execute(Store, Data)
    Actions->>DB: Insert address & pivot
    Actions-->>SAMC: Success
    SAMC-->>Admin: 201 Created
""", "sd_admin_manage_addresses", diagrams)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Missing Address System diagrams appended to {filename}")

if __name__ == "__main__":
    append_missing_address_diagrams()
