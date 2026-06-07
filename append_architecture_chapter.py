import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Architecture Diagram: {title}', level=3)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_architecture_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    # ================== CHAPTER 5: CODE ARCHITECTURE ==================
    doc.add_page_break()
    doc.add_heading('Chapter 5: Code Architecture & Design Patterns', level=1)
    
    intro_text = (
        "To ensure long-term maintainability, testability, and scalability, this platform strictly abandons the default Laravel MVC structure. "
        "Instead, it adopts a highly decoupled Domain-Driven Design (DDD) architecture. Business logic is completely isolated from HTTP delivery mechanisms, "
        "allowing the application to scale gracefully as new features are added without creating 'fat controllers' or 'god models'."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 5.1 Domain-Driven Design ------------------
    doc.add_heading('5.1 Domain-Driven Design (DDD)', level=2)
    doc.add_paragraph("The `app/` directory is split into two primary boundaries:")
    ddd = [
        "app/Application/: Contains the HTTP delivery layer. This includes all Controllers, FormRequests, API Resources, and Middleware. Controllers here are 'thin' and only responsible for receiving requests and returning HTTP responses.",
        "app/Domain/: The absolute core of the platform. It contains pure business logic separated into 15 isolated domains (e.g., User, Store, Product, Subscription, PonyAI, Notification).",
        "Domain Internal Structure: Each domain contains its own localized Folders: `Models`, `Actions`, `Repositories`, `DTOs`, `Events`, `Listeners`, `Enums`, and `Exceptions`."
    ]
    for d in ddd:
        doc.add_paragraph(d, style='List Bullet')

    # ------------------ 5.2 The Action Pattern ------------------
    doc.add_heading('5.2 The Action Pattern', level=2)
    doc.add_paragraph("Instead of dumping business logic into Controllers or bloated Services, the platform relies heavily on the Action pattern. Actions are single-responsibility classes that execute exactly one business process.")
    actions = [
        "Single Responsibility: Classes like `ConfirmPaymentAction` or `GetSellerDashboardAction` contain exactly one public `execute()` method.",
        "Reusability: Because Actions are decoupled from the HTTP Request, they can be safely triggered by a Controller, a Console Command, or an asynchronous Queue Worker without modification.",
        "Dependency Injection: Actions utilize constructor injection to resolve their dependencies (like Repositories or other Actions), making them incredibly easy to mock during Unit Testing."
    ]
    for a in actions:
        doc.add_paragraph(a, style='List Bullet')

    # ------------------ 5.3 The Repository Pattern ------------------
    doc.add_heading('5.3 The Repository Pattern', level=2)
    doc.add_paragraph("Direct Eloquent queries (`Model::where(...)`) are strictly prohibited inside Controllers. All database interactions flow through dedicated Repositories.")
    repo = [
        "Query Abstraction: Classes like `UserRepository` or `StoreRepository` abstract complex Eloquent queries and joins away from the business logic.",
        "Centralized Caching: Repositories handle automatic caching (e.g., caching `user.by_id.{id}` for 1 hour). If the database schema changes or a query needs optimization, it only needs to be updated in one isolated Repository file."
    ]
    for r in repo:
        doc.add_paragraph(r, style='List Bullet')

    # ------------------ 5.4 Data Transfer Objects (DTOs) ------------------
    doc.add_heading('5.4 Data Transfer Objects (DTOs)', level=2)
    doc.add_paragraph("To ensure Type Safety, data passed from the Application layer down into the Domain Actions must be wrapped in DTOs.")
    dtos = [
        "Type Safety: Instead of passing unpredictable associative arrays (`$request->all()`), Controllers map validated request data into strict PHP classes (e.g., `CreateProductDTO`).",
        "Predictability: Actions type-hint the DTO, giving developers full IDE autocomplete and guaranteeing that the required data structures exist before execution begins."
    ]
    for d in dtos:
        doc.add_paragraph(d, style='List Bullet')

    # ======================== ARCHITECTURE DIAGRAMS ========================
    doc.add_heading('System Architecture Diagrams', level=2)

    add_diagram(doc, "Request Flow (DDD + Action Pattern)", """sequenceDiagram
    participant Client
    participant App as Application Layer (Controller)
    participant DTO as Data Transfer Object
    participant Domain as Domain Layer (Action)
    participant Repo as Repository Layer
    participant DB as Database

    Client->>App: POST /api/v1/subscriptions (JSON payload)
    App->>App: FormRequest Validation
    App->>DTO: Map validated array to CreateSubscriptionDTO
    App->>Domain: execute(CreateSubscriptionDTO)
    
    Domain->>Domain: Execute core business logic
    Domain->>Repo: store(SubscriptionEntity)
    
    Repo->>DB: Execute Eloquent INSERT
    DB-->>Repo: Saved Record
    Repo-->>Domain: Returned Model
    
    Domain-->>App: Result
    App->>App: Wrap in API Resource
    App-->>Client: 200 OK (Standard JSON Wrapper)
""", "sd_ddd_flow", diagrams, 6.0)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Code Architecture Chapter appended to {filename}")

if __name__ == "__main__":
    append_architecture_chapter()
