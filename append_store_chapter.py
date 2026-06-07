import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_store_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== STORE & PROFILE MANAGEMENT FEATURE ==================
    doc.add_heading('1.1.3 Store & Profile Management', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "Store & Profile Management allows users to manage their personal identities and create business entities (Stores). "
        "Users can update their personal profiles (like first name, last name, and avatar) via the '/me' endpoint. "
        "When a user creates a new Store, the system records the core store information and handles the upload of store assets (logo and banner) as well as verification documents. "
        "Store owners and managers can separately update public-facing store profiles and manage business hours, social links, and locations. "
        "To provide analytics to store owners, the system automatically logs a 'Store Profile View' record with the viewer's IP address every time a public store profile is visited."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-17: User Profile Update - A user updates their personal profile information (name, avatar, bio) via the authenticated /me endpoint.",
        "UC-18: Create Store - A user creates a new store, providing store details and securely uploading initial logos and verification documents.",
        "UC-19: Update Store Profile - A store manager explicitly updates the public-facing store profile assets (logo, banner) separate from core store business details.",
        "UC-20: View Public Store Profile - Customers view a public store profile. This action automatically logs a view record for store analytics.",
        "UC-21: Update Verification Documents - Store managers submit or replace business verification documents to achieve a 'Verified' status."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Update User Profile Flow
    doc.add_heading('Sequence Diagram: Update User Profile Flow', level=5)
    profile_mermaid = """sequenceDiagram
    participant User
    participant LC as LoginController
    participant DB as Database
    participant Storage
    
    User->>LC: PATCH /me (first_name, avatar)
    LC->>DB: Fetch authenticated User & Profile
    alt Avatar provided
        LC->>Storage: Delete old avatar if exists
        LC->>Storage: Upload new avatar
    end
    LC->>DB: Update user_profiles table
    LC-->>User: 200 OK (Updated User Profile)
"""
    profile_img_path = "profile_diagram.png"
    download_image(generate_kroki_url(profile_mermaid), profile_img_path)
    doc.add_picture(profile_img_path, width=Inches(5.0))
    diagrams.append(profile_img_path)

    # 2. Create Store Flow
    doc.add_heading('Sequence Diagram: Create Store Flow', level=5)
    create_store_mermaid = """sequenceDiagram
    participant User
    participant SC as StoreController
    participant CSA as CreateStore Action
    participant Storage
    participant DB as Database
    
    User->>SC: POST /stores (StoreData, logo, docs)
    SC->>CSA: execute(StoreData)
    CSA->>DB: Insert new Store record
    CSA-->>SC: Store Instance
    
    SC->>Storage: Upload logo_url & banner_url
    SC->>DB: Update Store with logo paths
    
    SC->>Storage: Upload Verification Documents
    SC->>DB: Insert Store Verifications
    
    SC-->>User: 201 Created (Store details)
"""
    create_store_img_path = "create_store_diagram.png"
    download_image(generate_kroki_url(create_store_mermaid), create_store_img_path)
    doc.add_picture(create_store_img_path, width=Inches(6.0))
    diagrams.append(create_store_img_path)

    # 3. View Public Store Profile Flow
    doc.add_heading('Sequence Diagram: View Public Store Profile Flow', level=5)
    view_store_mermaid = """sequenceDiagram
    participant Client
    participant SC as StoreController
    participant DB as Database
    
    Client->>SC: GET /public/stores/{store}
    SC->>DB: Fetch Store with public relations
    
    alt Store is Active
        SC->>DB: Insert into store_profile_views (IP, store_id, user_id)
        SC-->>Client: 200 OK (Public Store Data)
    else Store is Inactive/Deleted
        SC-->>Client: 404 Not Found
    end
"""
    view_store_img_path = "view_store_diagram.png"
    download_image(generate_kroki_url(view_store_mermaid), view_store_img_path)
    doc.add_picture(view_store_img_path, width=Inches(5.0))
    diagrams.append(view_store_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "Store and Profile Management utilizes specific tables to handle multi-tenant business data alongside user demographics and engagement tracking."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "profiles: Stores extended user details. Fields: id, user_id (FK to users), first_name, last_name, avatar_url, bio.",
        "stores: Main business entity. Fields: id (UUID), owner_user_id (FK to users), name, description, logo_url, banner_url, status, is_verified, rating_avg, followers_count.",
        "store_verifications: Tracks uploaded legal documents. Fields: id, store_id, document_type, document_path, status.",
        "store_profile_views: Analytics table tracking public views. Fields: id, store_id (FK to stores), user_id (nullable FK for logged-in users), ip_address."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--o| PROFILES : "has_one"
    USERS ||--o{ STORES : "owns"
    STORES ||--o{ STORE_VERIFICATIONS : "validates_via"
    STORES ||--o{ STORE_PROFILE_VIEWS : "receives"
    USERS ||--o{ STORE_PROFILE_VIEWS : "generates (optional)"
    
    USERS {
        char id PK "UUID"
        string email
    }
    PROFILES {
        bigint id PK
        uuid user_id FK
        string avatar_url
    }
    STORES {
        uuid id PK
        char owner_user_id FK
        string name
        string logo_url
        string status
    }
    STORE_VERIFICATIONS {
        bigint id PK
        uuid store_id FK
        string document_type
        string status
    }
    STORE_PROFILE_VIEWS {
        bigint id PK
        uuid store_id FK
        char user_id FK "nullable"
        string ip_address
    }
"""
    er_img_path = "store_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(6.0))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Store & Profile Management chapter appended to {filename}")

if __name__ == "__main__":
    append_store_chapter()
