import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def append_db_strategy_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)

    # ================== CHAPTER 8: DATABASE STRATEGY ==================
    doc.add_page_break()
    doc.add_heading('Chapter 8: Database Strategy & Query Optimization', level=1)
    
    intro_text = (
        "As a large-scale e-commerce and loyalty platform, preventing database bottlenecking is the highest priority. "
        "The backend strictly enforces heavy optimization rules, leveraging advanced Eloquent eager loading, pessimistic locking, "
        "and pushing complex mathematical calculations directly to the MySQL engine."
    )
    doc.add_paragraph(intro_text)

    # ------------------ 8.1 Eager Loading ------------------
    doc.add_heading('8.1 Eager Loading & N+1 Prevention', level=2)
    doc.add_paragraph("N+1 queries (where looping over data causes hundreds of hidden database hits) are actively prevented. All Repositories and Services proactively preload massive relationship trees in a single query:")
    eager = [
        "Deep Eager Loading (Banners): The `BannerManagementController` resolves deep hierarchies instantly using `with(['store', 'branches', 'offers.product.images', 'offers.product.variants.attributes'])`.",
        "Closure-Based Loading: The `ProductRepository` uses conditional eager loading `fn ($query) => $query->where('is_active', true)->with('attributes')` to filter massive catalogs in memory before returning them to the controller."
    ]
    for e in eager:
        doc.add_paragraph(e, style='List Bullet')

    # ------------------ 8.2 Database Transactions ------------------
    doc.add_heading('8.2 Transactions & Pessimistic Locking', level=2)
    doc.add_paragraph("To absolutely guarantee data consistency and prevent race conditions (e.g., two users trying to claim the exact same final flash offer simultaneously), the platform uses atomic `DB::transaction()` wrappers with Pessimistic Locking (`lockForUpdate()`):")
    trans = [
        "RedeemOfferClaim Action: Completely locks the `OfferClaim` row, locks the associated `ProductVariant`, decrements the `stock_qty`, and awards loyalty points atomically. If any step fails, the entire transaction rolls back instantly.",
        "PointsService: When users or stores are awarded/deducted points, their balances are locked using `lockForUpdate()` during the transaction to ensure mathematical accuracy during concurrent API hits."
    ]
    for t in trans:
        doc.add_paragraph(t, style='List Bullet')

    # ------------------ 8.3 Complex Raw Queries ------------------
    doc.add_heading('8.3 Offloading Math to MySQL (DB::raw)', level=2)
    doc.add_paragraph("Fetching thousands of rows into PHP just to calculate a value will crash the server. Instead, heavy math is offloaded directly to the MySQL engine using `DB::raw()`:")
    raw = [
        "The Haversine Formula: Geospatial distance for 'Nearby Offers' is calculated natively in SQL using `(6371 * ACOS(LEAST(1, GREATEST(-1, COS(RADIANS(?)) * COS(...) + SIN(...)))))`. The `LEAST/GREATEST` clamping prevents a known MySQL crashing bug with float precision.",
        "The Trending Algorithm: The Explore Service sorts popular items dynamically via SQL using `(products.favorites_count * 1 + {views} * 0.5 + {discount} * 0.2)`."
    ]
    for r in raw:
        doc.add_paragraph(r, style='List Bullet')

    # ------------------ 8.4 Database Indexing ------------------
    doc.add_heading('8.4 Database Indexing', level=2)
    doc.add_paragraph("The database schema applies explicit Composite Indexes (`$table->index()`) to speed up read-heavy operations across the platform:")
    indexes = [
        "Addresses: A composite index on `['latitude', 'longitude']` specifically to speed up the Haversine geospatial calculations.",
        "Products: Composite indexes on `['store_id', 'status']` and `['is_active', 'sort_order']` ensure the store profiles load their active catalog in milliseconds.",
        "Users: Critical indexes on `email`, `phone_number`, and `['provider', 'provider_id']` for instantaneous social-login lookup."
    ]
    for i in indexes:
        doc.add_paragraph(i, style='List Bullet')

    # Save document
    doc.save(filename)
    print(f"Database Strategy Chapter appended to {filename}")

if __name__ == "__main__":
    append_db_strategy_chapter()
