import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_more_points_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    doc.add_heading('Additional Sequence Diagrams for Points System', level=4)
    
    diagrams = []
    
    # 3. Deduct Store Points Flow
    doc.add_heading('Sequence Diagram: Deduct Store Points Flow', level=5)
    deduct_mermaid = """sequenceDiagram
    participant Admin
    participant APC as AdminPointController
    participant PS as PointsService
    participant DB as Database
    
    Admin->>APC: POST /stores/{store}/points/deduct (points=100)
    APC->>PS: deductStorePoints(Store, 100)
    PS->>DB: Begin Transaction
    PS->>DB: UPDATE store_points SET current_balance -= 100, lifetime_spent += 100
    PS->>DB: INSERT INTO store_point_transactions (amount=100, type='debit')
    DB-->>PS: Commit Transaction
    PS-->>APC: Updated StorePoints Model
    APC-->>Admin: 200 OK (New Balance)
"""
    deduct_img_path = "deduct_points_diagram.png"
    download_image(generate_kroki_url(deduct_mermaid), deduct_img_path)
    doc.add_picture(deduct_img_path, width=Inches(5.0))
    diagrams.append(deduct_img_path)

    # 4. Set Points Override Flow
    doc.add_heading('Sequence Diagram: Set Points Override Flow', level=5)
    set_override_mermaid = """sequenceDiagram
    participant Admin
    participant APC as AdminPointController
    participant PS as PointsService
    participant DB as Database
    
    Admin->>APC: POST /users/{user}/points/set (points=500)
    APC->>PS: setUserPoints(User, 500)
    PS->>DB: Fetch current_balance (e.g., 200)
    PS->>PS: Calculate Delta (500 - 200 = +300)
    PS->>DB: Begin Transaction
    PS->>DB: UPDATE user_points SET current_balance = 500
    alt Delta > 0
        PS->>DB: UPDATE lifetime_earned += 300
        PS->>DB: INSERT user_point_transactions (amount=300, type='credit')
    else Delta < 0
        PS->>DB: UPDATE lifetime_spent += |Delta|
        PS->>DB: INSERT user_point_transactions (amount=|Delta|, type='debit')
    end
    DB-->>PS: Commit Transaction
    PS-->>APC: Updated UserPoints Model
    APC-->>Admin: 200 OK
"""
    set_override_img_path = "set_override_diagram.png"
    download_image(generate_kroki_url(set_override_mermaid), set_override_img_path)
    doc.add_picture(set_override_img_path, width=Inches(5.5))
    diagrams.append(set_override_img_path)

    # 5. View Points Balance Flow
    doc.add_heading('Sequence Diagram: View Points Balance Flow', level=5)
    view_balance_mermaid = """sequenceDiagram
    participant Frontend
    participant MPC as MyPointController
    participant PS as PointsService
    participant DB as Database
    
    Frontend->>MPC: GET /me/points
    MPC->>PS: getOrCreateUserPoints(User)
    PS->>DB: Query user_points where user_id = User.id
    alt Record Exists
        DB-->>PS: user_points record
    else Record Does Not Exist
        PS->>DB: Insert default record (balance=0)
        DB-->>PS: new user_points record
    end
    PS-->>MPC: UserPoints Model
    MPC-->>Frontend: 200 OK (current_balance, lifetime_earned, etc)
"""
    view_balance_img_path = "view_balance_diagram.png"
    download_image(generate_kroki_url(view_balance_mermaid), view_balance_img_path)
    doc.add_picture(view_balance_img_path, width=Inches(5.5))
    diagrams.append(view_balance_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Additional Points System diagrams appended to {filename}")

if __name__ == "__main__":
    append_more_points_diagrams()
