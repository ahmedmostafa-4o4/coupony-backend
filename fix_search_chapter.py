import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def fix_search_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # Remove old inaccurate chapter
    found = False
    for p in list(doc.paragraphs):
        if p.text == '1.1.14 Global Search System (Meilisearch)':
            found = True
        if found:
            p._element.getparent().remove(p._element)
            
    diagrams = []

    # ================== SEARCH SYSTEM FEATURE ==================
    doc.add_heading('1.1.14 Native Search & Offers Discovery System', level=3)

    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Search System is designed around native database queries (Eloquent) optimized for precise matching without requiring external full-text engines. "
        "It features custom Arabic Text Normalization, which strips out Arabic diacritics (tashkeel) and unifies variations of the 'Alef' character "
        "before building the query to significantly improve hit rates. "
        "When a user performs a search, the system matches the query against Product titles and descriptions, Store names, and Category names using combined SQL WHERE clauses. "
        "Crucially, the system supports advanced geospatial filtering. By utilizing raw SQL and the Haversine formula against linked store addresses, "
        "users can filter offers by geographical proximity (e.g., 'nearby'). "
        "Users can also seamlessly toggle their 'favorite' status on any discovered offer directly from the search interface."
    )
    doc.add_paragraph(logic_text)

    # ======================== USE CASES ========================
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-108: Search Offers with Filters - User submits a search query along with optional filters (category, min_price, max_price, sort_by, lat/lng). The system normalizes Arabic text, applies the Haversine formula for proximity if coordinates are provided, and returns paginated matching offers along with mock facet data.",
        "UC-109: Toggle Offer Favorite - Authenticated user marks or unmarks a discovered offer (Product) as a favorite. The system securely updates the pivot table and returns the new favorite status."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')

    # ======================== SEQUENCE DIAGRAMS ========================
    doc.add_heading('Sequence Diagrams', level=4)

    # --- UC-108: Search Offers ---
    add_diagram(doc, "Search Offers with Filters & Geolocation (UC-108)", """sequenceDiagram
    participant User
    participant SOC as SearchOffersController
    participant SOS as SearchOfferService
    participant DB as Database

    User->>SOC: GET /search?q=coffee&lat=X&lng=Y
    SOC->>SOC: Validate SearchOffersRequest
    SOC->>SOS: search(query, filters, lat/lng)
    SOS->>SOS: normalizeArabicText("coffee")
    SOS->>DB: Build Base Query (Products where ACTIVE)
    Note over DB: Apply LIKE clauses (Product title/desc OR Store name OR Category name)
    alt If Lat/Lng Provided
        SOS->>DB: Apply Haversine formula on stores->addressables->addresses
        SOS->>DB: Filter where distance <= threshold
    end
    SOS->>DB: Apply Price / Category / Sorting filters
    DB-->>SOS: Paginated Product Matches
    SOS-->>SOC: Paginator
    SOC-->>User: 200 OK (SearchOfferResource + Facets)
""", "sd_search_offers", diagrams, 6.0)

    # --- UC-109: Toggle Offer Favorite ---
    add_diagram(doc, "Toggle Offer Favorite (UC-109)", """sequenceDiagram
    participant User
    participant SOC as SearchOffersController
    participant TFA as ToggleOfferFavoriteAction
    participant DB as Database

    User->>SOC: POST /search/offers/{offerId}/favorite
    SOC->>TFA: execute(Offer/Product, User)
    TFA->>DB: Check if favorite exists in product_favorites
    alt Favorite Exists
        TFA->>DB: Delete from product_favorites
        TFA->>DB: Decrement Product.favorites_count
        TFA-->>SOC: is_favorite = false
    else Favorite Does Not Exist
        TFA->>DB: Insert into product_favorites
        TFA->>DB: Increment Product.favorites_count
        TFA-->>SOC: is_favorite = true
    end
    SOC-->>User: 200 OK (Favorite Status)
""", "sd_toggle_favorite", diagrams, 5.5)

    # ======================== DATABASE SCHEMA ========================
    doc.add_heading('Database Schema', level=4)
    doc.add_paragraph(
        "The native search engine relies dynamically on existing platform tables and does not require dedicated indexing or history tables."
    )
    schemas = [
        "products: Matched against 'title' and 'description'. Uses base_price/price for range filtering.",
        "stores: Matched against 'name'. Provides relationship to physical addresses.",
        "categories: Matched against 'name_en' and 'name_ar'. Used for categorization filtering.",
        "addresses & addressables: Queried using raw SQL (Haversine formula) utilizing the 'latitude' and 'longitude' columns to calculate physical proximity to the user.",
        "product_favorites: Pivot table enabling users to bookmark discovered offers."
    ]
    for s in schemas:
        doc.add_paragraph(s, style='List Bullet')

    # ======================== ER DIAGRAM ========================
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    PRODUCTS ||--o{ CATEGORIES : "categorized_by"
    PRODUCTS ||--|| STORES : "sold_by"
    STORES ||--o{ ADDRESSABLES : "morphs_to"
    ADDRESSABLES ||--|| ADDRESSES : "maps_to"
    PRODUCTS ||--o{ PRODUCT_FAVORITES : "favorited_in"
    USERS ||--o{ PRODUCT_FAVORITES : "favorites"

    PRODUCTS {
        uuid id PK
        string title
        text description
        decimal price
    }
    STORES {
        uuid id PK
        string name
    }
    CATEGORIES {
        bigint id PK
        string name_en
        string name_ar
    }
    ADDRESSES {
        uuid id PK
        decimal latitude
        decimal longitude
    }
    PRODUCT_FAVORITES {
        uuid id PK
        uuid product_id FK
        char user_id FK
    }
"""
    er_img_path = "search_er_diagram_fixed.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Search System chapter fixed and appended to {filename}")

if __name__ == "__main__":
    fix_search_chapter()
