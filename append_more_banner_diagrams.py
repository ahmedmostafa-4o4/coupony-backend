import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def add_diagram(doc, title, mermaid_text, img_name, diagrams, width=5.5):
    doc.add_heading(f'Sequence Diagram: {title}', level=5)
    img_path = f"{img_name}.png"
    download_image(generate_kroki_url(mermaid_text), img_path)
    doc.add_picture(img_path, width=Inches(width))
    diagrams.append(img_path)

def append_missing_banner_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    diagrams = []

    doc.add_heading('Additional Banners System Sequence Diagrams', level=4)

    # --- UC-115: Manage Store Banners ---
    add_diagram(doc, "Manage Store Banners (UC-115)", """sequenceDiagram
    participant Owner as Store Owner
    participant SBC as StoreBannerController
    participant BS as BannerService
    participant DB as Database

    Owner->>SBC: POST /stores/{store}/banners (data, offers, branches)
    SBC->>BS: create()
    BS->>DB: INSERT INTO banners (status = 'pending')
    BS->>DB: Attach to banner_offers & banner_branches
    BS-->>SBC: Banner Model
    SBC-->>Owner: 201 Created (Pending)
    
    Owner->>SBC: PUT /stores/{store}/banners/{id}
    SBC->>BS: update()
    BS->>DB: UPDATE fields, reset status = 'pending'
    BS-->>SBC: Banner Model
    SBC-->>Owner: 200 OK
""", "sd_manage_store_banners_separated", diagrams, 5.5)

    # --- UC-116: Admin Review Banner ---
    add_diagram(doc, "Admin Review Banner (UC-116)", """sequenceDiagram
    participant Admin
    participant ABC as AdminBannerController
    participant BS as BannerService
    participant DB as Database

    Admin->>ABC: POST /admin/banners/{id}/approve
    ABC->>BS: approve()
    BS->>DB: UPDATE banners SET status='approved', is_active=true
    
    loop Every Linked Offer
        BS->>DB: Auto-approve pending product revisions
        BS->>DB: UPDATE product & offer SET is_active=true
    end
    
    BS-->>ABC: Success
    ABC-->>Admin: 200 OK
    
    Admin->>ABC: POST /admin/banners/{id}/reject (reason)
    ABC->>BS: reject(reason)
    BS->>DB: UPDATE banners SET status='rejected', rejection_reason=X
    BS-->>ABC: Success
    ABC-->>Admin: 200 OK
""", "sd_admin_review_banner_separated", diagrams, 6.0)

    # --- UC-117: Manage Travel Banners ---
    add_diagram(doc, "Manage Travel Banners (UC-117)", """sequenceDiagram
    participant Admin
    participant TBC as Admin\TravelBannerController
    participant DB as Database

    Admin->>TBC: POST /admin/travel-banners (product_id, image, cta)
    TBC->>TBC: Validate Product is Active
    TBC->>DB: INSERT INTO travel_banners
    TBC-->>Admin: 201 Created
    
    Admin->>TBC: PUT /admin/travel-banners/{id}
    TBC->>DB: UPDATE travel_banners fields
    TBC-->>Admin: 200 OK
    
    Admin->>TBC: DELETE /admin/travel-banners/{id}
    TBC->>DB: DELETE FROM travel_banners
    TBC-->>Admin: 200 OK
""", "sd_manage_travel_banners", diagrams, 5.5)

    # --- UC-119: Interact with Banner ---
    add_diagram(doc, "Interact with Banner (UC-119)", """sequenceDiagram
    participant Customer
    participant CBC as CustomerBannerController
    participant IBA as InteractWithBannerAction
    participant DB as Database

    Customer->>CBC: POST /banners/{id}/{action} (like/favorite/share)
    CBC->>IBA: execute(Banner, action, User)
    
    IBA->>DB: Begin Transaction
    IBA->>DB: Check Interaction Table (e.g., banner_likes) for User
    alt Already Interacted
        IBA-->>CBC: Skip or Un-toggle
        CBC-->>Customer: 200 OK (Toggled Off)
    else New Interaction
        IBA->>DB: INSERT INTO banner_likes (banner_id, user_id)
        IBA->>DB: Increment Banner.likes_count
        DB-->>IBA: Commit Transaction
        IBA-->>CBC: Success
        CBC-->>Customer: 200 OK (Toggled On)
    end
""", "sd_interact_banner_separated", diagrams, 6.0)

    # Save document
    doc.save(filename)

    # Cleanup
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)

    print(f"Missing Banner System diagrams appended to {filename}")

if __name__ == "__main__":
    append_missing_banner_diagrams()
