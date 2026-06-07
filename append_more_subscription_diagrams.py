import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_more_subscription_diagrams():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    doc.add_heading('Additional Sequence Diagrams for Subscriptions System', level=4)
    
    diagrams = []
    
    # 4. View Paginated Subscriptions List Flow
    doc.add_heading('Sequence Diagram: View Paginated Subscriptions List Flow', level=5)
    view_list_mermaid = """sequenceDiagram
    participant Admin
    participant SSMC as StoreSubscriptionManagementController
    participant DB as Database
    
    Admin->>SSMC: GET /subscriptions?status=ACTIVE&per_page=15
    SSMC->>DB: Query Subscriptions
    SSMC->>DB: Apply Filters (status, store_id, plan_id)
    SSMC->>DB: Eager Load -> store, plan
    SSMC->>DB: Order by updated_at DESC (latest)
    DB-->>SSMC: Paginated Collection
    SSMC-->>Admin: 200 OK (List of Subscriptions + Meta)
"""
    view_list_img_path = "view_sub_list_diagram.png"
    download_image(generate_kroki_url(view_list_mermaid), view_list_img_path)
    doc.add_picture(view_list_img_path, width=Inches(5.0))
    diagrams.append(view_list_img_path)

    # 5. View Subscription Details Flow
    doc.add_heading('Sequence Diagram: View Subscription Details Flow', level=5)
    view_details_mermaid = """sequenceDiagram
    participant Admin
    participant SSMC as StoreSubscriptionManagementController
    participant DB as Database
    
    Admin->>SSMC: GET /subscriptions/{id}
    SSMC->>DB: Query Subscription by ID
    SSMC->>DB: Eager Load -> store, plan, auditLogs
    alt Record Found
        DB-->>SSMC: Subscription Record with Logs
        SSMC-->>Admin: 200 OK (Detailed Snapshot)
    else Record Not Found
        SSMC-->>Admin: 404 Not Found
    end
"""
    view_details_img_path = "view_sub_details_diagram.png"
    download_image(generate_kroki_url(view_details_mermaid), view_details_img_path)
    doc.add_picture(view_details_img_path, width=Inches(5.5))
    diagrams.append(view_details_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Additional Subscriptions System diagrams appended to {filename}")

if __name__ == "__main__":
    append_more_subscription_diagrams()
