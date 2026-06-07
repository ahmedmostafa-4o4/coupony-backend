import os
import zlib
import base64
import urllib.request
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    print("python-docx not installed.")
    sys.exit(1)

def generate_kroki_url(diagram_text, diagram_type='mermaid', output_format='png'):
    compressed = zlib.compress(diagram_text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    return f"https://kroki.io/{diagram_type}/{output_format}/{encoded}"

def download_image(url, path):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
        out_file.write(response.read())

def append_points_chapter():
    filename = "Features_Documentation_Diagrams.docx"
    doc = Document(filename)
    
    # ================== POINTS SYSTEM FEATURE ==================
    doc.add_heading('1.1.8 Points System', level=3)
    
    doc.add_heading('Description of Logic', level=4)
    logic_text = (
        "The Points System functions as a dual-ledger digital wallet for both Users and Stores. "
        "It maintains an aggregated balance (tracking current_balance, lifetime_earned, and lifetime_spent) while strictly recording every change in an immutable transaction ledger. "
        "A dedicated 'PointsService' acts as the core domain service, ensuring atomic operations when awarding, deducting, or manually setting points. "
        "Administrators have endpoints to manually adjust balances with specific reasons and notes, creating an auditable trail. Users and Stores can retrieve their own balances and a paginated history of their transactions."
    )
    doc.add_paragraph(logic_text)
    
    doc.add_heading('Use Cases', level=4)
    use_cases = [
        "UC-46: Retrieve Points Balance - A User or Store retrieves their current aggregated points balance and lifetime statistics.",
        "UC-47: Add Points - The system automatically, or an Admin manually, awards points to a User or Store (e.g., for purchases, engagement, or compensation).",
        "UC-48: Deduct Points - The system deducts points when a User redeems a reward, updating both the current_balance and lifetime_spent.",
        "UC-49: View Point Transactions - A User or Store fetches a paginated, historical ledger of all credit and debit point transactions.",
        "UC-50: Set Points Override - An Administrator manually overrides a balance to an exact value, automatically generating the necessary adjustment transaction to balance the ledger."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Bullet')
        
    doc.add_heading('Sequence Diagrams', level=4)
    
    diagrams = []
    
    # 1. Add User Points Flow
    doc.add_heading('Sequence Diagram: Add User Points Flow', level=5)
    add_points_mermaid = """sequenceDiagram
    participant Admin
    participant APC as AdminPointController
    participant PS as PointsService
    participant DB as Database
    
    Admin->>APC: POST /users/{user}/points/add (points=50, reason="Reward")
    APC->>PS: addUserPoints(User, 50, reason)
    PS->>DB: Begin Transaction
    PS->>DB: UPDATE user_points SET current_balance += 50, lifetime_earned += 50
    PS->>DB: INSERT INTO user_point_transactions (amount=50, type='credit', reason)
    DB-->>PS: Commit Transaction
    PS-->>APC: Updated UserPoints Model
    APC-->>Admin: 200 OK (New Balance)
"""
    add_points_img_path = "add_points_diagram.png"
    download_image(generate_kroki_url(add_points_mermaid), add_points_img_path)
    doc.add_picture(add_points_img_path, width=Inches(5.5))
    diagrams.append(add_points_img_path)

    # 2. View Transactions Ledger Flow
    doc.add_heading('Sequence Diagram: View Transactions Ledger Flow', level=5)
    view_transactions_mermaid = """sequenceDiagram
    participant User
    participant APC as AdminPointController
    participant DB as Database
    
    User->>APC: GET /users/{user}/points/transactions
    APC->>DB: Query user_point_transactions
    APC->>DB: Filter where user_id = User.id
    APC->>DB: Order by created_at DESC (latest)
    DB-->>APC: Paginated Collection
    APC-->>User: 200 OK (Paginated Ledger Data)
"""
    view_transactions_img_path = "view_transactions_diagram.png"
    download_image(generate_kroki_url(view_transactions_mermaid), view_transactions_img_path)
    doc.add_picture(view_transactions_img_path, width=Inches(5.0))
    diagrams.append(view_transactions_img_path)

    # Database Schema
    doc.add_heading('Database Schema', level=4)
    db_text = (
        "The schema relies on a fast-read balance table linked to an append-only transaction ledger table, duplicated for both Users and Stores to prevent data mixing."
    )
    doc.add_paragraph(db_text)
    
    schemas = [
        "user_points / store_points: Stores the live balance. Fields: id, user_id/store_id (FK), current_balance (int), lifetime_earned (int), lifetime_spent (int).",
        "user_point_transactions / store_point_transactions: The immutable ledger. Fields: id, user_id/store_id (FK), amount (int), type (enum: credit/debit), reason (string), reference_id (nullable, for linking to orders/actions), note, created_at."
    ]
    for schema in schemas:
        doc.add_paragraph(schema, style='List Bullet')

    # ER Diagram
    doc.add_heading('Entity-Relationship (ER) Diagram', level=5)
    er_mermaid = """erDiagram
    USERS ||--|| USER_POINTS : "has_balance"
    USERS ||--o{ USER_POINT_TRANSACTIONS : "has_ledger_entries"
    STORES ||--|| STORE_POINTS : "has_balance"
    STORES ||--o{ STORE_POINT_TRANSACTIONS : "has_ledger_entries"
    
    USER_POINTS {
        bigint id PK
        char user_id FK
        int current_balance
        int lifetime_earned
        int lifetime_spent
    }
    USER_POINT_TRANSACTIONS {
        bigint id PK
        char user_id FK
        int amount
        string type "credit/debit"
        string reason
        string reference_id "Nullable"
    }
"""
    er_img_path = "points_er_diagram.png"
    download_image(generate_kroki_url(er_mermaid), er_img_path)
    doc.add_picture(er_img_path, width=Inches(5.5))
    diagrams.append(er_img_path)

    # Save document
    doc.save(filename)
    
    # Cleanup downloaded images
    for img in diagrams:
        if os.path.exists(img):
            os.remove(img)
            
    print(f"Points System chapter appended to {filename}")

if __name__ == "__main__":
    append_points_chapter()
